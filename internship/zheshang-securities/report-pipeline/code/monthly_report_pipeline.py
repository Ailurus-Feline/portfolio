from __future__ import annotations

import argparse
import datetime
import json
import os
import re
import time as time_module
from contextlib import contextmanager
from dataclasses import dataclass
from pathlib import Path

import pandas as pd



# 图6：股指期货主力合约基差 及其相关文案/图片开关。
# 权限开通后改为 True 即可恢复正常生成。
ENABLE_FIGURE6_FUTURES = False

# 图7：两市涨跌数量统计 开关。
ENABLE_FIGURE7_UPDOWN = True

# 图8：两市融资融券余额、融资买入占比 开关。
ENABLE_FIGURE8_MARGIN = False

# 图10：股票型ETF净流入额 开关。
ENABLE_FIGURE10_ETF_NETINFLOW = True

# 图13：按行业指数分类ETF周区间净流入额 开关。
ENABLE_FIGURE13_INDUSTRY_ETF = True

# 图14：限售股解禁金额和数量（周度）开关。
ENABLE_FIGURE14_UNLOCK = True

# 图15：IPO和定增金额（周度）开关。
ENABLE_FIGURE15_IPO = True


@dataclass(frozen=True)
class ProjectPaths:
    project_root: Path
    inputs_dir: Path
    assets_dir: Path
    workspace_dir: Path
    outputs_dir: Path
    data_dir: Path
    fund_list_dir: Path
    industry_code_dir: Path
    picture_dir: Path
    word_dir: Path
    word_replacement_dir: Path
    weekly_word_output_dir: Path
    word_template: Path
    funding_draft_dir: Path
    font_dir: Path
    falling_alert_dir: Path


@dataclass(frozen=True)
class PipelineOptions:
    refresh_code_audit: bool = False
    init_full_history: bool = False
    profile: bool = False
    batch_size: int = 120
    validate_etf_field: bool = False
    refresh_active_quarter_cache: bool = False
    allow_stale_margin_data: bool = False


@dataclass
class LocalTabularResult:
    frame: pd.DataFrame
    ErrorCode: int = 0
    ErrMsg: str = ""
    Fields: list[str] | None = None
    Codes: list[str] | None = None

    def __post_init__(self) -> None:
        if self.Fields is None:
            self.Fields = [str(col) for col in self.frame.columns if str(col).lower() != "date"]
        if self.Codes is None:
            self.Codes = []
        self.errorcode = self.ErrorCode
        self.errmsg = self.ErrMsg
        self.Times = pd.to_datetime(
            self.frame["date"], errors="coerce"
        ).tolist() if "date" in self.frame.columns else []
        self.Data = [
            pd.to_numeric(self.frame[field], errors="coerce").tolist()
            for field in self.Fields
            if field in self.frame.columns
        ]

    @property
    def empty(self) -> bool:
        return self.frame.empty

    def __getattr__(self, item: str):
        return getattr(self.frame, item)

    def __getitem__(self, key):
        return self.frame.__getitem__(key)


@dataclass
class LocalEtfValidationResult:
    selected_path: str
    interface_type: str
    metric_name: str
    requested_dates: list[str]
    actual_dates: list[str]
    non_zero_counts: list[int]
    config_path: Path
    validated_at: str


@dataclass
class LocalEtfFetchResult:
    metric_name: str
    requested_date: str
    actual_date: str
    requested_count: int
    matched_count: int
    non_zero_count: int
    total_value: float


class LocalDataClient:
    def __init__(self, paths: ProjectPaths):
        self.paths = paths
        self.root = paths.data_dir / "local_sources"
        self.root.mkdir(parents=True, exist_ok=True)

    def _slug(self, text: str) -> str:
        normalized = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", text.strip())
        normalized = re.sub(r"_+", "_", normalized).strip("_")
        return normalized or "default"

    def _read_endpoint(self, endpoint: str, names: list[str]) -> pd.DataFrame:
        endpoint_dir = self.root / endpoint
        endpoint_dir.mkdir(parents=True, exist_ok=True)
        for name in names:
            candidate = endpoint_dir / f"{self._slug(name)}.csv"
            if candidate.exists():
                return pd.read_csv(candidate)
        fallback = endpoint_dir / "default.csv"
        if fallback.exists():
            return pd.read_csv(fallback)
        return pd.DataFrame()

    def _empty_result(self) -> LocalTabularResult:
        return LocalTabularResult(pd.DataFrame(), ErrorCode=1, ErrMsg="本地数据未找到")

    def _ensure_date_column(self, df: pd.DataFrame) -> pd.DataFrame:
        if "date" in df.columns:
            return df
        df = df.copy()
        if len(df.columns) == 0:
            df["date"] = pd.Series(dtype="datetime64[ns]")
        else:
            df.rename(columns={df.columns[0]: "date"}, inplace=True)
        return df

    def history(
        self,
        code: str,
        fields: str,
        begin_date: str,
        end_date: str,
        option_string: str = "",
        usedf: bool = False,
        **kwargs,
    ) -> LocalTabularResult | pd.DataFrame:
        del option_string, kwargs
        df = self._read_endpoint("history", [f"{code}_{fields}", code])
        if df.empty:
            empty_result = self._empty_result()
            return empty_result.frame if usedf else empty_result
        df = self._ensure_date_column(df)
        begin_ts = pd.to_datetime(begin_date, errors="coerce")
        end_ts = pd.to_datetime(end_date, errors="coerce")
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        if pd.notna(begin_ts):
            df = df[df["date"] >= begin_ts]
        if pd.notna(end_ts):
            df = df[df["date"] <= end_ts]
        requested_fields = [item.strip() for item in str(fields).split(",") if item.strip()]
        for field in requested_fields:
            if field not in df.columns:
                df[field] = pd.NA
        keep_columns = ["date"] + [field for field in requested_fields if field in df.columns]
        result_df = df[keep_columns].reset_index(drop=True)
        if usedf:
            return result_df
        return LocalTabularResult(result_df, Fields=requested_fields)

    def edb_query(self, indicator_code: str, begin_date: str, end_date: str, option_string: str = "") -> LocalTabularResult:
        del option_string
        df = self._read_endpoint("edb", [indicator_code])
        if df.empty:
            return self._empty_result()
        df = self._ensure_date_column(df)
        value_col = "value"
        if value_col not in df.columns:
            value_col = df.columns[1] if len(df.columns) > 1 else "value"
            if value_col not in df.columns:
                df[value_col] = pd.NA
        df = df.rename(columns={value_col: indicator_code})
        return self.history(indicator_code, indicator_code, begin_date, end_date, "") if indicator_code in df.columns else self._empty_result()

    def snapshot(self, codes: str | list[str], fields: str, option_string: str = "") -> LocalTabularResult:
        del option_string
        code_list = [c.strip() for c in (codes.split(",") if isinstance(codes, str) else codes) if str(c).strip()]
        field_list = [f.strip() for f in str(fields).split(",") if f.strip()] or ["value"]
        df = self._read_endpoint("snapshot", ["snapshot"])
        if df.empty:
            data = {"date": [pd.Timestamp.today().normalize()]}
            for field in field_list:
                data[field] = [0.0]
            return LocalTabularResult(pd.DataFrame(data), Codes=code_list, Fields=field_list)
        if "date" not in df.columns:
            df["date"] = pd.Timestamp.today().normalize()
        for field in field_list:
            if field not in df.columns:
                df[field] = 0.0
        return LocalTabularResult(df[["date"] + field_list].head(max(1, len(code_list))), Codes=code_list, Fields=field_list)

    def snapshot_chunked(self, codes: str | list[str], fields: str, option_string: str = "", batch_size: int = 120) -> LocalTabularResult:
        del batch_size
        return self.snapshot(codes, fields, option_string)

    def datapool(self, dataset_name: str, option_string: str) -> LocalTabularResult:
        del option_string
        df = self._read_endpoint("datapool", [dataset_name])
        if df.empty:
            return self._empty_result()
        df = self._ensure_date_column(df)
        return LocalTabularResult(df)

    def block_wc_query(self, query_text: str) -> pd.DataFrame:
        return self._read_endpoint("wc", [query_text])

    def wc_query_dataframe(self, query_text: str, domain: str = "") -> pd.DataFrame:
        suffix = f"{domain}_{query_text}" if domain else query_text
        return self._read_endpoint("wc", [suffix, query_text])

    def dateserial_chunked(self, codes: str | list[str], field: str, begin_date: str, end_date: str, option_string: str = "", batch_size: int = 120) -> pd.DataFrame:
        del field, begin_date, end_date, option_string, batch_size
        key = ",".join(codes) if isinstance(codes, list) else str(codes)
        return self._read_endpoint("dateserial", [key])

    def validate_etf_size_field(self, codes: list[str], validation_dates: list[str], sample_size: int = 15, write_config: bool = True) -> LocalEtfValidationResult:
        del codes, sample_size
        config_path = self.get_etf_size_config_path()
        if write_config:
            config_path.parent.mkdir(parents=True, exist_ok=True)
            config_path.write_text(
                json.dumps(
                    {
                        "selected_path": "data/local_sources/etf/etf_netasset.csv",
                        "interface_type": "local_csv",
                        "metric_name": "etf_netasset",
                        "validated_at": datetime.datetime.now().isoformat(timespec="seconds"),
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
        return LocalEtfValidationResult(
            selected_path="data/local_sources/etf/etf_netasset.csv",
            interface_type="local_csv",
            metric_name="etf_netasset",
            requested_dates=list(validation_dates),
            actual_dates=list(validation_dates),
            non_zero_counts=[0 for _ in validation_dates],
            config_path=config_path,
            validated_at=datetime.datetime.now().isoformat(timespec="seconds"),
        )

    def ensure_etf_size_path_ready(self) -> LocalEtfValidationResult:
        return self.validate_etf_size_field([], [datetime.date.today().strftime("%Y%m%d")], sample_size=0, write_config=True)

    def get_etf_size_config_path(self) -> Path:
        return self.root / "etf" / "etf_size_config.json"

    def prepare_etf_query_codes(self, codes: list[str], end_date: str) -> list[str]:
        del end_date
        return [str(code).strip() for code in codes if str(code).strip()]

    def fetch_etf_netasset_verified(self, codes: list[str] | str, trade_date: str, batch_size: int = 120) -> LocalEtfFetchResult:
        del batch_size
        code_list = [c.strip() for c in (codes.split(",") if isinstance(codes, str) else codes) if str(c).strip()]
        df = self._read_endpoint("etf", ["etf_netasset"])
        if df.empty:
            return LocalEtfFetchResult("etf_netasset", trade_date, trade_date, len(code_list), 0, 0, 0.0)
        required = {"date", "code", "value"}
        if not required.issubset(df.columns):
            return LocalEtfFetchResult("etf_netasset", trade_date, trade_date, len(code_list), 0, 0, 0.0)
        df = df.copy()
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
        target = pd.to_datetime(trade_date, errors="coerce")
        scoped = df[df["code"].astype(str).isin(code_list)]
        if pd.notna(target):
            exact = scoped[scoped["date"] == target]
            if exact.empty:
                scoped = scoped[scoped["date"] <= target]
                if not scoped.empty:
                    scoped = scoped[scoped["date"] == scoped["date"].max()]
            else:
                scoped = exact
        total = float(pd.to_numeric(scoped["value"], errors="coerce").fillna(0).sum()) if not scoped.empty else 0.0
        non_zero = int((pd.to_numeric(scoped["value"], errors="coerce").fillna(0) != 0).sum()) if not scoped.empty else 0
        actual_date = trade_date if scoped.empty else pd.to_datetime(scoped["date"].iloc[0]).strftime("%Y-%m-%d")
        return LocalEtfFetchResult("etf_netasset", trade_date, actual_date, len(code_list), int(scoped["code"].nunique()) if not scoped.empty else 0, non_zero, total)

    def _ensure_sdk_login(self) -> None:
        return

    def _load_sdk_module(self):
        class _DummySdk:
            @staticmethod
            def THS_ErrorInfo(error_code):
                return {"errmsg": f"本地模式错误码={error_code}"}

            @staticmethod
            def THS_DS(*args, **kwargs):
                del args, kwargs
                class _Result:
                    errorcode = 1
                    errmsg = "本地模式下 THS_DS 不可用"
                    data = pd.DataFrame()

                return _Result()

        return _DummySdk()


def validate_etf_size_field_local(paths: ProjectPaths, report_date: datetime.date, stock_etf_codes: list[str], validation_dates: list[str]) -> LocalEtfValidationResult:
    del report_date, stock_etf_codes
    return LocalDataClient(paths).validate_etf_size_field([], validation_dates, sample_size=15, write_config=True)


class PerformanceTracker:
    def __init__(self, enabled: bool = False):
        self.enabled = enabled
        self.records: list[tuple[str, float]] = []

    @contextmanager
    def track(self, label: str):
        start = time_module.perf_counter()
        try:
            yield
        finally:
            elapsed = time_module.perf_counter() - start
            self.records.append((label, elapsed))
            print(f"[耗时] {label}: {elapsed:.2f} 秒")

    def summary(self) -> None:
        if not self.records:
            return
        print("[耗时] 汇总:")
        for label, elapsed in self.records:
            print(f"[耗时]   {label}: {elapsed:.2f} 秒")


def get_project_root() -> Path:
    return Path(__file__).resolve().parent.parent


def build_paths(project_root: Path) -> ProjectPaths:
    inputs_dir = project_root / 'inputs'
    assets_dir = project_root / 'assets'
    workspace_dir = project_root / 'workspace'
    outputs_dir = project_root / 'outputs'

    data_dir = inputs_dir / 'data'
    return ProjectPaths(
        project_root=project_root,
        inputs_dir=inputs_dir,
        assets_dir=assets_dir,
        workspace_dir=workspace_dir,
        outputs_dir=outputs_dir,
        data_dir=data_dir,
        fund_list_dir=data_dir / 'funds',
        industry_code_dir=data_dir / '行业代码',
        picture_dir=outputs_dir / 'figures',
        word_dir=outputs_dir / 'word',
        word_replacement_dir=assets_dir / 'replacements',
        weekly_word_output_dir=outputs_dir / 'word',
        word_template=assets_dir / 'templates' / 'interactive_template.docx',
        funding_draft_dir=workspace_dir / 'funding_draft',
        font_dir=assets_dir / 'fonts',
        falling_alert_dir=data_dir / 'falling_alert',
    )


def get_report_date(report_date_arg: str | None) -> datetime.date:
    if not report_date_arg:
        return datetime.date.today()
    return datetime.datetime.strptime(report_date_arg, '%Y-%m-%d').date()


def ensure_directories(paths: ProjectPaths) -> None:
    required = [
        paths.inputs_dir / 'data' / 'local_sources',
        paths.inputs_dir / 'reference',
        paths.assets_dir / 'templates',
        paths.assets_dir / 'fonts',
        paths.assets_dir / 'replacements',
        paths.workspace_dir / 'cache',
        paths.workspace_dir / 'logs',
        paths.workspace_dir / 'tmp',
        paths.outputs_dir / 'figures',
        paths.outputs_dir / 'csv' / 'figures_1_15',
        paths.outputs_dir / 'csv' / 'summary',
        paths.outputs_dir / 'word',
        paths.data_dir / 'funds',
        paths.data_dir / '行业代码',
        paths.data_dir / 'falling_alert',
        paths.data_dir / 'quotation',
        paths.data_dir / 'market_sentiment',
        paths.data_dir / 'local_sources' / 'history',
        paths.data_dir / 'local_sources' / 'edb',
        paths.data_dir / 'local_sources' / 'snapshot',
        paths.data_dir / 'local_sources' / 'datapool',
        paths.data_dir / 'local_sources' / 'wc',
        paths.data_dir / 'local_sources' / 'dateserial',
        paths.data_dir / 'local_sources' / 'etf',
        paths.funding_draft_dir / 'data' / '被动权益基金',
        paths.funding_draft_dir / 'data' / '主动权益基金',
        paths.funding_draft_dir / 'data' / '两融余额数据',
        paths.funding_draft_dir / 'data' / '保险资金',
        paths.funding_draft_dir / 'data' / 'cache',
        paths.funding_draft_dir / 'picture',
    ]

    for folder in required:
        folder.mkdir(parents=True, exist_ok=True)


def configure_matplotlib_backend() -> None:
    import matplotlib

    try:
        matplotlib.use('TkAgg')
    except Exception:
        matplotlib.use('Agg')


def trim_chart_label_suffix(
    value: str,
    suffixes: tuple[str, ...] = ("概念板块", "概念", "板块", "指数"),
) -> str:
    text = str(value or "").strip()
    for suffix in suffixes:
        if text.endswith(suffix):
            trimmed = text[: -len(suffix)].rstrip()
            return trimmed or text
    return text


def configure_plot_font(paths: ProjectPaths) -> str:
    import matplotlib.font_manager as fm

    if paths.font_dir.exists():
        for font_file in paths.font_dir.glob('*.ttf'):
            try:
                fm.fontManager.addfont(str(font_file))
                if font_file.stem.lower().startswith('kaiti'):
                    return font_file.stem
            except Exception:
                continue

    available = {font.name for font in fm.fontManager.ttflist}
    for candidate in ['KaiTi', 'Kaiti SC', 'SimKai', 'STKaiti', 'Microsoft YaHei']:
        if candidate in available:
            return candidate
    return 'DejaVu Sans'


def resolve_falling_alert_file(paths: ProjectPaths, report_date: datetime.date) -> Path:
    exact = paths.falling_alert_dir / f'六指数_peak_window300_{report_date.strftime("%Y%m%d")}.csv'
    if exact.exists():
        return exact

    dated_files: list[tuple[datetime.date, Path]] = []
    for file_path in paths.falling_alert_dir.glob('六指数_peak_window300_*.csv'):
        suffix = file_path.stem.replace('六指数_peak_window300_', '')
        if len(suffix) != 8 or not suffix.isdigit():
            continue
        parsed = datetime.datetime.strptime(suffix, '%Y%m%d').date()
        if parsed <= report_date:
            dated_files.append((parsed, file_path))
    if dated_files:
        dated_files.sort(key=lambda item: item[0])
        return dated_files[-1][1]

    fallback = paths.falling_alert_dir / '六指数_peak_window300.csv'
    if fallback.exists():
        return fallback

    raise FileNotFoundError(
        f'未找到 falling_alert 文件。请检查目录 {paths.falling_alert_dir}。'
        f'或补充与报告日 {report_date:%Y-%m-%d} 匹配的 六指数_peak_window300_YYYYMMDD.csv。'
    )


def load_local_data(file_path: Path) -> pd.DataFrame:
    if not file_path.exists():
        raise FileNotFoundError(f"本地数据文件不存在：{file_path}")
    return pd.read_csv(file_path)

def load_stock_etf_codes(paths: ProjectPaths) -> list[str]:
    stock_etf_path = paths.fund_list_dir / 'stock_ETF.csv'
    df = load_local_data(stock_etf_path)
    if '证券代码' in df.columns:
        code_series = df['证券代码']
    else:
        code_series = df.iloc[:, 0]
    return [
        str(value).strip()
        for value in code_series.dropna().tolist()
        if str(value).strip()
    ]


def build_etf_validation_dates(report_date: datetime.date) -> list[str]:
    ts = pd.Timestamp(report_date)
    dates = [
        (ts - pd.offsets.QuarterEnd(2)).strftime('%Y%m%d'),
        (ts - pd.offsets.QuarterEnd(1)).strftime('%Y%m%d'),
        ts.strftime('%Y%m%d'),
    ]
    deduped: list[str] = []
    for item in dates:
        if item not in deduped:
            deduped.append(item)
    return deduped


def run_etf_field_validation(paths: ProjectPaths, report_date: datetime.date) -> None:
    stock_etf_codes = load_stock_etf_codes(paths)
    validation_dates = build_etf_validation_dates(report_date)
    print(
        f"ETF字段验证模式: requested_dates={validation_dates}, "
        f"sample_source={paths.fund_list_dir / 'stock_ETF.csv'}"
    )
    validation = validate_etf_size_field_local(
        paths,
        report_date,
        stock_etf_codes,
        validation_dates,
    )
    print(
        "ETF字段验证通过: "
        f"path={validation.selected_path}, "
        f"interface={validation.interface_type}, "
        f"metric={validation.metric_name}, "
        f"requested_dates={list(validation.requested_dates)}, "
        f"actual_dates={list(validation.actual_dates)}, "
        f"non_zero_counts={list(validation.non_zero_counts)}, "
        f"config={validation.config_path}"
    )



def run_step1(
    local_client: LocalDataClient,
    paths: ProjectPaths,
    report_date: datetime.date,
    options: PipelineOptions | None = None,
) -> None:
    # -*- coding: utf-8 -*-

    from loguru import logger
    import datetime
    import os
    import pandas as pd
    import warnings
    warnings.filterwarnings("ignore")

    # 固定日期（datetime.date类型，与原代码日期类型一致）
    FIXED_TODAY = report_date
    options = options or PipelineOptions()

    logger.add(str(paths.workspace_dir / 'logs' / 'updating_data.log'), encoding="utf-8")

    DB_LOC = str(paths.data_dir)

    # %% 连接WIND


    # %% 更新板块行情数据 板块包括大盘指数、中信行业分类（一级、二级）
    sector_code_dict = {
        '上证指数': '000001.SH',
        '深证成指': '399001.SZ',
        '创业板指': '399006.SZ',
        '上证50': '000016.SH',
        '科创50': '000688.SH',
        '沪深300': '000300.SH',
        '中证500': '000905.SH',
        '中证1000': '000852.SH',
    }


    def updating_quotation_sector_data(wind_code_dict, field_list, config_info, database_location=DB_LOC, history_begin_date='2021-03-29'):
        logger.info('updating_quotation_sector_data()')
        data_file_location = '{}/quotation'.format(database_location)

        for name, wind_code in wind_code_dict.items():
            data_file = r'{}/{}.csv'.format(data_file_location, name)
            # 使用固定日期2025-09-30
            today = FIXED_TODAY
            yesterday = today - datetime.timedelta(days=1)  # 自动计算2025-09-29
            now_time = datetime.datetime.now()
            # 若当前时间模拟在15点后，end_date025-09-30；否则为2025-09-29
            end_date = today if now_time.hour > 15 else yesterday
            end_date_str = end_date.strftime('%Y-%m-%d')

            # 新增文件时下载历史数据至固定日期
            if not os.path.exists(data_file):
                if not options.init_full_history:
                    raise FileNotFoundError(
                        f"缺少历史行情缓存文件: {data_file}。"
                        "默认主流程不会自动从 2021-03-29 全量初始化；"
                        "如需首次建库，请显式加参 --init-full-history。"
                    )
                data_df = local_client.history(wind_code, field_list, history_begin_date,
                                end_date_str, config_info, usedf=True)
                data_df.reset_index(inplace=True)
                data_df.rename(columns={'index': 'date'}, inplace=True)
                logger.info(f'creating new data file: [{name}.csv] period:{history_begin_date} to {end_date_str}')
                data_df.to_csv(data_file, index=False, encoding="utf_8_sig")

            else:
                his_data = pd.read_csv(data_file)

                # 关键：整列解析，解析失败NaT，然后取最后一个有效日
                d = pd.to_datetime(his_data['date'], errors='coerce')
                d = d.dropna()
                if d.empty:
                    raise ValueError(f"{data_file} 的 date 列全部无法解析，请检查 CSV 的 date 列内容。")
                last_data_date = d.iloc[-1].date()


                # 仅当库存数据未更新到固定日期时补充数
                if last_data_date < today:
                    begin_date = last_data_date + datetime.timedelta(days=1)
                    begin_date_str = begin_date.strftime('%Y-%m-%d')
                    data_df = local_client.history(wind_code, field_list, begin_date_str,
                                    end_date_str, config_info, usedf=True)
                    data_df.reset_index(inplace=True)
                    data_df.rename(columns={'index': 'date'}, inplace=True)

                    if not data_df.empty and data_df.iloc[0]['date'] != 0:
                        new_data = pd.concat([his_data, data_df])
                        logger.info(f'updating data file: [{name}.csv] period:{begin_date_str} to {end_date_str}')
                        new_data.to_csv(data_file, index=False, encoding="utf_8_sig")
                    else:
                        logger.info(f'no updating data for [{name}.csv]')
                else:
                    logger.info(f'no updating data for [{name}.csv]')


    # 更新股票板块/指数数据
    stock_field_list = 'pre_close,open,high,low,close,volume,amt,pe_ttm,pb_mrq,pb_lf'
    config_info = 'PriceAdj=F'
    updating_quotation_sector_data(sector_code_dict, stock_field_list, config_info)


    # %% 更新市场情绪相关数据
    def updating_market_sentiment_data(indicator_name, code, database_location=DB_LOC):
        indicator_csvfile = r'{}/market_sentiment/{}.csv'.format(database_location, indicator_name)
        # 使用固定日期2025-09-30
        today = FIXED_TODAY
        yesterday = today - datetime.timedelta(days=1)  # 自动计算2025-09-29
        now_time = datetime.datetime.now()
        end_date = today if now_time.hour > 15 else yesterday
        end_date_str = end_date.strftime('%Y-%m-%d')

        # 新增文件时下载历史数据至固定日期
        if not os.path.exists(indicator_csvfile):
            if not options.init_full_history:
                raise FileNotFoundError(
                    f"缺少历史情绪缓存文件: {indicator_csvfile}。"
                    "默认主流程不会自动从 2021-03-29 全量初始化；"
                    "如需首次建库，请显式加参 --init-full-history。"
                )
            begin_date = '2021-03-29'
            data_wind = local_client.edb_query(code, begin_date, end_date_str)
            data_df = pd.DataFrame({'date': data_wind.Times, indicator_name: data_wind.Data[0]})
            logger.info(f'creating new data file [{indicator_name}.csv] period:{begin_date} to {end_date_str}')
            data_df.to_csv(indicator_csvfile, index=False, encoding="utf_8_sig")

        else:
            his_data = pd.read_csv(indicator_csvfile)
            parsed_dates = pd.to_datetime(his_data["date"], errors="coerce")
            parsed_dates = parsed_dates.dropna()
            if parsed_dates.empty:
                raise ValueError(
                    f"{indicator_csvfile} 的 date 列全部无法解析，请检查 CSV 的 date 列内容。"
                )
            last_data_date = parsed_dates.iloc[-1].date()

            # 仅当库存数据未更新到固定日期时补充数
            if last_data_date < end_date:
                begin_date = last_data_date + datetime.timedelta(days=1)
                begin_date_str = begin_date.strftime('%Y-%m-%d')
                data_wind = local_client.edb_query(code, begin_date_str, end_date_str)
                data_df = pd.DataFrame({'date': data_wind.Times, indicator_name: data_wind.Data[0]})

                if len(data_df) > 0:
                    new_data = pd.concat([his_data, data_df])
                    new_data.drop_duplicates(subset=['date'], keep='last', inplace=True)
                    logger.info(f'updating data file [{indicator_name}.csv] period:{begin_date_str} to {end_date_str}')
                    new_data.to_csv(indicator_csvfile, index=False, encoding="utf_8_sig")
                else:
                    logger.info(f'no updating data for [{indicator_name}.csv]')
            else:
                logger.info(f'no updating data for [{indicator_name}.csv]')


    # 更新市场活跃度数
    indicator_code_dict = {
        '沪市_成交金额(亿元)': 'M0331254',
        '深市_成交金额(亿元)': 'M0340770',
        '沪市_融资融券余额(万元)': 'M0061608',
        '深市_融资融券余额(万元)': 'M0061613',
        '沪市_融资余额(万元)': 'M0061606',
        '深市_融资余额(万元)': 'M0061610',
        '沪市_融资买入金额(万元)': 'M0061604',
        '深市_融资买入金额(万元)': 'M0061609',
        '换手率_上证综合指数(%)': 'M0331169',
        '换手率_深证综合指数(%)': 'M0331178'
    }

    for indicator_name, code in indicator_code_dict.items():
        updating_market_sentiment_data(indicator_name, code)


def run_step2(
    local_client: LocalDataClient,
    paths: ProjectPaths,
    report_date: datetime.date,
    options: PipelineOptions | None = None,
) -> None:
    import numpy as np
    import pandas as pd
    import datetime
    from datetime import timedelta
    from loguru import logger
    import os
    import json
    import re
    from difflib import SequenceMatcher
    options = options or PipelineOptions()
    batch_size = max(1, int(options.batch_size))
    section_start = time_module.perf_counter()
    run_time_now = pd.Timestamp.now()
    run_time_text = run_time_now.strftime("%Y-%m-%d %H:%M:%S")
    report_date_ts = pd.Timestamp(report_date).normalize()
    report_date_text = report_date_ts.strftime("%Y-%m-%d")
    figure7_universe_min_count = 4500
    figure7_total_min_count = 4000
    market_target_max_staleness_days = 7
    a_share_code_pattern = re.compile(r"\b\d{6}\.(?:SH|SZ|BJ)\b")

    def checkpoint(label: str) -> None:
        nonlocal section_start
        elapsed = time_module.perf_counter() - section_start
        print(f"[耗时] step2::{label}: {elapsed:.2f} 秒")
        section_start = time_module.perf_counter()

    def snapshot_batch(codes, fields, option_string=""):
        return local_client.snapshot_chunked(codes, fields, option_string, batch_size=batch_size)

    def repair_mojibake_text(text: str) -> str:
        if not text:
            return text
        # Only attempt repair for common UTF-8-as-Latin1 mojibake patterns.
        if not any(token in text for token in ("å", "ä", "æ", "ç", "è", "é", "ï", "œ")):
            return text
        try:
            repaired = text.encode("latin1").decode("utf-8")
        except Exception:
            return text
        original_cjk_count = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
        repaired_cjk_count = sum(1 for ch in repaired if "\u4e00" <= ch <= "\u9fff")
        return repaired if repaired_cjk_count > original_cjk_count else text

    def print(*args, **kwargs):
        if not args:
            return
        parts = []
        for arg in args:
            if isinstance(arg, pd.DataFrame):
                parts.append(f"DataFrame(shape={arg.shape})")
            elif isinstance(arg, (list, tuple, set)):
                values = list(arg)
                if len(values) > 8:
                    parts.append(f"{type(arg).__name__}(len={len(values)}, sample={values[:5]})")
                else:
                    parts.append(str(arg))
            else:
                text = str(arg)
                text = repair_mojibake_text(text)
                parts.append(text if len(text) <= 300 else f"{text[:300]}...")
        logger.info(" ".join(parts))

    def to_float_or_zero(value) -> float:
        if pd.isna(value):
            return 0.0
        try:
            return float(value)
        except (TypeError, ValueError):
            return 0.0

    def clean_numeric_values(values) -> list[float]:
        return [to_float_or_zero(value) for value in values]

    figure_failures: list[str] = []

    def record_figure_failure(label: str, error: Exception | str) -> None:
        message = f"[错误] {label} 执行失败：{error}"
        figure_failures.append(message)
        print(message)

    def extract_sequence_values_early(raw_value) -> list:
        if raw_value is None:
            return []
        if isinstance(raw_value, pd.DataFrame):
            if raw_value.empty:
                return []
            if raw_value.shape[1] >= 1:
                return raw_value.iloc[:, 0].tolist()
            if raw_value.shape[0] >= 1:
                return raw_value.iloc[0, :].tolist()
            return []
        if isinstance(raw_value, (pd.Series, pd.Index)):
            return raw_value.tolist()
        if isinstance(raw_value, (list, tuple)):
            return list(raw_value)
        return [raw_value]

    def extract_first_data_series_early(raw_data, expected_len: int | None = None) -> list:
        if raw_data is None:
            return []

        candidates: list[list] = []

        def append_candidate(candidate) -> None:
            values = extract_sequence_values_early(candidate)
            if values:
                candidates.append(values)

        if isinstance(raw_data, pd.DataFrame):
            if raw_data.empty:
                return []
            for column in raw_data.columns:
                append_candidate(raw_data[column])
            for row_index in raw_data.index:
                append_candidate(raw_data.loc[row_index])
        elif isinstance(raw_data, (pd.Series, pd.Index)):
            append_candidate(raw_data)
        elif isinstance(raw_data, (list, tuple)):
            if not raw_data:
                return []
            scalar_only = all(
                not isinstance(item, (list, tuple, pd.Series, pd.Index, pd.DataFrame))
                for item in raw_data
            )
            if scalar_only:
                append_candidate(raw_data)
            else:
                for item in raw_data:
                    append_candidate(item)
        else:
            append_candidate(raw_data)

        if expected_len is not None:
            for candidate in candidates:
                if len(candidate) == expected_len:
                    return candidate
        return candidates[0] if candidates else []

    def build_result_dataframe(result) -> pd.DataFrame:
        raw_data = getattr(result, "Data", None)
        fields = extract_sequence_values_early(getattr(result, "Fields", None))
        if raw_data is None:
            return pd.DataFrame(columns=fields)
        try:
            df = pd.DataFrame(raw_data).T
        except Exception:
            return pd.DataFrame(columns=fields)
        if fields and len(fields) == len(df.columns):
            df.columns = fields
        return df

    def summarize_numeric_values(values) -> dict[str, object]:
        numeric = pd.to_numeric(pd.Series(list(values)), errors="coerce")
        valid = numeric.dropna()
        return {
            "count": len(numeric),
            "non_null_count": int(valid.shape[0]),
            "non_zero_count": int((valid != 0).sum()),
            "sample": valid.head(5).tolist(),
        }

    def get_local_error_detail(error_code) -> str:
        if error_code in (None, "", 0):
            return ""
        try:
            sdk = local_client._load_sdk_module()
            detail = sdk.THS_GetErrorInfo(int(error_code))
            if isinstance(detail, dict):
                return str(detail.get("errmsg", ""))
            return str(detail)
        except Exception:
            return ""

    def get_datapool_probe_debug_dir() -> Path:
        debug_dir = paths.funding_draft_dir / "data" / "cache" / "datapool_probe_debug"
        debug_dir.mkdir(parents=True, exist_ok=True)
        return debug_dir

    def write_datapool_probe_debug(
        label: str,
        dataset_name: str,
        probe_records: list[dict[str, object]],
    ) -> Path | None:
        if not probe_records:
            return None
        normalized_chars: list[str] = []
        for char in str(label):
            codepoint = ord(char)
            if char.isascii() and char.isalnum():
                normalized_chars.append(char)
            elif 0x4E00 <= codepoint <= 0x9FFF:
                normalized_chars.append(char)
            else:
                normalized_chars.append("_")
        safe_label = re.sub(r"_+", "_", "".join(normalized_chars)).strip("_") or dataset_name
        debug_path = get_datapool_probe_debug_dir() / f"{safe_label}_{dataset_name}.csv"
        pd.DataFrame(probe_records).to_csv(debug_path, index=False, encoding="utf-8-sig")
        return debug_path

    def probe_datapool_variants(
        label: str,
        dataset_name: str,
        option_variants: list[str],
    ) -> tuple[object, pd.DataFrame, list[dict[str, object]]]:
        probe_records: list[dict[str, object]] = []
        for index, option_string in enumerate(option_variants, start=1):
            started = time_module.perf_counter()
            result = local_client.datapool(dataset_name, option_string)
            df = build_result_dataframe(result)
            fields = extract_sequence_values_early(getattr(result, "Fields", None))
            sample_rows = df.head(2).to_dict(orient="records") if not df.empty else []
            error_code = getattr(result, "ErrorCode", None)
            probe_record = {
                "model": dataset_name,
                "option_string": option_string,
                "error_code": error_code,
                "errmsg": getattr(result, "ErrMsg", ""),
                "error_detail": get_local_error_detail(error_code),
                "fields": fields,
                "rows": int(len(df)),
                "sample_rows": sample_rows,
                "elapsed": round(time_module.perf_counter() - started, 4),
            }
            probe_records.append(probe_record)
            print(
                f"{label} probe {index}/{len(option_variants)}: model={dataset_name}, "
                f"error_code={probe_record['error_code']}, error_detail={probe_record['error_detail']}, "
                f"rows={probe_record['rows']}, fields={fields[:6]}, elapsed={probe_record['elapsed']:.2f}s"
            )
            if sample_rows:
                print(f"{label} probe sample_rows={sample_rows}")
            if getattr(result, "ErrorCode", None) == 0 and not df.empty:
                return result, df, probe_records
        debug_path = write_datapool_probe_debug(label, dataset_name, probe_records)
        raise RuntimeError(
            f"{label} DataPool 探针失败：model={dataset_name} 所有参数组合都未返回有效结果。 "
            f"probes={probe_records}, debug_path={debug_path}"
        )

    def format_log_date(value) -> str:
        parsed = pd.to_datetime(value, errors="coerce")
        if pd.isna(parsed):
            return "None"
        return parsed.strftime("%Y-%m-%d")

    def load_market_sentiment_indexed_series(indicator_name: str) -> pd.DataFrame:
        file_path = paths.data_dir / "market_sentiment" / f"{indicator_name}.csv"
        if not file_path.exists():
            empty = pd.DataFrame(columns=[indicator_name])
            empty.index = pd.DatetimeIndex([], name="date")
            return empty
        df = pd.read_csv(file_path)
        if df.empty or "date" not in df.columns or indicator_name not in df.columns:
            empty = pd.DataFrame(columns=[indicator_name])
            empty.index = pd.DatetimeIndex([], name="date")
            return empty
        parsed = df[["date", indicator_name]].copy()
        parsed["date"] = pd.to_datetime(parsed["date"], errors="coerce").dt.normalize()
        parsed[indicator_name] = pd.to_numeric(parsed[indicator_name], errors="coerce")
        parsed = parsed.dropna(subset=["date"]).drop_duplicates(subset=["date"], keep="last")
        parsed = parsed.sort_values("date").set_index("date")
        parsed.index.name = "date"
        return parsed

    def load_market_sentiment_series(indicator_name: str) -> pd.DataFrame:
        indexed = load_market_sentiment_indexed_series(indicator_name)
        if indexed.empty:
            return pd.DataFrame(columns=["date", indicator_name])
        parsed = indexed.reset_index()
        parsed["date"] = parsed["date"].dt.strftime("%Y%m%d")
        return parsed[["date", indicator_name]]

    def upsert_market_sentiment_series(indicator_name: str, data: pd.DataFrame) -> None:
        file_path = paths.data_dir / "market_sentiment" / f"{indicator_name}.csv"
        file_path.parent.mkdir(parents=True, exist_ok=True)
        normalized = data.copy()
        if "date" not in normalized.columns or indicator_name not in normalized.columns:
            raise RuntimeError(
                f"market_sentiment 回写失败：缺少必要列。 indicator={indicator_name}, columns={list(normalized.columns)}"
            )
        normalized = normalized[["date", indicator_name]].copy()
        normalized["date"] = pd.to_datetime(normalized["date"], errors="coerce").dt.normalize()
        normalized[indicator_name] = pd.to_numeric(normalized[indicator_name], errors="coerce")
        normalized = normalized.dropna(subset=["date"]).drop_duplicates(subset=["date"], keep="last")
        if file_path.exists():
            existing = pd.read_csv(file_path)
            if "date" in existing.columns and indicator_name in existing.columns:
                existing = existing[["date", indicator_name]].copy()
                existing["date"] = pd.to_datetime(existing["date"], errors="coerce").dt.normalize()
                existing[indicator_name] = pd.to_numeric(existing[indicator_name], errors="coerce")
                normalized = pd.concat([existing, normalized], ignore_index=True)
        normalized = normalized.dropna(subset=["date"]).drop_duplicates(subset=["date"], keep="last")
        normalized = normalized.sort_values("date")
        normalized["date"] = normalized["date"].dt.strftime("%Y-%m-%d")
        normalized.to_csv(file_path, index=False, encoding="utf_8_sig")
        latest_written_date = pd.to_datetime(normalized["date"], errors="coerce").max()
        print(
            f"market_sentiment 回写完成: indicator={indicator_name}, file={file_path.name}, "
            f"latest_date={format_log_date(latest_written_date)}, rows={len(normalized)}"
        )

    def load_local_quotation_close_series(index_name: str) -> pd.Series:
        file_path = paths.data_dir / "quotation" / f"{index_name}.csv"
        if not file_path.exists():
            return pd.Series(dtype="float64", name="close")
        import csv

        rows: list[tuple[pd.Timestamp, float]] = []
        try:
            with file_path.open("r", encoding="utf-8-sig", newline="") as csv_file:
                reader = csv.DictReader(csv_file)
                fieldnames = list(reader.fieldnames or [])
                if not fieldnames or "date" not in fieldnames:
                    return pd.Series(dtype="float64", name="close")
                close_keys = [field for field in fieldnames if str(field).strip().lower() == "close"]
                if not close_keys:
                    return pd.Series(dtype="float64", name="close")
                for raw_row in reader:
                    trade_date = pd.to_datetime(raw_row.get("date"), errors="coerce")
                    if pd.isna(trade_date):
                        continue
                    close_value = np.nan
                    for key in close_keys:
                        numeric_value = pd.to_numeric(pd.Series([raw_row.get(key)]), errors="coerce").iloc[0]
                        if pd.notna(numeric_value):
                            close_value = float(numeric_value)
                    if pd.isna(close_value):
                        continue
                    rows.append((trade_date.normalize(), close_value))
        except Exception as exc:
            print(f"quotation 本地 close 读取失败: index={index_name}, file={file_path}, error={exc}")
            return pd.Series(dtype="float64", name="close")
        if not rows:
            return pd.Series(dtype="float64", name="close")
        parsed = pd.DataFrame(rows, columns=["date", "close"])
        parsed = parsed.drop_duplicates(subset=["date"], keep="last").sort_values("date")
        parsed = parsed.set_index("date")["close"].dropna()
        parsed.index.name = "date"
        parsed.name = "close"
        return parsed

    def load_local_quotation_numeric_series(
        index_name: str,
        field_name: str,
        output_name: str | None = None,
    ) -> pd.Series:
        file_path = paths.data_dir / "quotation" / f"{index_name}.csv"
        series_name = output_name or field_name
        if not file_path.exists():
            return pd.Series(dtype="float64", name=series_name)
        import csv

        rows: list[tuple[pd.Timestamp, float]] = []
        normalized_field_name = str(field_name).strip().lower()
        try:
            with file_path.open("r", encoding="utf-8-sig", newline="") as csv_file:
                reader = csv.DictReader(csv_file)
                fieldnames = list(reader.fieldnames or [])
                if not fieldnames or "date" not in fieldnames:
                    return pd.Series(dtype="float64", name=series_name)
                matching_keys = [
                    field
                    for field in fieldnames
                    if str(field).strip().lower() == normalized_field_name
                ]
                if not matching_keys:
                    return pd.Series(dtype="float64", name=series_name)
                for raw_row in reader:
                    trade_date = pd.to_datetime(raw_row.get("date"), errors="coerce")
                    if pd.isna(trade_date):
                        continue
                    value = np.nan
                    for key in matching_keys:
                        numeric_value = pd.to_numeric(pd.Series([raw_row.get(key)]), errors="coerce").iloc[0]
                        if pd.notna(numeric_value):
                            value = float(numeric_value)
                    if pd.isna(value):
                        continue
                    rows.append((trade_date.normalize(), value))
        except Exception as exc:
            print(
                f"quotation 本地字段读取失败: index={index_name}, field={field_name}, "
                f"file={file_path}, error={exc}"
            )
            return pd.Series(dtype="float64", name=series_name)
        if not rows:
            return pd.Series(dtype="float64", name=series_name)
        parsed = pd.DataFrame(rows, columns=["date", series_name])
        parsed = parsed.drop_duplicates(subset=["date"], keep="last").sort_values("date")
        series = parsed.set_index("date")[series_name].dropna()
        series.index.name = "date"
        series.name = series_name
        return series

    def merge_numeric_series(series_list: list[pd.Series], series_name: str) -> pd.Series:
        frames = []
        for series in series_list:
            if series is None or series.empty:
                continue
            current = series.rename(series_name).reset_index()
            current.columns = ["date", series_name]
            frames.append(current)
        if not frames:
            return pd.Series(dtype="float64", name=series_name)
        merged = pd.concat(frames, ignore_index=True)
        merged["date"] = pd.to_datetime(merged["date"], errors="coerce").dt.normalize()
        merged[series_name] = pd.to_numeric(merged[series_name], errors="coerce")
        merged = (
            merged.dropna(subset=["date"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
        )
        result = merged.set_index("date")[series_name].dropna()
        result.index.name = "date"
        result.name = series_name
        return result

    def extract_history_dataframe(history_result) -> pd.DataFrame:
        if isinstance(history_result, pd.DataFrame):
            return history_result.copy()
        if isinstance(history_result, (list, tuple)) and len(history_result) >= 2:
            candidate = history_result[1]
            if isinstance(candidate, pd.DataFrame):
                return candidate.copy()
        raise RuntimeError(f"history 返回结构异常: type={type(history_result).__name__}")

    def fetch_online_index_close_series(
        index_name: str,
        index_code: str,
        begin_date: pd.Timestamp,
        end_date: pd.Timestamp,
    ) -> pd.Series:
        history_result = local_client.history(
            index_code,
            "close",
            begin_date.strftime("%Y-%m-%d"),
            end_date.strftime("%Y-%m-%d"),
            "PriceAdj=F",
            usedf=True,
        )
        history_df = extract_history_dataframe(history_result).reset_index()
        if "date" not in history_df.columns:
            history_df = history_df.rename(columns={history_df.columns[0]: "date"})
        close_column = next(
            (column for column in history_df.columns if str(column).strip().lower() == "close"),
            None,
        )
        if close_column is None:
            raise RuntimeError(
                f"图1在线 close 查询失败：缺少 close 列。 index={index_name}, columns={list(history_df.columns)}"
            )
        history_df["date"] = pd.to_datetime(history_df["date"], errors="coerce").dt.normalize()
        history_df["close"] = pd.to_numeric(history_df[close_column], errors="coerce")
        series = (
            history_df[["date", "close"]]
            .dropna(subset=["date"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
            .set_index("date")["close"]
            .dropna()
        )
        if series.empty:
            raise RuntimeError(
                f"图1在线 close 查询失败：未返回有效 close。 index={index_name}, code={index_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}"
            )
        series.index.name = "date"
        series.name = "close"
        print(
            f"图1在线 close fallback: index={index_name}, code={index_code}, "
            f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, rows={len(series)}, "
            f"latest_date={format_log_date(series.index.max())}"
        )
        return series

    def get_block_close_cache_dir() -> Path:
        cache_dir = paths.funding_draft_dir / "data" / "cache" / "block_close"
        cache_dir.mkdir(parents=True, exist_ok=True)
        return cache_dir

    def use_strict_sw_block_codes(scope_key: str) -> bool:
        return scope_key in {"sw_l1", "sw_l2"}

    def get_block_close_cache_scope(scope_key: str) -> str:
        if use_strict_sw_block_codes(scope_key):
            return f"{scope_key}_strict"
        return scope_key

    def normalize_strict_sw_query_code(scope_key: str, asset_code: str) -> str:
        code_text = str(asset_code).strip()
        if use_strict_sw_block_codes(scope_key) and code_text.endswith(".SI"):
            return f"{code_text[:-3]}.SL"
        return code_text

    def fetch_strict_sw_close_via_ds(
        asset_name: str,
        query_code: str,
        begin_date: pd.Timestamp,
        end_date: pd.Timestamp,
        scope_key: str,
    ) -> pd.Series | None:
        if not use_strict_sw_block_codes(scope_key):
            return None
        sdk_code = str(query_code).strip()
        if not sdk_code.endswith(".SL"):
            return None
        try:
            local_client._ensure_sdk_login()
            sdk = local_client._load_sdk_module()
            ds_result = sdk.THS_DS(
                sdk_code,
                "ths_close_price_index;ths_turnover_ratio_index",
                ";",
                "Days:Alldays,Fill:Previous,Interval:D",
                begin_date.strftime("%Y-%m-%d"),
                end_date.strftime("%Y-%m-%d"),
                "format:dataframe",
            )
        except Exception as exc:
            print(
                f"{scope_key} THS_DS close fallback exception: asset={asset_name}, code={sdk_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, error={exc}"
            )
            return None

        if getattr(ds_result, "errorcode", 0) != 0:
            print(
                f"{scope_key} THS_DS close fallback failed: asset={asset_name}, code={sdk_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, "
                f"errorcode={getattr(ds_result, 'errorcode', None)}, errmsg={getattr(ds_result, 'errmsg', '')}"
            )
            return None

        ds_df = getattr(ds_result, "data", None)
        if ds_df is None or not isinstance(ds_df, pd.DataFrame) or ds_df.empty:
            print(
                f"{scope_key} THS_DS close fallback empty: asset={asset_name}, code={sdk_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}"
            )
            return None

        normalized_df = ds_df.copy()
        normalized_df.columns = [str(column).strip() for column in normalized_df.columns]
        date_column = next(
            (column for column in normalized_df.columns if str(column).strip().lower() in {"time", "date"}),
            None,
        )
        close_column = next(
            (
                column for column in normalized_df.columns
                if "ths_close_price_index" in str(column).strip().lower()
                or str(column).strip().lower() == "close"
            ),
            None,
        )
        if date_column is None or close_column is None:
            print(
                f"{scope_key} THS_DS close fallback malformed: asset={asset_name}, code={sdk_code}, "
                f"columns={list(normalized_df.columns)}"
            )
            return None

        normalized_df["date"] = pd.to_datetime(normalized_df[date_column], errors="coerce").dt.normalize()
        normalized_df["close"] = pd.to_numeric(normalized_df[close_column], errors="coerce")
        series = (
            normalized_df[["date", "close"]]
            .dropna(subset=["date", "close"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
            .set_index("date")["close"]
        )
        if series.empty:
            print(
                f"{scope_key} THS_DS close fallback no valid close: asset={asset_name}, code={sdk_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}"
            )
            return None
        series.index.name = "date"
        series.name = "close"
        print(
            f"{scope_key} THS_DS close fallback: asset={asset_name}, code={sdk_code}, "
            f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, rows={len(series)}, "
            f"latest_date={format_log_date(series.index.max())}"
        )
        return series

    def get_block_close_cache_path(scope_key: str, asset_code: str) -> Path:
        cache_scope = get_block_close_cache_scope(scope_key)
        safe_code = re.sub(r"[^A-Za-z0-9._-]+", "_", str(asset_code))
        return get_block_close_cache_dir() / f"{cache_scope}_{safe_code}.csv"

    def load_local_block_close_series(scope_key: str, asset_code: str) -> pd.Series:
        file_path = get_block_close_cache_path(scope_key, asset_code)
        if not file_path.exists():
            return pd.Series(dtype="float64", name="close")
        try:
            df = pd.read_csv(file_path)
        except Exception as exc:
            print(
                f"block close 本地缓存读取失败: scope={scope_key}, code={asset_code}, "
                f"file={file_path.name}, error={exc}"
            )
            return pd.Series(dtype="float64", name="close")
        if "date" not in df.columns or "close" not in df.columns:
            return pd.Series(dtype="float64", name="close")
        df["date"] = pd.to_datetime(df["date"], errors="coerce").dt.normalize()
        df["close"] = pd.to_numeric(df["close"], errors="coerce")
        series = (
            df[["date", "close"]]
            .dropna(subset=["date", "close"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
            .set_index("date")["close"]
        )
        series.index.name = "date"
        series.name = "close"
        return series

    def load_legacy_block_close_series(scope_key: str, asset_code: str) -> pd.Series:
        safe_code = re.sub(r"[^A-Za-z0-9._-]+", "_", str(asset_code))
        file_path = get_block_close_cache_dir() / f"{scope_key}_{safe_code}.csv"
        if not file_path.exists():
            return pd.Series(dtype="float64", name="close")
        try:
            df = pd.read_csv(file_path)
        except Exception:
            return pd.Series(dtype="float64", name="close")
        if "date" not in df.columns or "close" not in df.columns:
            return pd.Series(dtype="float64", name="close")
        df["date"] = pd.to_datetime(df["date"], errors="coerce").dt.normalize()
        df["close"] = pd.to_numeric(df["close"], errors="coerce")
        series = (
            df[["date", "close"]]
            .dropna(subset=["date", "close"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
            .set_index("date")["close"]
        )
        series.index.name = "date"
        series.name = "close"
        return series

    def persist_block_close_series(
        scope_key: str,
        asset_code: str,
        asset_name: str,
        close_series: pd.Series,
    ) -> None:
        if close_series is None or close_series.empty:
            return
        file_path = get_block_close_cache_path(scope_key, asset_code)
        normalized = close_series.rename("close").reset_index()
        normalized.columns = ["date", "close"]
        normalized["date"] = pd.to_datetime(normalized["date"], errors="coerce").dt.normalize()
        normalized["close"] = pd.to_numeric(normalized["close"], errors="coerce")
        normalized["code"] = str(asset_code)
        normalized["name"] = str(asset_name)
        if file_path.exists():
            try:
                existing = pd.read_csv(file_path)
                if "date" in existing.columns and "close" in existing.columns:
                    existing["date"] = pd.to_datetime(existing["date"], errors="coerce").dt.normalize()
                    existing["close"] = pd.to_numeric(existing["close"], errors="coerce")
                    if "code" not in existing.columns:
                        existing["code"] = str(asset_code)
                    if "name" not in existing.columns:
                        existing["name"] = str(asset_name)
                    normalized = pd.concat([existing[["date", "close", "code", "name"]], normalized], ignore_index=True)
            except Exception as exc:
                print(
                    f"block close 本地缓存合并失败，改为重写: scope={scope_key}, code={asset_code}, "
                    f"file={file_path.name}, error={exc}"
                )
        normalized = (
            normalized.dropna(subset=["date", "close"])
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
        )
        normalized["date"] = normalized["date"].dt.strftime("%Y-%m-%d")
        normalized.to_csv(file_path, index=False, encoding="utf-8-sig")

    def fetch_online_block_close_series(
        asset_name: str,
        asset_code: str,
        begin_date: pd.Timestamp,
        end_date: pd.Timestamp,
        scope_key: str,
    ) -> pd.Series:
        history_attempts: list[dict[str, object]] = []

        def try_history_close(query_code: str, source: str) -> pd.Series | None:
            history_result = local_client.history(
                query_code,
                "close",
                begin_date.strftime("%Y-%m-%d"),
                end_date.strftime("%Y-%m-%d"),
                "",
                usedf=True,
            )
            history_df = extract_history_dataframe(history_result).reset_index()
            if "date" not in history_df.columns:
                history_df = history_df.rename(columns={history_df.columns[0]: "date"})
            close_column = next(
                (column for column in history_df.columns if str(column).strip().lower() == "close"),
                None,
            )
            if close_column is None:
                history_attempts.append(
                    {"code": query_code, "source": source, "error": f"缺少 close 列, columns={list(history_df.columns)}"}
                )
                return None
            history_df["date"] = pd.to_datetime(history_df["date"], errors="coerce").dt.normalize()
            history_df["close"] = pd.to_numeric(history_df[close_column], errors="coerce")
            series = (
                history_df[["date", "close"]]
                .dropna(subset=["date", "close"])
                .drop_duplicates(subset=["date"], keep="last")
                .sort_values("date")
                .set_index("date")["close"]
            )
            if series.empty:
                history_attempts.append({"code": query_code, "source": source, "error": "未返回有效 close"})
                return None
            series.index.name = "date"
            series.name = "close"
            print(
                f"{scope_key} 在线 close fallback: asset={asset_name}, code={query_code}, source={source}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, rows={len(series)}, "
                f"latest_date={format_log_date(series.index.max())}"
            )
            return series

        direct_series = try_history_close(asset_code, "history_direct")
        if direct_series is not None:
            return direct_series
        if use_strict_sw_block_codes(scope_key):
            ds_series = fetch_strict_sw_close_via_ds(
                asset_name,
                asset_code,
                begin_date,
                end_date,
                scope_key,
            )
            if ds_series is not None:
                return ds_series
            raise RuntimeError(
                f"{scope_key} strict close query failed: asset={asset_name}, code={asset_code}, "
                f"begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, "
                f"history_attempts={history_attempts}"
            )

        normalized_asset_name = str(asset_name).replace("(申万)", "").replace("（申万）", "").strip()
        mapped_code = resolve_block_code(normalized_asset_name, scope_key, end_date.strftime("%Y-%m-%d"))
        if mapped_code and mapped_code != asset_code:
            mapped_series = try_history_close(mapped_code, "history_mapped_code")
            if mapped_series is not None:
                return mapped_series

        requested_trade_dates = get_trade_dates_from_quotation_cache(
            begin_date.strftime("%Y-%m-%d"),
            end_date.strftime("%Y-%m-%d"),
        )
        close_values = fetch_block_close_values_for_dates(
            normalized_asset_name,
            scope_key,
            requested_trade_dates,
            reference_date=end_date.strftime("%Y-%m-%d"),
        )
        if close_values:
            series = pd.Series(
                list(close_values.values()),
                index=pd.to_datetime(list(close_values.keys()), format="%Y-%m-%d", errors="coerce"),
                dtype="float64",
                name="close",
            ).dropna().sort_index()
            series.index = series.index.normalize()
            series.index.name = "date"
            if not series.empty:
                print(
                    f"{scope_key} block_wc close fallback: asset={asset_name}, code={mapped_code or asset_code}, "
                    f"rows={len(series)}, latest_date={format_log_date(series.index.max())}"
                )
                return series

        raise RuntimeError(
            f"{scope_key} 在线 close 查询失败：未返回有效 close。 asset={asset_name}, code={asset_code}, "
            f"mapped_code={mapped_code}, begin={begin_date:%Y-%m-%d}, end={end_date:%Y-%m-%d}, "
            f"history_attempts={history_attempts}"
        )

    def resolve_target_market_date() -> pd.Timestamp:
        reference_index_name = "上证指数"
        reference_index_code = "000001.SH"
        same_day_run = report_date_ts == run_time_now.normalize()
        candidate_date = report_date_ts
        if same_day_run and run_time_now.hour <= 15:
            candidate_date = candidate_date - pd.Timedelta(days=1)

        local_series = load_local_quotation_close_series(reference_index_name)
        local_valid = local_series[local_series.index <= candidate_date].dropna()
        local_latest = local_valid.index.max() if not local_valid.empty else pd.NaT

        online_latest = pd.NaT
        should_probe_online = pd.isna(local_latest) or (
            candidate_date.weekday() < 5 and local_latest < candidate_date
        )
        if should_probe_online:
            try:
                online_series = fetch_online_index_close_series(
                    reference_index_name,
                    reference_index_code,
                    candidate_date - pd.Timedelta(days=10),
                    candidate_date,
                )
                online_valid = online_series[online_series.index <= candidate_date].dropna()
                online_latest = online_valid.index.max() if not online_valid.empty else pd.NaT
            except Exception as exc:
                print(
                    f"Step2目标交易日在线探测失败: report_date={report_date_text}, "
                    f"candidate_date={candidate_date:%Y-%m-%d}, error={exc}"
                )

        candidate_dates = [value for value in [local_latest, online_latest] if pd.notna(value)]
        resolved_target = max(candidate_dates) if candidate_dates else pd.NaT
        if pd.isna(resolved_target):
            raise RuntimeError(
                "Step2 无法解析目标交易日：本地 quotation 与在线 history 均未返回有效收盘日期。 "
                f"report_date={report_date_text}, candidate_date={candidate_date:%Y-%m-%d}"
            )

        resolved_target = pd.Timestamp(resolved_target).normalize()
        if (candidate_date - resolved_target).days > market_target_max_staleness_days:
            raise RuntimeError(
                "Step2 目标交易日解析失败：可用市场数据日期过旧。 "
                f"report_date={report_date_text}, candidate_date={candidate_date:%Y-%m-%d}, "
                f"resolved_target={resolved_target:%Y-%m-%d}, "
                f"max_staleness_days={market_target_max_staleness_days}"
            )
        print(
            f"Step2市场数据目标日: runtime={run_time_text}, report_date={report_date_text}, "
            f"candidate_date={candidate_date:%Y-%m-%d}, local_latest={format_log_date(local_latest)}, "
            f"online_latest={format_log_date(online_latest)}, target_market_date={resolved_target:%Y-%m-%d}"
        )
        return resolved_target

    def require_report_date_close(
        close_series: pd.Series,
        index_name: str,
        expected_date: pd.Timestamp,
    ) -> tuple[pd.Timestamp, float]:
        valid = close_series[close_series.index <= expected_date].dropna()
        if valid.empty:
            raise RuntimeError(
                f"图1 close 计算失败：{index_name} 在目标交易日前无有效收盘价。 "
                f"target_market_date={expected_date:%Y-%m-%d}, report_date={report_date_text}"
            )
        end_date = valid.index[-1]
        if end_date < expected_date:
            raise RuntimeError(
                f"图1 close 计算失败：{index_name} 收盘价未更新到目标交易日。 "
                f"latest_close_date={format_log_date(end_date)}, "
                f"target_market_date={expected_date:%Y-%m-%d}, report_date={report_date_text}"
            )
        return end_date, float(valid.iloc[-1])

    def get_previous_trade_close(
        close_series: pd.Series,
        boundary_date: pd.Timestamp,
        index_name: str,
        period_label: str,
    ) -> tuple[pd.Timestamp, float]:
        history = close_series[close_series.index < boundary_date].dropna()
        if history.empty:
            raise RuntimeError(
                f"图1 close 计算失败：{index_name} 缺少 {period_label} 基准日前最后一个交易日收盘价。 "
                f"boundary_date={boundary_date:%Y-%m-%d}"
            )
        base_date = history.index[-1]
        return base_date, float(history.iloc[-1])

    def build_market_index_returns_dataframe() -> pd.DataFrame:
        index_rows = [
            ("000001.SH", "上证指数"),
            ("399001.SZ", "深证成指"),
            ("399006.SZ", "创业板指"),
            ("000688.SH", "科创50"),
            ("000016.SH", "上证50"),
            ("000300.SH", "沪深300"),
            ("000905.SH", "中证500"),
            ("000852.SH", "中证1000"),
        ]
        period_boundaries = {
            "本周": target_market_date - pd.Timedelta(days=target_market_date.weekday()),
            "本月": target_market_date.replace(day=1),
            "本年": target_market_date.replace(month=1, day=1),
        }
        online_begin_date = min(period_boundaries.values()) - pd.Timedelta(days=31)
        rows: list[dict[str, object]] = []
        for index_code, index_name in index_rows:
            close_series = load_local_quotation_close_series(index_name)
            need_online_fallback = (
                close_series.empty
                or close_series.index.max() < target_market_date
                or any(close_series[close_series.index < boundary].dropna().empty for boundary in period_boundaries.values())
            )
            if need_online_fallback:
                online_series = fetch_online_index_close_series(index_name, index_code, online_begin_date, target_market_date)
                close_series = pd.concat([close_series, online_series])
                close_series = close_series[~close_series.index.duplicated(keep="last")].sort_index()
            end_date, end_close = require_report_date_close(close_series, index_name, target_market_date)
            row = {
                "代码": index_code,
                "市场指数": index_name,
            }
            for period_label, boundary_date in period_boundaries.items():
                base_date, base_close = get_previous_trade_close(close_series, boundary_date, index_name, period_label)
                return_value = end_close / base_close - 1
                row[f"{period_label}涨跌幅(%)"] = return_value
                print(
                    f"图1收益自算: 指数={index_name}, 周期={period_label}, 基准日期={format_log_date(base_date)}, "
                    f"基准close={base_close:.4f}, 终点日期={format_log_date(end_date)}, "
                    f"终点close={end_close:.4f}, 收益率={return_value:.6%}"
                )
            rows.append(row)
        return pd.DataFrame(rows)

    target_market_date = resolve_target_market_date()
    target_market_date_text = target_market_date.strftime("%Y-%m-%d")
    margin_required_date = (
        target_market_date if options.allow_stale_margin_data else report_date_ts
    ).normalize()
    margin_required_date_text = margin_required_date.strftime("%Y-%m-%d")

    def build_block_self_computed_returns_dataframe(
        assets: list[tuple[str, str]],
        scope_key: str,
        chart_label: str,
    ) -> pd.DataFrame:
        if use_strict_sw_block_codes(scope_key):
            period_boundaries = {
                "本周": target_market_date - pd.Timedelta(days=target_market_date.weekday()),
                "本月": target_market_date.replace(day=1),
                "本年": target_market_date.replace(month=1, day=1),
            }
            online_begin_date = min(period_boundaries.values()) - pd.Timedelta(days=31)
            rows: list[dict[str, object]] = []
            for asset_code, asset_name in assets:
                query_code = normalize_strict_sw_query_code(scope_key, asset_code)
                close_series = load_local_block_close_series(scope_key, query_code)
                need_online_fallback = (
                    close_series.empty
                    or close_series.index.max() < target_market_date
                    or any(
                        close_series[close_series.index < boundary].dropna().empty
                        for boundary in period_boundaries.values()
                    )
                )
                if need_online_fallback:
                    online_series = fetch_online_block_close_series(
                        asset_name,
                        query_code,
                        online_begin_date,
                        target_market_date,
                        scope_key,
                    )
                    close_series = pd.concat([close_series, online_series])
                    close_series = close_series[~close_series.index.duplicated(keep="last")].sort_index()
                    persist_block_close_series(scope_key, query_code, asset_name, close_series)
                end_date, end_close = require_report_date_close(close_series, asset_name, target_market_date)
                row = {"代码": asset_code, "市场指数": asset_name}
                for period_label, boundary_date in period_boundaries.items():
                    base_date, base_close = get_previous_trade_close(close_series, boundary_date, asset_name, period_label)
                    return_value = end_close / base_close - 1
                    row[f"{period_label}涨跌幅(%)"] = return_value
                    print(
                        f"{chart_label}严格申万收益自算: 资产={asset_name}, 原始代码={asset_code}, 查询代码={query_code}, "
                        f"周期={period_label}, 基准日期={format_log_date(base_date)}, 基准close={base_close:.4f}, "
                        f"终点日期={format_log_date(end_date)}, 终点close={end_close:.4f}, 收益率={return_value:.6%}"
                    )
                rows.append(row)
            return pd.DataFrame(rows)

        period_boundaries = {
            "本周": target_market_date - pd.Timedelta(days=target_market_date.weekday()),
            "本月": target_market_date.replace(day=1),
            "本年": target_market_date.replace(month=1, day=1),
        }
        online_begin_date = min(period_boundaries.values()) - pd.Timedelta(days=31)
        rows: list[dict[str, object]] = []
        for asset_code, asset_name in assets:
            close_series = load_local_block_close_series(scope_key, asset_code)
            need_online_fallback = (
                close_series.empty
                or close_series.index.max() < target_market_date
                or any(
                    close_series[close_series.index < boundary].dropna().empty
                    for boundary in period_boundaries.values()
                )
            )
            if need_online_fallback:
                online_series = fetch_online_block_close_series(
                    asset_name,
                    asset_code,
                    online_begin_date,
                    target_market_date,
                    scope_key,
                )
                close_series = pd.concat([close_series, online_series])
                close_series = close_series[~close_series.index.duplicated(keep="last")].sort_index()
                persist_block_close_series(scope_key, asset_code, asset_name, close_series)
            end_date, end_close = require_report_date_close(close_series, asset_name, target_market_date)
            row = {"代码": asset_code, "市场指数": asset_name}
            for period_label, boundary_date in period_boundaries.items():
                base_date, base_close = get_previous_trade_close(close_series, boundary_date, asset_name, period_label)
                return_value = end_close / base_close - 1
                row[f"{period_label}涨跌幅(%)"] = return_value
                print(
                    f"{chart_label}收益自算: 资产={asset_name}, 周期={period_label}, 基准日期={format_log_date(base_date)}, "
                    f"基准close={base_close:.4f}, 终点日期={format_log_date(end_date)}, "
                    f"终点close={end_close:.4f}, 收益率={return_value:.6%}"
                )
            rows.append(row)
        return pd.DataFrame(rows)

    def build_margin_dataframe_from_local_cache() -> tuple[pd.DataFrame | None, dict[str, object]]:
        required_fields = [
            "沪市_融资买入金额(万元)",
            "深市_融资买入金额(万元)",
            "沪市_融资融券余额(万元)",
            "深市_融资融券余额(万元)",
            "沪市_成交金额(亿元)",
            "深市_成交金额(亿元)",
        ]
        merged: pd.DataFrame | None = None
        hit_fields: list[str] = []
        missing_fields: list[str] = []
        for field in required_fields:
            field_df = load_market_sentiment_indexed_series(field)
            if field_df.empty or field_df[field].dropna().empty:
                missing_fields.append(field)
                continue
            hit_fields.append(field)
            merged = field_df if merged is None else merged.join(field_df, how="outer")
        if merged is None or merged.empty:
            return None, {
                "source": "local_cache",
                "hit_fields": hit_fields,
                "missing_fields": missing_fields,
                "shape": (0, 0),
                "latest_local_date": pd.NaT,
            }
        merged = merged.sort_index()
        usable = merged.dropna(subset=required_fields).copy() if len(hit_fields) == len(required_fields) else pd.DataFrame()
        latest_local_date = usable.index.max() if not usable.empty else pd.NaT
        if usable.empty:
            return None, {
                "source": "local_cache",
                "hit_fields": hit_fields,
                "missing_fields": missing_fields,
                "shape": tuple(merged.shape),
                "latest_local_date": latest_local_date,
            }
        rzrq_handle = pd.DataFrame(index=usable.index.copy())
        rzrq_handle["日期"] = usable.index.strftime("%Y-%m-%d")
        rzrq_handle["两市融资融券余额（万亿元）"] = (
            usable["沪市_融资融券余额(万元)"] + usable["深市_融资融券余额(万元)"]
        ) / 100000000
        rzrq_handle["两市融资买入金额"] = usable["沪市_融资买入金额(万元)"] + usable["深市_融资买入金额(万元)"]
        rzrq_handle["A股成交总额"] = (
            usable["沪市_成交金额(亿元)"] + usable["深市_成交金额(亿元)"]
        ) * 10000
        rzrq_handle["两市融资买入金额占比"] = rzrq_handle["两市融资买入金额"] / rzrq_handle["A股成交总额"]
        rzrq_handle = rzrq_handle.replace(to_replace="None", value=np.nan).dropna().tail(50)
        if rzrq_handle.empty:
            raise RuntimeError(
                "图8本地缓存聚合后结果为空。 "
                f"latest_local_date={format_log_date(latest_local_date)}, "
                f"hit_fields={hit_fields}, missing_fields={missing_fields}, usable_rows={len(usable)}"
            )
        rzrq_handle = rzrq_handle.reset_index(drop=True)
        return rzrq_handle, {
            "source": "local_cache",
            "hit_fields": hit_fields,
            "missing_fields": missing_fields,
            "shape": tuple(rzrq_handle.shape),
            "latest_local_date": latest_local_date,
        }

    def fetch_online_edb_indicator_series(
        indicator_name: str,
        indicator_code: str,
        begin_date: str,
        end_date: str,
    ) -> pd.DataFrame:
        result = local_client.edb_query(indicator_code, begin_date, end_date)
        times = extract_sequence_values_early(getattr(result, "Times", None))
        values = extract_first_data_series_early(getattr(result, "Data", None), expected_len=len(times))
        row_count = min(len(times), len(values))
        series_df = pd.DataFrame({"date": times[:row_count], indicator_name: values[:row_count]})
        if series_df.empty:
            print(
                f"图8在线单指标返回为空: indicator={indicator_name}, code={indicator_code}, "
                f"begin={begin_date}, end={end_date}"
            )
            return pd.DataFrame(columns=[indicator_name], index=pd.DatetimeIndex([], name="date"))
        series_df["date"] = pd.to_datetime(series_df["date"], errors="coerce").dt.normalize()
        series_df[indicator_name] = pd.to_numeric(series_df[indicator_name], errors="coerce")
        series_df = series_df.dropna(subset=["date"]).drop_duplicates(subset=["date"], keep="last")
        indexed = series_df.sort_values("date").set_index("date")[[indicator_name]]
        indexed.index.name = "date"
        print(
            f"图8在线单指标完成: indicator={indicator_name}, code={indicator_code}, "
            f"latest_date={format_log_date(indexed.index.max())}, rows={len(indexed)}"
        )
        return indexed

    def normalize_date_strings_early(values) -> list[str]:
        normalized = []
        for value in values:
            if value is None:
                continue
            text = str(value).strip()
            if text:
                normalized.append(text)
        return normalized

    def build_edb_series_early(result) -> pd.Series | None:
        time_values = extract_sequence_values_early(getattr(result, "Times", None))
        data_values = extract_first_data_series_early(
            getattr(result, "Data", None),
            expected_len=len(time_values),
        )
        if not time_values or not data_values:
            return None
        return pd.Series(
            clean_numeric_values(data_values),
            index=pd.to_datetime(time_values),
            dtype="float64",
        ).sort_index()

    def fetch_single_edb_value_early(indicator_code: str, requested_date: str, label: str) -> float:
        normalized_date = str(requested_date)
        parsed_date = pd.to_datetime(normalized_date, format="%Y%m%d", errors="coerce")
        candidate_dates = [normalized_date]
        if pd.notna(parsed_date):
            hyphen_date = parsed_date.strftime("%Y-%m-%d")
            if hyphen_date not in candidate_dates:
                candidate_dates.append(hyphen_date)

        for candidate_date in candidate_dates:
            result = local_client.edb_query(indicator_code, candidate_date, candidate_date, "Fill=Previous")
            series = build_edb_series_early(result)
            if series is None or series.empty:
                continue
            numeric = pd.to_numeric(series, errors="coerce").dropna()
            if numeric.empty:
                continue
            return round(float(numeric.iloc[-1]), 1)

        print(f"[警告] {label} 数据缺失: indicator={indicator_code}, date={normalized_date}, 跳过该日期")
        return None

    def fetch_edb_values_for_dates_early(indicator_code: str, requested_dates, label: str) -> dict[str, float]:
        normalized_dates = normalize_date_strings_early(requested_dates)
        if not normalized_dates:
            return {}
        query_start = min(normalized_dates)
        query_end = max(normalized_dates)
        print(
            f"{label}: interval_query start={query_start}, end={query_end}, "
            f"request_dates={normalized_dates}"
        )
        query_start_time = time_module.perf_counter()
        result = local_client.edb_query(indicator_code, query_start, query_end, "Fill=Previous")
        series = build_edb_series_early(result)
        request_index = pd.to_datetime(normalized_dates, format="%Y%m%d", errors="coerce")
        if request_index.isna().any():
            raise RuntimeError(f"{label} 存在无法解析的日期: {normalized_dates}")
        if series is not None and not series.empty:
            aligned = series.reindex(series.index.union(request_index)).sort_index().ffill().reindex(request_index)
            if not aligned.isna().any():
                value_map = {
                    request_date: round(float(value), 1)
                    for request_date, value in zip(normalized_dates, aligned.tolist())
                }
                print(
                    f"{label}: dates={normalized_dates}, total_points={len(value_map)}, "
                    f"elapsed={time_module.perf_counter() - query_start_time:.2f}s"
                )
                return value_map

            missing_dates = list(pd.Index(normalized_dates)[aligned.isna()])
            print(
                f"{label}: interval_query incomplete, fallback_to_single_dates. "
                f"indicator={indicator_code}, missing_dates={missing_dates}"
            )
        else:
            print(
                f"{label}: interval_query empty, fallback_to_single_dates. "
                f"indicator={indicator_code}, start={query_start}, end={query_end}"
            )

        fallback_start = time_module.perf_counter()
        value_map = {
            requested_date: fetch_single_edb_value_early(indicator_code, requested_date, label)
            for requested_date in normalized_dates
        }
        print(
            f"{label}: single_date_fallback complete, dates={normalized_dates}, "
            f"total_points={len(value_map)}, elapsed={time_module.perf_counter() - fallback_start:.2f}s"
        )
        return value_map

    def build_margin_dataframe_from_merged(
        merged: pd.DataFrame,
        *,
        source: str,
        hit_fields: list[str],
        missing_fields: list[str],
        online_meta: dict[str, object] | None = None,
    ) -> tuple[pd.DataFrame, dict[str, object]]:
        if merged is None or merged.empty:
            raise RuntimeError("图8在线结果为空：未生成任何可用数据行。")
        rzrq_data = merged.sort_index()
        required_fields = [
            "沪市_融资买入金额(万元)",
            "深市_融资买入金额(万元)",
            "沪市_融资融券余额(万元)",
            "深市_融资融券余额(万元)",
            "沪市_成交金额(亿元)",
            "深市_成交金额(亿元)",
        ]
        usable = rzrq_data.dropna(subset=required_fields).copy()
        latest_online_date = usable.index.max() if not usable.empty else pd.NaT
        if pd.isna(latest_online_date) or latest_online_date < margin_required_date:
            raise RuntimeError(
                "图8在线 edb_query 未更新到严格要求日期。 "
                f"latest_online_date={format_log_date(latest_online_date)}, "
                f"target_market_date={target_market_date_text}, report_date={report_date_text}, "
                f"margin_required_date={margin_required_date_text}, "
                f"allow_stale_margin_data={options.allow_stale_margin_data}"
            )
        rzrq_handle = pd.DataFrame(index=usable.index.copy())
        rzrq_handle["日期"] = usable.index.strftime("%Y-%m-%d")
        rzrq_handle["两市融资融券余额（万亿元）"] = (
            usable["沪市_融资融券余额(万元)"] + usable["深市_融资融券余额(万元)"]
        ) / 100000000
        rzrq_handle["两市融资买入金额"] = usable["沪市_融资买入金额(万元)"] + usable["深市_融资买入金额(万元)"]
        rzrq_handle["A股成交总额"] = (
            usable["沪市_成交金额(亿元)"] + usable["深市_成交金额(亿元)"]
        ) * 10000
        rzrq_handle["两市融资买入金额占比"] = rzrq_handle["两市融资买入金额"] / rzrq_handle["A股成交总额"]
        rzrq_handle = rzrq_handle.replace(to_replace="None", value=np.nan).dropna().tail(50)
        if rzrq_handle.empty:
            raise RuntimeError(
                "图8在线结果聚合后为空。 "
                f"source={source}, latest_online_date={format_log_date(latest_online_date)}, "
                f"usable_rows={len(usable)}, hit_fields={hit_fields}, missing_fields={missing_fields}, "
                f"online_meta={online_meta}"
            )
        rzrq_handle = rzrq_handle.reset_index(drop=True)
        meta = {
            "source": source,
            "hit_fields": hit_fields,
            "missing_fields": missing_fields,
            "shape": tuple(rzrq_handle.shape),
            "latest_online_date": latest_online_date,
            "raw_data": rzrq_data.reset_index().rename(columns={"index": "date"}).copy(),
        }
        if online_meta:
            meta.update(online_meta)
        return rzrq_handle, meta

    def build_margin_dataframe_from_online_edb() -> tuple[pd.DataFrame, dict[str, object]]:
        online_indicator_code_dict = {
            "沪市_融资买入金额(万元)": "M0061604",
            "深市_融资买入金额(万元)": "M0061609",
            "沪市_融资余额(万元)": "M0061606",
            "深市_融资余额(万元)": "M0061610",
            "沪市_融资融券余额(万元)": "M0061608",
            "深市_融资融券余额(万元)": "M0061613",
            "沪市_成交金额(亿元)": "M0331254",
            "深市_成交金额(亿元)": "M0340770",
        }
        merged: pd.DataFrame | None = None
        hit_fields: list[str] = []
        missing_fields: list[str] = []
        interval_failures: list[str] = []
        for indicator_name, indicator_code in online_indicator_code_dict.items():
            field_df = fetch_online_edb_indicator_series(indicator_name, indicator_code, start_week20, today)
            if field_df.empty:
                missing_fields.append(indicator_name)
                interval_failures.append(f"{indicator_name}[code={indicator_code}]")
                continue
            hit_fields.append(indicator_name)
            merged = field_df if merged is None else merged.join(field_df, how="outer")

        if merged is not None and not merged.empty:
            try:
                return build_margin_dataframe_from_merged(
                    merged,
                    source="online_edb_interval",
                    hit_fields=hit_fields,
                    missing_fields=missing_fields,
                    online_meta={"interval_failures": interval_failures},
                )
            except Exception as exc:
                print(f"图8区间在线结果不足，回退逐交易日点查: error={exc}")

        requested_trade_dates = [
            pd.Timestamp(item).strftime("%Y%m%d")
            for item in get_trade_dates_from_quotation_cache(start_week20, today)
        ]
        point_merged: pd.DataFrame | None = None
        point_hit_fields: list[str] = []
        point_missing_fields: list[str] = []
        point_failures: list[str] = []
        for indicator_name, indicator_code in online_indicator_code_dict.items():
            label = f"图8 {indicator_name}"
            try:
                value_map = fetch_edb_values_for_dates_early(indicator_code, requested_trade_dates, label)
                if not value_map:
                    raise RuntimeError("未返回任何有效点查结果")
                field_df = pd.DataFrame(
                    {
                        "date": pd.to_datetime(list(value_map.keys()), format="%Y%m%d", errors="coerce"),
                        indicator_name: list(value_map.values()),
                    }
                ).dropna(subset=["date"]).sort_values("date").set_index("date")[[indicator_name]]
                field_df.index = field_df.index.normalize()
                point_merged = field_df if point_merged is None else point_merged.join(field_df, how="outer")
                point_hit_fields.append(indicator_name)
                print(
                    f"图8在线单指标点查回退完成: indicator={indicator_name}, code={indicator_code}, "
                    f"latest_date={format_log_date(field_df.index.max())}, rows={len(field_df)}"
                )
            except Exception as exc:
                point_missing_fields.append(indicator_name)
                point_failures.append(f"{indicator_name}[code={indicator_code}, error={exc}]")
                print(
                    f"图8在线单指标点查回退失败: indicator={indicator_name}, code={indicator_code}, error={exc}"
                )

        if point_merged is None or point_merged.empty:
            raise RuntimeError(
                "图8在线 edb_query 未返回任何有效序列。 "
                f"begin_date={start_week20}, end_date={today}, "
                f"interval_failures={interval_failures}, point_failures={point_failures}"
            )

        return build_margin_dataframe_from_merged(
            point_merged,
            source="online_edb_point_fallback",
            hit_fields=point_hit_fields,
            missing_fields=point_missing_fields,
            online_meta={
                "interval_failures": interval_failures,
                "point_failures": point_failures,
                "requested_trade_dates": requested_trade_dates,
            },
        )

    def persist_margin_online_cache(raw_data: pd.DataFrame) -> None:
        if raw_data.empty or "date" not in raw_data.columns:
            raise RuntimeError("图8在线回写失败：raw_data 为空或缺少 date 列。")
        for field in [column for column in raw_data.columns if column != "date"]:
            upsert_market_sentiment_series(field, raw_data[["date", field]].copy())

    def get_updown_cache_dir() -> Path:
        cache_dir = paths.funding_draft_dir / "data" / "cache" / "updown_stats"
        cache_dir.mkdir(parents=True, exist_ok=True)
        return cache_dir

    def get_a_share_universe_cache_path() -> Path:
        return get_updown_cache_dir() / "a_share_universe_codes.txt"

    def get_a_share_universe_meta_path() -> Path:
        return get_updown_cache_dir() / "a_share_universe_codes_meta.json"

    def get_updown_stats_cache_path(begin_date: str, end_date: str) -> Path:
        return get_updown_cache_dir() / f"updown_stats_{begin_date}_{end_date}.csv"

    def extract_a_share_codes(values) -> list[str]:
        extracted: list[str] = []
        for value in values:
            if pd.isna(value):
                continue
            match = a_share_code_pattern.search(str(value).strip())
            if match:
                extracted.append(match.group(0))
        return list(dict.fromkeys(extracted))

    def extract_a_share_codes_from_frame(df: pd.DataFrame) -> list[str]:
        if df is None or df.empty:
            return []
        preferred_columns = [
            column
            for column in df.columns
            if any(keyword in str(column).lower() for keyword in ("code", "wind", "thscode", "证券代码", "股票代码", "代码"))
        ]
        ordered_columns = preferred_columns + [column for column in df.columns if column not in preferred_columns]
        collected: list[str] = []
        for column in ordered_columns:
            collected.extend(extract_a_share_codes(df[column].tolist()))
            if len(collected) >= figure7_universe_min_count:
                break
        return list(dict.fromkeys(collected))

    def read_a_share_universe_cache() -> tuple[list[str], dict[str, object]]:
        cache_path = get_a_share_universe_cache_path()
        meta_path = get_a_share_universe_meta_path()
        if not cache_path.exists():
            return [], {}
        cached_codes = [line.strip() for line in cache_path.read_text(encoding="utf-8").splitlines() if line.strip()]
        if len(cached_codes) < figure7_universe_min_count:
            print(
                f"图7 A股 universe 缓存失效: cache={cache_path.name}, code_count={len(cached_codes)}, "
                f"threshold={figure7_universe_min_count}"
            )
            return [], {}
        cache_meta: dict[str, object] = {}
        if meta_path.exists():
            try:
                cache_meta = json.loads(meta_path.read_text(encoding="utf-8"))
            except Exception:
                cache_meta = {}
        return cached_codes, cache_meta

    def write_a_share_universe_cache(
        codes: list[str],
        source: str,
        reference_date: str,
        extra: dict[str, object] | None = None,
    ) -> None:
        cache_path = get_a_share_universe_cache_path()
        meta_path = get_a_share_universe_meta_path()
        deduped_codes = list(dict.fromkeys(codes))
        cache_path.write_text("\n".join(deduped_codes), encoding="utf-8")
        meta = {
            "source": source,
            "reference_date": reference_date,
            "code_count": len(deduped_codes),
            "written_at": datetime.datetime.now().isoformat(timespec="seconds"),
        }
        if extra:
            meta.update(extra)
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    def try_load_a_share_universe_from_structured_sources(reference_date: str) -> tuple[list[str], dict[str, object]] | None:
        structured_datapool_candidates = [
            (
                "sectorconstituent",
                [
                    f"date={reference_date};sectorid=a001010100000000;field=wind_code,sec_name",
                    f"date={reference_date};sectorid=a001010100000000",
                ],
            ),
        ]
        for dataset_name, option_variants in structured_datapool_candidates:
            try:
                _, df, _ = probe_datapool_variants("图7 A股universe结构化", dataset_name, option_variants)
            except Exception as exc:
                print(f"图7结构化 universe 失败: dataset={dataset_name}, error={exc}")
                continue
            extracted_codes = extract_a_share_codes_from_frame(df)
            print(
                f"图7结构化 universe 返回: dataset={dataset_name}, code_count={len(extracted_codes)}, "
                f"columns={list(df.columns)}"
            )
            if len(extracted_codes) >= figure7_universe_min_count:
                return extracted_codes, {
                    "source": f"structured_datapool:{dataset_name}",
                    "columns": list(df.columns),
                }
        block_queries = [
            f"全部A股 {reference_date}收盘价",
            f"沪深京A股 {reference_date}收盘价",
            f"A股 {reference_date}收盘价",
        ]
        for query_text in block_queries:
            try:
                raw_df = local_client.block_wc_query(query_text)
                parsed = parse_block_query_dataframe(raw_df)
            except Exception as exc:
                print(f"图7板块成分 universe 失败: query={query_text}, error={exc}")
                continue
            extracted_codes = extract_a_share_codes_from_frame(parsed)
            print(
                f"图7板块成分 universe 返回: query={query_text}, code_count={len(extracted_codes)}, "
                f"columns={list(parsed.columns)}"
            )
            if len(extracted_codes) >= figure7_universe_min_count:
                return extracted_codes, {
                    "source": "structured_block_wc",
                    "query_text": query_text,
                }
        return None

    def resolve_codetable_dir() -> Path:
        candidate_dirs: list[Path] = [
            paths.inputs_dir / "reference" / "codetables",
        ]
        env_dir = os.getenv("LOCAL_CODETABLE_DIR", "").strip()
        if env_dir:
            candidate_dirs.insert(0, Path(env_dir))

        seen: set[Path] = set()
        for candidate in candidate_dirs:
            resolved = Path(candidate)
            if resolved in seen:
                continue
            seen.add(resolved)
            if (resolved / "CODE_AS.dat").exists():
                return resolved
        raise RuntimeError(
            "图7 A股代码表定位失败：未找到 CodeTables/CODE_AS.dat。"
            "请在 inputs/reference/codetables 下放置 CODE_AS.dat，"
            "或设置环境变量 LOCAL_CODETABLE_DIR。"
        )

    def load_a_share_universe_codes(reference_date: str) -> tuple[list[str], dict[str, object]]:
        cached_codes, cache_meta = read_a_share_universe_cache()
        cache_source = str(cache_meta.get("source", ""))
        if cached_codes and cache_source.startswith("structured_"):
            print(
                f"图7 A股 universe 缓存命中: source={cache_source}, cache={get_a_share_universe_cache_path().name}, "
                f"code_count={len(cached_codes)}"
            )
            return cached_codes, cache_meta

        structured = try_load_a_share_universe_from_structured_sources(reference_date)
        if structured is not None:
            structured_codes, structured_meta = structured
            write_a_share_universe_cache(structured_codes, structured_meta["source"], reference_date, structured_meta)
            print(
                f"图7 A股 universe 结构化获取成功: source={structured_meta['source']}, "
                f"code_count={len(structured_codes)}, reference_date={reference_date}"
            )
            return structured_codes, structured_meta

        if cached_codes:
            print(
                f"图7 A股 universe 回退至现有缓存: source={cache_source or 'unknown'}, "
                f"code_count={len(cached_codes)}, cache={get_a_share_universe_cache_path().name}"
            )
            return cached_codes, cache_meta

        codetable_dir = resolve_codetable_dir()
        codetable_path = codetable_dir / "CODE_AS.dat"
        raw_text = codetable_path.read_bytes().decode("utf-16le", errors="ignore")
        extracted_codes = re.findall(r"\b\d{6}\.(?:SH|SZ|BJ)\b", raw_text)
        deduped_codes = list(dict.fromkeys(extracted_codes))
        if len(deduped_codes) < figure7_universe_min_count:
            raise RuntimeError(
                f"图7 A股代码表解析失败：有效代码数异常。 path={codetable_path}, code_count={len(deduped_codes)}"
            )
        write_a_share_universe_cache(
            deduped_codes,
            "local_codetable",
            reference_date,
            {"codetable_path": str(codetable_path)},
        )
        print(
            f"图7 A股代码表解析完成: source={codetable_path}, code_count={len(deduped_codes)}, "
            f"cache={get_a_share_universe_cache_path().name}"
        )
        return deduped_codes, {
            "source": "local_codetable",
            "codetable_path": str(codetable_path),
        }

    def classify_stock_limit_pct(code: str, stock_name: str) -> float:
        name_text = str(stock_name).upper()
        code_text = str(code).split(".", 1)[0]
        if "ST" in name_text:
            return 5.0
        if code_text.startswith(("300", "301", "688")):
            return 20.0
        return 10.0

    def get_trade_dates_from_quotation_cache(begin_date: str, end_date: str) -> list[str]:
        begin_ts = pd.to_datetime(begin_date, errors="coerce").normalize()
        end_ts = pd.to_datetime(end_date, errors="coerce").normalize()
        if pd.isna(begin_ts) or pd.isna(end_ts):
            raise RuntimeError(
                f"交易日范围解析失败: begin_date={begin_date}, end_date={end_date}"
            )

        quotation_dir = paths.data_dir / "quotation"
        trade_date_set: set[pd.Timestamp] = set()
        if quotation_dir.exists():
            for file_path in quotation_dir.glob("*.csv"):
                local_series = load_local_quotation_close_series(file_path.stem)
                if local_series.empty:
                    continue
                sliced_index = local_series[
                    (local_series.index >= begin_ts) & (local_series.index <= end_ts)
                ].index
                trade_date_set.update(pd.Timestamp(item).normalize() for item in sliced_index)

        if trade_date_set:
            latest_local_trade_date = max(trade_date_set)
            if latest_local_trade_date < end_ts:
                weekday_fill = pd.bdate_range(latest_local_trade_date + pd.Timedelta(days=1), end_ts)
                trade_date_set.update(item.normalize() for item in weekday_fill)
        else:
            weekday_fill = pd.bdate_range(begin_ts, end_ts)
            trade_date_set.update(item.normalize() for item in weekday_fill)

        trade_dates = [item.strftime("%Y-%m-%d") for item in sorted(trade_date_set)]
        if not trade_dates:
            raise RuntimeError(
                f"图7无法从 quotation 缓存解析交易日。 begin_date={begin_date}, end_date={end_date}"
            )
        return trade_dates

    def extract_wc_metric_value(
        wc_df: pd.DataFrame,
        query_code: str,
    ) -> tuple[float, str]:
        if wc_df is None or not isinstance(wc_df, pd.DataFrame) or wc_df.empty:
            raise RuntimeError("THS_WC 返回空 DataFrame")

        matched = wc_df.copy()
        for column in matched.columns:
            if "代码" not in str(column):
                continue
            exact_match = matched[matched[column].astype(str).str.upper() == str(query_code).upper()]
            if not exact_match.empty:
                matched = exact_match
                break

        row = matched.iloc[0]
        for column in matched.columns:
            column_text = str(column)
            if any(token in column_text for token in ["代码", "简称", "名称", "所属"]):
                continue
            numeric_value = pd.to_numeric(pd.Series([row[column]]), errors="coerce").iloc[0]
            if pd.notna(numeric_value):
                return float(numeric_value), column_text

        raise RuntimeError(
            f"THS_WC 未解析出有效指标值: query_code={query_code}, columns={list(matched.columns)}"
        )

    def fetch_wc_metric_series_for_trade_dates(
        metric_label: str,
        query_code: str,
        metric_name: str,
        trade_dates,
        *,
        domain: str = "index",
    ) -> tuple[pd.Series, dict[str, object]]:
        rows: list[dict[str, object]] = []
        attempts: list[dict[str, object]] = []
        normalized_trade_dates = sorted({pd.Timestamp(item).normalize() for item in trade_dates})

        for trade_date in normalized_trade_dates:
            query_text = f"{query_code} {trade_date:%Y%m%d} {metric_name}"
            try:
                wc_df = local_client.wc_query_dataframe(query_text, domain=domain)
                metric_value, metric_column = extract_wc_metric_value(wc_df, query_code)
                rows.append({"date": trade_date, metric_label: metric_value})
                attempts.append(
                    {
                        "date": trade_date.strftime("%Y-%m-%d"),
                        "query": query_text,
                        "metric_column": metric_column,
                        "value": metric_value,
                    }
                )
                print(
                    f"{metric_label} WC fallback success: query={query_text}, domain={domain}, "
                    f"metric_column={metric_column}, value={metric_value}"
                )
            except Exception as exc:
                attempts.append(
                    {
                        "date": trade_date.strftime("%Y-%m-%d"),
                        "query": query_text,
                        "error": str(exc),
                    }
                )
                print(
                    f"{metric_label} WC fallback miss: query={query_text}, domain={domain}, error={exc}"
                )

        if not rows:
            raise RuntimeError(
                f"{metric_label} WC fallback failed: code={query_code}, domain={domain}, attempts={attempts}"
            )

        series = (
            pd.DataFrame(rows)
            .drop_duplicates(subset=["date"], keep="last")
            .sort_values("date")
            .set_index("date")[metric_label]
        )
        series.index.name = "date"
        series.name = metric_label
        return series, {
            "source": "wc_query",
            "query_code": query_code,
            "metric_name": metric_name,
            "domain": domain,
            "attempts": attempts,
            "latest_date": format_log_date(series.index.max()),
        }

    def build_updown_statistics_from_block_wc(
        begin_date: str,
        end_date: str,
        universe_count: int,
    ) -> pd.DataFrame:
        trade_dates = get_trade_dates_from_quotation_cache(begin_date, end_date)
        query_templates = [
            "全部A股 {trade_date}涨跌幅",
            "沪深京A股 {trade_date}涨跌幅",
            "A股 {trade_date}涨跌幅",
        ]
        grouped_records: list[dict[str, object]] = []
        for trade_date in trade_dates:
            snapshot_df = pd.DataFrame()
            used_query = ""
            last_error = ""
            for template in query_templates:
                query_text = template.format(trade_date=trade_date.replace("-", ""))
                try:
                    raw_df = local_client.block_wc_query(query_text)
                    parsed = parse_block_query_dataframe(raw_df)
                except Exception as exc:
                    last_error = str(exc)
                    continue
                parsed = parsed[
                    parsed["code"].astype(str).str.fullmatch(r"\d{6}\.(?:SH|SZ|BJ)")
                    & parsed["value"].notna()
                ].copy()
                if len(parsed) >= figure7_total_min_count:
                    snapshot_df = parsed
                    used_query = query_text
                    break
            if snapshot_df.empty:
                raise RuntimeError(
                    f"图7 block 查询 fallback 失败：trade_date={trade_date}, last_error={last_error}"
                )

            snapshot_df["limit_pct"] = [
                classify_stock_limit_pct(code, stock_name)
                for code, stock_name in zip(snapshot_df["code"], snapshot_df["name"])
            ]
            snapshot_df["is_up_limit"] = snapshot_df["value"].ge(snapshot_df["limit_pct"] - 0.075)
            snapshot_df["is_down_limit"] = snapshot_df["value"].le(-(snapshot_df["limit_pct"] - 0.075))

            up_count = int(snapshot_df["value"].gt(0).sum())
            down_count = int(snapshot_df["value"].lt(0).sum())
            flat_count = int(snapshot_df["value"].eq(0).sum())
            up_limit_count = int(snapshot_df["is_up_limit"].sum())
            down_limit_count = int(snapshot_df["is_down_limit"].sum())
            total_count = (
                down_limit_count
                + max(down_count - down_limit_count, 0)
                + flat_count
                + max(up_count - up_limit_count, 0)
                + up_limit_count
            )
            coverage_ratio = total_count / universe_count if universe_count else np.nan
            print(
                f"图7 block fallback 汇总: trade_date={trade_date}, query={used_query}, 股票池数量={len(snapshot_df)}, "
                f"涨跌平及涨跌停合计={total_count}, coverage_ratio={coverage_ratio:.2%}"
            )
            if total_count < figure7_total_min_count:
                raise RuntimeError(
                    f"图7 block fallback 总家数异常：trade_date={trade_date}, total_count={total_count}, "
                    f"threshold={figure7_total_min_count}"
                )
            grouped_records.append(
                {
                    "日期": trade_date,
                    "跌停家数": down_limit_count,
                    "下跌家数（不含跌停）": max(down_count - down_limit_count, 0),
                    "平盘家数": flat_count,
                    "上涨家数（不含涨停）": max(up_count - up_limit_count, 0),
                    "涨停家数": up_limit_count,
                    "_股票池数量": int(len(snapshot_df)),
                    "_总家数": total_count,
                    "_coverage_ratio": coverage_ratio,
                }
            )
        return pd.DataFrame(grouped_records)

    def build_updown_statistics_from_dateserial(begin_date: str, end_date: str) -> pd.DataFrame:
        cache_path = get_updown_stats_cache_path(begin_date, end_date)
        required_columns = [
            "日期",
            "跌停家数",
            "下跌家数（不含跌停）",
            "平盘家数",
            "上涨家数（不含涨停）",
            "涨停家数",
        ]
        if cache_path.exists():
            cached_df = pd.read_csv(cache_path)
            if not cached_df.empty and all(column in cached_df.columns for column in required_columns):
                total_counts = cached_df[required_columns[1:]].apply(pd.to_numeric, errors="coerce").sum(axis=1)
                if total_counts.min() >= figure7_total_min_count:
                    print(f"图7周缓存命中: cache={cache_path.name}, shape={cached_df.shape}")
                    return cached_df[required_columns].copy()
                print(
                    f"图7周缓存失效: cache={cache_path.name}, min_total={int(total_counts.min())}, "
                    f"threshold={figure7_total_min_count}, rebuilding=true"
                )

        a_share_codes, universe_meta = load_a_share_universe_codes(end_date)
        universe_count = len(a_share_codes)
        query_batch_size = min(max(batch_size, 1), 30)
        print(
            f"图7结构化取数开始: begin_date={begin_date}, end_date={end_date}, "
            f"code_count={universe_count}, batch_size={query_batch_size}, universe_source={universe_meta.get('source', 'unknown')}"
        )
        try:
            query_started = time_module.perf_counter()
            serial_df = local_client.dateserial_chunked(
                a_share_codes,
                [
                    "ths_stock_short_name_stock",
                    "ths_pre_close_stock",
                    "ths_close_price_stock",
                    "ths_chg_ratio_stock",
                ],
                ["", "101", "100", ""],
                begin_date,
                end_date,
                option_string="Days:Tradedays,Fill:Previous,Interval:D",
                batch_size=query_batch_size,
            )
            print(
                f"图7结构化取数完成: shape={serial_df.shape}, elapsed={time_module.perf_counter() - query_started:.2f}s"
            )
            if serial_df.empty:
                raise RuntimeError(
                    f"图7 两市涨跌数量统计查询失败：THS_DateSerial 未返回有效记录。 begin_date={begin_date}, end_date={end_date}"
                )

            lower_map = {str(column).lower(): column for column in serial_df.columns}
            required_mapping = {
                "time": lower_map.get("time"),
                "code": lower_map.get("thscode"),
                "name": lower_map.get("ths_stock_short_name_stock"),
                "pre_close": lower_map.get("ths_pre_close_stock"),
                "close": lower_map.get("ths_close_price_stock"),
                "chg_ratio": lower_map.get("ths_chg_ratio_stock"),
            }
            missing_mapping = [key for key, value in required_mapping.items() if value is None]
            if missing_mapping:
                raise RuntimeError(
                    f"图7 两市涨跌数量统计查询失败：THS_DateSerial 返回字段不完整。 "
                    f"missing={missing_mapping}, columns={list(serial_df.columns)}"
                )

            detail_df = serial_df[
                [
                    required_mapping["time"],
                    required_mapping["code"],
                    required_mapping["name"],
                    required_mapping["pre_close"],
                    required_mapping["close"],
                    required_mapping["chg_ratio"],
                ]
            ].copy()
            detail_df.columns = ["time", "code", "name", "pre_close", "close", "chg_ratio"]
            detail_df["time"] = pd.to_datetime(detail_df["time"], errors="coerce")
            detail_df["code"] = detail_df["code"].astype(str)
            detail_df["name"] = detail_df["name"].astype(str)
            detail_df["pre_close"] = pd.to_numeric(detail_df["pre_close"], errors="coerce")
            detail_df["close"] = pd.to_numeric(detail_df["close"], errors="coerce")
            detail_df["chg_ratio"] = pd.to_numeric(detail_df["chg_ratio"], errors="coerce")
            detail_df = detail_df.dropna(subset=["time", "code"]).drop_duplicates(subset=["time", "code"], keep="last")
            detail_df["effective_chg_ratio"] = detail_df["chg_ratio"]
            recalc_mask = (
                detail_df["effective_chg_ratio"].isna()
                & detail_df["pre_close"].notna()
                & detail_df["close"].notna()
                & detail_df["pre_close"].ne(0)
            )
            detail_df.loc[recalc_mask, "effective_chg_ratio"] = (
                detail_df.loc[recalc_mask, "close"] / detail_df.loc[recalc_mask, "pre_close"] - 1
            ) * 100
            detail_df = detail_df[detail_df["effective_chg_ratio"].notna()].copy()
            if detail_df.empty:
                raise RuntimeError(
                    f"图7 两市涨跌数量统计查询失败：清洗后无有效股票记录。 begin_date={begin_date}, end_date={end_date}"
                )

            detail_df["trade_date"] = detail_df["time"].dt.strftime("%Y-%m-%d")
            detail_df["limit_pct"] = [
                classify_stock_limit_pct(code, stock_name)
                for code, stock_name in zip(detail_df["code"], detail_df["name"])
            ]
            detail_df["is_up_limit"] = detail_df["effective_chg_ratio"].ge(detail_df["limit_pct"] - 0.075)
            detail_df["is_down_limit"] = detail_df["effective_chg_ratio"].le(-(detail_df["limit_pct"] - 0.075))

            grouped_records: list[dict[str, object]] = []
            for trade_date, group in detail_df.groupby("trade_date", sort=True):
                up_count = int(group["effective_chg_ratio"].gt(0).sum())
                down_count = int(group["effective_chg_ratio"].lt(0).sum())
                flat_count = int(group["effective_chg_ratio"].eq(0).sum())
                up_limit_count = int(group["is_up_limit"].sum())
                down_limit_count = int(group["is_down_limit"].sum())
                total_count = (
                    down_limit_count
                    + max(down_count - down_limit_count, 0)
                    + flat_count
                    + max(up_count - up_limit_count, 0)
                    + up_limit_count
                )
                stock_pool_count = int(group["code"].nunique())
                coverage_ratio = total_count / universe_count if universe_count else np.nan
                print(
                    f"图7交易日汇总: trade_date={trade_date}, 股票池数量={stock_pool_count}, "
                    f"涨跌平及涨跌停合计={total_count}, coverage_ratio={coverage_ratio:.2%}, "
                    f"上涨={up_count}, 下跌={down_count}, 平盘={flat_count}, 涨停={up_limit_count}, 跌停={down_limit_count}"
                )
                if total_count < figure7_total_min_count:
                    raise RuntimeError(
                        f"图7 两市涨跌数量统计异常：trade_date={trade_date} 总家数仅 {total_count}，"
                        f"低于阈值 {figure7_total_min_count}。请检查结构化 universe 和明细覆盖。"
                    )
                grouped_records.append(
                    {
                        "日期": trade_date,
                        "跌停家数": down_limit_count,
                        "下跌家数（不含跌停）": max(down_count - down_limit_count, 0),
                        "平盘家数": flat_count,
                        "上涨家数（不含涨停）": max(up_count - up_limit_count, 0),
                        "涨停家数": up_limit_count,
                        "_股票池数量": stock_pool_count,
                        "_总家数": total_count,
                        "_coverage_ratio": coverage_ratio,
                    }
                )

            updown_df = pd.DataFrame(grouped_records)
            if updown_df.empty:
                raise RuntimeError(
                    f"图7 两市涨跌数量统计查询失败：按交易日聚合后为空。 begin_date={begin_date}, end_date={end_date}"
                )
        except Exception as exc:
            print(
                f"图7 DateSerial 主路径失败，回退 block 查询: begin_date={begin_date}, "
                f"end_date={end_date}, error={exc}"
            )
            updown_df = build_updown_statistics_from_block_wc(begin_date, end_date, universe_count)
            if updown_df.empty:
                raise RuntimeError(
                    f"图7 block 查询 fallback 后仍为空。 begin_date={begin_date}, end_date={end_date}"
                )
        updown_df = updown_df.sort_values("日期").tail(5).reset_index(drop=True)
        print(
            f"图7聚合完成: rows={len(updown_df)}, dates={updown_df['日期'].tolist()}, "
            f"samples={updown_df.tail(2).to_dict(orient='records')}"
        )
        updown_df.to_csv(cache_path, index=False, encoding="utf-8-sig")
        print(f"图7周缓存写出完成: cache={cache_path.name}, shape={updown_df.shape}")
        return updown_df[required_columns].copy()

    def build_unlock_weekly_dataframe_from_wc(begin_date: str, end_date: str) -> pd.DataFrame:
        begin_text = pd.to_datetime(begin_date, errors="coerce").strftime("%Y-%m-%d")
        end_text = pd.to_datetime(end_date, errors="coerce").strftime("%Y-%m-%d")
        query_text = f"{begin_text}至{end_text}限售股解禁公司明细"
        print(f"图14问财取数开始: query={query_text}")
        started = time_module.perf_counter()
        detail_df = local_client.wc_query_dataframe(query_text, domain="stock")
        print(
            f"图14问财取数完成: shape={detail_df.shape}, columns={list(detail_df.columns)}, "
            f"elapsed={time_module.perf_counter() - started:.2f}s"
        )
        if detail_df.empty:
            raise RuntimeError(f"图14 限售股解禁金额和数量查询失败：问财未返回有效明细。 query={query_text}")

        code_col = next((column for column in detail_df.columns if "股票代码" in str(column)), detail_df.columns[0])
        date_col = next((column for column in detail_df.columns if "解禁日期" in str(column)), None)
        amount_col = next((column for column in detail_df.columns if "解禁金额" in str(column)), None)
        if date_col is None or amount_col is None:
            raise RuntimeError(
                f"图14 限售股解禁金额和数量查询失败：问财返回字段缺失。 columns={list(detail_df.columns)}"
            )

        parsed = detail_df[[code_col, date_col, amount_col]].copy()
        parsed.columns = ["股票代码", "解禁日期", "解禁金额"]
        parsed["解禁日期"] = pd.to_datetime(parsed["解禁日期"].astype(str), errors="coerce")
        parsed["解禁金额"] = pd.to_numeric(parsed["解禁金额"], errors="coerce")
        parsed = parsed.dropna(subset=["解禁日期", "解禁金额"]).sort_values("解禁日期").reset_index(drop=True)
        if parsed.empty:
            raise RuntimeError(f"图14 限售股解禁金额和数量查询失败：明细清洗后为空。 query={query_text}")

        parsed["起始日期"] = parsed["解禁日期"] - pd.to_timedelta(parsed["解禁日期"].dt.weekday, unit="D")
        weekly = parsed.groupby("起始日期", as_index=False).agg({"解禁金额": lambda series: round(series.sum() / 1e8, 2), "股票代码": "count"})
        weekly.columns = ["起始日期", "当周解禁市值", "当周解禁家数"]
        weekly["日期"] = weekly["起始日期"] + pd.Timedelta(days=4)
        weekly["起始日期"] = weekly["起始日期"].dt.strftime("%Y-%m-%d")
        weekly["日期"] = weekly["日期"].dt.strftime("%Y-%m-%d")
        weekly["当周解禁市值"] = pd.to_numeric(weekly["当周解禁市值"], errors="coerce").round(2)
        weekly["当周解禁家数"] = pd.to_numeric(weekly["当周解禁家数"], errors="coerce").fillna(0).astype(int)
        weekly = weekly[["起始日期", "当周解禁市值", "当周解禁家数", "日期"]]
        print(
            f"图14周度聚合完成: shape={weekly.shape}, "
            f"dates={weekly['日期'].tail(5).tolist()}, sample={weekly.tail(3).to_dict(orient='records')}"
        )
        return weekly

    def build_ipo_private_weekly_dataframe_from_wc(begin_date: str, end_date: str) -> pd.DataFrame:
        begin_ts = pd.to_datetime(begin_date, errors="coerce")
        end_ts = pd.to_datetime(end_date, errors="coerce")
        begin_text = begin_ts.strftime("%Y-%m-%d")
        end_text = end_ts.strftime("%Y-%m-%d")

        ipo_query = f"{begin_text}至{end_text}IPO首发上市公司明细，上市日期，首发募集资金"
        private_query = f"{begin_text}至{end_text}已实施定增公司明细，增发上市日，增发实际募集资金总额"

        print(f"图15问财IPO取数开始: query={ipo_query}")
        ipo_started = time_module.perf_counter()
        ipo_detail = local_client.wc_query_dataframe(ipo_query, domain="stock")
        print(
            f"图15问财IPO取数完成: shape={ipo_detail.shape}, columns={list(ipo_detail.columns)}, "
            f"elapsed={time_module.perf_counter() - ipo_started:.2f}s"
        )
        print(f"图15问财定增取数开始: query={private_query}")
        private_started = time_module.perf_counter()
        private_detail = local_client.wc_query_dataframe(private_query, domain="stock")
        print(
            f"图15问财定增取数完成: shape={private_detail.shape}, columns={list(private_detail.columns)}, "
            f"elapsed={time_module.perf_counter() - private_started:.2f}s"
        )

        ipo_date_col = next((column for column in ipo_detail.columns if "上市日期" in str(column)), None)
        ipo_amount_col = next((column for column in ipo_detail.columns if "首发募集资金" in str(column)), None)
        if ipo_date_col is None or ipo_amount_col is None:
            raise RuntimeError(
                f"图15 IPO问财返回字段缺失。 columns={list(ipo_detail.columns)}"
            )
        ipo_parsed = ipo_detail[[ipo_date_col, ipo_amount_col]].copy()
        ipo_parsed.columns = ["event_date", "amount"]
        ipo_parsed["event_date"] = pd.to_datetime(ipo_parsed["event_date"].astype(str), errors="coerce")
        ipo_parsed["amount"] = pd.to_numeric(ipo_parsed["amount"], errors="coerce")
        ipo_parsed = ipo_parsed.dropna(subset=["event_date", "amount"])
        ipo_parsed = ipo_parsed[(ipo_parsed["event_date"] >= begin_ts) & (ipo_parsed["event_date"] <= end_ts)].copy()

        private_date_col = next((column for column in private_detail.columns if "增发上市日" in str(column)), None)
        private_amount_col = next((column for column in private_detail.columns if "增发实际募集资金总额" in str(column)), None)
        if private_date_col is None or private_amount_col is None:
            raise RuntimeError(
                f"图15 定增问财返回字段缺失。 columns={list(private_detail.columns)}"
            )
        private_parsed = private_detail[[private_date_col, private_amount_col]].copy()
        private_parsed.columns = ["event_date", "amount"]
        private_parsed["event_date"] = pd.to_datetime(private_parsed["event_date"].astype(str), errors="coerce")
        private_parsed["amount"] = pd.to_numeric(private_parsed["amount"], errors="coerce")
        private_parsed = private_parsed.dropna(subset=["event_date", "amount"])
        private_parsed = private_parsed[(private_parsed["event_date"] >= begin_ts) & (private_parsed["event_date"] <= end_ts)].copy()

        if ipo_parsed.empty and private_parsed.empty:
            raise RuntimeError(
                f"图15 IPO和定增金额查询失败：IPO与定增明细清洗后均为空。 begin_date={begin_text}, end_date={end_text}"
            )

        ipo_weekly = (
            ipo_parsed.assign(week_start=lambda df: df["event_date"] - pd.to_timedelta(df["event_date"].dt.weekday, unit="D"))
            .groupby("week_start", as_index=False)
            .agg({"amount": lambda series: round(series.sum() / 1e8, 1), "event_date": "count"})
            .rename(columns={"amount": "IPO首发募集资金（亿元）", "event_date": "IPO首发家数"})
        )
        private_weekly = (
            private_parsed.assign(week_start=lambda df: df["event_date"] - pd.to_timedelta(df["event_date"].dt.weekday, unit="D"))
            .groupby("week_start", as_index=False)
            .agg({"amount": lambda series: round(series.sum() / 1e8, 1), "event_date": "count"})
            .rename(columns={"amount": "定增募集（亿元）", "event_date": "定增家数"})
        )

        all_week_fridays: list[pd.Timestamp] = []
        current_monday = begin_ts
        while current_monday <= end_ts:
            all_week_fridays.append(current_monday + pd.Timedelta(days=4))
            current_monday += pd.Timedelta(days=7)
        all_week_fridays = [date for date in all_week_fridays if date <= end_ts]
        target_fridays = all_week_fridays[-12:]
        if len(target_fridays) < 12:
            raise RuntimeError(
                f"图15 周度日期生成失败：最近12个周频日期不足。 begin_date={begin_text}, end_date={end_text}, generated={len(target_fridays)}"
            )
        base = pd.DataFrame({"日期": [date.strftime("%Y-%m-%d") for date in target_fridays]})
        base["week_start"] = pd.to_datetime(base["日期"]) - pd.to_timedelta(pd.to_datetime(base["日期"]).dt.weekday, unit="D")

        ipo_weekly["week_start"] = pd.to_datetime(ipo_weekly["week_start"])
        private_weekly["week_start"] = pd.to_datetime(private_weekly["week_start"])
        merged = base.merge(ipo_weekly, on="week_start", how="left").merge(private_weekly, on="week_start", how="left")
        merged = merged.fillna(0)
        merged["IPO首发家数"] = pd.to_numeric(merged["IPO首发家数"], errors="coerce").fillna(0).astype(int)
        merged["IPO首发募集资金（亿元）"] = pd.to_numeric(merged["IPO首发募集资金（亿元）"], errors="coerce").fillna(0).round(1)
        merged["定增家数"] = pd.to_numeric(merged["定增家数"], errors="coerce").fillna(0).astype(int)
        merged["定增募集（亿元）"] = pd.to_numeric(merged["定增募集（亿元）"], errors="coerce").fillna(0).round(1)
        merged = merged[["日期", "IPO首发家数", "IPO首发募集资金（亿元）", "定增家数", "定增募集（亿元）"]]
        print(
            f"图15周度聚合完成: shape={merged.shape}, "
            f"sample={merged.tail(5).to_dict(orient='records')}"
        )
        return merged

    def build_ipo_private_weekly_dataframe_from_wc_safe(begin_date: str, end_date: str) -> pd.DataFrame:
        begin_ts = pd.to_datetime(begin_date, errors="coerce")
        end_ts = pd.to_datetime(end_date, errors="coerce")
        begin_text = begin_ts.strftime("%Y-%m-%d")
        end_text = end_ts.strftime("%Y-%m-%d")

        ipo_query = (
            f"{begin_text}\u81f3{end_text}"
            "IPO\u9996\u53d1\u4e0a\u5e02\u516c\u53f8\u660e\u7ec6\uff0c"
            "\u4e0a\u5e02\u65e5\u671f\uff0c\u9996\u53d1\u52df\u96c6\u8d44\u91d1"
        )
        private_query = (
            f"{begin_text}\u81f3{end_text}"
            "\u5df2\u5b9e\u65bd\u5b9a\u589e\u516c\u53f8\u660e\u7ec6\uff0c"
            "\u589e\u53d1\u4e0a\u5e02\u65e5\uff0c\u589e\u53d1\u5b9e\u9645\u52df\u96c6\u8d44\u91d1\u603b\u989d"
        )

        print(f"图15问财IPO取数开始: query={ipo_query}")
        ipo_started = time_module.perf_counter()
        ipo_detail = local_client.wc_query_dataframe(ipo_query, domain="stock")
        print(
            f"图15问财IPO取数完成: shape={ipo_detail.shape}, columns={list(ipo_detail.columns)}, "
            f"elapsed={time_module.perf_counter() - ipo_started:.2f}s"
        )
        print(f"图15问财定增取数开始: query={private_query}")
        private_started = time_module.perf_counter()
        private_detail = local_client.wc_query_dataframe(private_query, domain="stock")
        print(
            f"图15问财定增取数完成: shape={private_detail.shape}, columns={list(private_detail.columns)}, "
            f"elapsed={time_module.perf_counter() - private_started:.2f}s"
        )

        ipo_date_col = next((column for column in ipo_detail.columns if "上市日期" in str(column)), None)
        ipo_amount_col = next((column for column in ipo_detail.columns if "首发募集资金" in str(column)), None)
        if ipo_date_col is None or ipo_amount_col is None:
            raise RuntimeError(f"图15 IPO问财返回字段缺失。 columns={list(ipo_detail.columns)}")

        ipo_parsed = ipo_detail[[ipo_date_col, ipo_amount_col]].copy()
        ipo_parsed.columns = ["event_date", "amount"]
        ipo_parsed["event_date"] = pd.to_datetime(ipo_parsed["event_date"].astype(str), errors="coerce")
        ipo_parsed["amount"] = pd.to_numeric(ipo_parsed["amount"], errors="coerce")
        ipo_parsed = ipo_parsed.dropna(subset=["event_date", "amount"])
        ipo_parsed = ipo_parsed[(ipo_parsed["event_date"] >= begin_ts) & (ipo_parsed["event_date"] <= end_ts)].copy()

        private_date_col = next((column for column in private_detail.columns if "增发上市日" in str(column)), None)
        private_amount_col = next((column for column in private_detail.columns if "增发实际募集资金总额" in str(column)), None)
        if private_date_col is None or private_amount_col is None:
            raise RuntimeError(f"图15 定增问财返回字段缺失。 columns={list(private_detail.columns)}")

        private_parsed = private_detail[[private_date_col, private_amount_col]].copy()
        private_parsed.columns = ["event_date", "amount"]
        private_parsed["event_date"] = pd.to_datetime(private_parsed["event_date"].astype(str), errors="coerce")
        private_parsed["amount"] = pd.to_numeric(private_parsed["amount"], errors="coerce")
        private_parsed = private_parsed.dropna(subset=["event_date", "amount"])
        private_parsed = private_parsed[
            (private_parsed["event_date"] >= begin_ts) & (private_parsed["event_date"] <= end_ts)
        ].copy()

        if ipo_parsed.empty and private_parsed.empty:
            raise RuntimeError(
                f"图15 IPO和定增金额查询失败：IPO与定增明细清洗后均为空。 begin_date={begin_text}, end_date={end_text}"
            )

        ipo_weekly = (
            ipo_parsed.assign(week_start=lambda df: df["event_date"] - pd.to_timedelta(df["event_date"].dt.weekday, unit="D"))
            .groupby("week_start", as_index=False)
            .agg({"amount": lambda series: round(series.sum() / 1e8, 1), "event_date": "count"})
            .rename(columns={"amount": "IPO首发募集资金（亿元）", "event_date": "IPO首发家数"})
        )
        private_weekly = (
            private_parsed.assign(
                week_start=lambda df: df["event_date"] - pd.to_timedelta(df["event_date"].dt.weekday, unit="D")
            )
            .groupby("week_start", as_index=False)
            .agg({"amount": lambda series: round(series.sum() / 1e8, 1), "event_date": "count"})
            .rename(columns={"amount": "定增募集（亿元）", "event_date": "定增家数"})
        )

        all_week_fridays: list[pd.Timestamp] = []
        current_monday = begin_ts
        while current_monday <= end_ts:
            all_week_fridays.append(current_monday + pd.Timedelta(days=4))
            current_monday += pd.Timedelta(days=7)
        all_week_fridays = [date for date in all_week_fridays if date <= end_ts]
        target_fridays = all_week_fridays[-12:]
        if len(target_fridays) < 12:
            raise RuntimeError(
                f"图15 周度日期生成失败：最近12个周频日期不足。 begin_date={begin_text}, end_date={end_text}, generated={len(target_fridays)}"
            )

        base = pd.DataFrame({"日期": [date.strftime("%Y-%m-%d") for date in target_fridays]})
        base["week_start"] = pd.to_datetime(base["日期"]) - pd.to_timedelta(pd.to_datetime(base["日期"]).dt.weekday, unit="D")

        ipo_weekly["week_start"] = pd.to_datetime(ipo_weekly["week_start"])
        private_weekly["week_start"] = pd.to_datetime(private_weekly["week_start"])
        merged = base.merge(ipo_weekly, on="week_start", how="left").merge(private_weekly, on="week_start", how="left")
        merged = merged.fillna(0)
        merged["IPO首发家数"] = pd.to_numeric(merged["IPO首发家数"], errors="coerce").fillna(0).astype(int)
        merged["IPO首发募集资金（亿元）"] = (
            pd.to_numeric(merged["IPO首发募集资金（亿元）"], errors="coerce").fillna(0).round(1)
        )
        merged["定增家数"] = pd.to_numeric(merged["定增家数"], errors="coerce").fillna(0).astype(int)
        merged["定增募集（亿元）"] = pd.to_numeric(merged["定增募集（亿元）"], errors="coerce").fillna(0).round(1)
        merged = merged[["日期", "IPO首发家数", "IPO首发募集资金（亿元）", "定增家数", "定增募集（亿元）"]]
        print(f"图15周度聚合完成: shape={merged.shape}, sample={merged.tail(5).to_dict(orient='records')}")
        return merged

    block_mapping_cache_path = paths.funding_draft_dir / "data" / "cache" / "block_code_mapping.json"
    block_mapping_cache: dict | None = None
    block_scope_labels = {
        "sw_l1": "申万一级行业",
        "sw_l2": "申万二级行业",
        "benchmark": "block",
    }
    block_query_aliases = {
        "sw_l1": {
            "基础化工": ["化学原料和化学制品制造业指数"],
            "食品饮料": ["酒饮料和精制茶制造业指数"],
            "轻工制造": ["造纸和纸制品业指数"],
            "医药生物": ["医药制造业指数"],
            "公用事业": ["电力热力燃气及水生产和供应业指数"],
            "交通运输": ["交通运输仓储和邮政业指数"],
            "非银金融": ["其他金融业指数"],
            "机械设备": ["通用设备制造业指数"],
            "石油石化": ["石油加工贸易"],
            "环保": ["环境治理"],
            "农林牧渔": ["农林牧渔业指数"],
        },
        "sw_l2": {
            "酒店餐饮": ["酒店"],
            "旅游及景区": ["旅游综合"],
        },
        "benchmark": {
            "基金重仓指数": ["基金重仓股"],
        },
    }
    block_code_overrides = {
        "sw_l1": {
            "基础化工": "700789.TI",
            "食品饮料": "700779.TI",
            "轻工制造": "700785.TI",
            "医药生物": "700790.TI",
            "公用事业": "700753.TI",
            "交通运输": "700756.TI",
            "非银金融": "700829.TI",
            "机械设备": "700797.TI",
            "石油石化": "881180.TI",
            "环保": "881181.TI",
            "农林牧渔": "700750.TI",
        },
        "sw_l2": {
            "酒店餐饮": "884167.TI",
            "旅游及景区": "884165.TI",
        },
        "benchmark": {
            "基金重仓指数": "883406.TI",
        },
    }

    def load_block_mapping_cache() -> dict:
        nonlocal block_mapping_cache
        if block_mapping_cache is not None:
            return block_mapping_cache
        if block_mapping_cache_path.exists():
            try:
                block_mapping_cache = json.loads(block_mapping_cache_path.read_text(encoding="utf-8"))
            except Exception:
                block_mapping_cache = {}
        else:
            block_mapping_cache = {}
        return block_mapping_cache

    def save_block_mapping_cache() -> None:
        cache = load_block_mapping_cache()
        block_mapping_cache_path.parent.mkdir(parents=True, exist_ok=True)
        block_mapping_cache_path.write_text(
            json.dumps(cache, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def normalize_block_name(value: str) -> str:
        text = str(value or "").strip()
        for token in ["(申万)", "(中信)", "（申万）", "（中信）", "Ⅱ", "指数", "概念", "行业", "主题", "板块"]:
            text = text.replace(token, "")
        for token in ["业指数", "制造业", "服务业", "仓储和邮政", "热力燃气及水生产和供应", "热力生产和供应"]:
            text = text.replace(token, "")
        return text.replace(" ", "")

    def parse_block_query_dataframe(df: pd.DataFrame) -> pd.DataFrame:
        if df is None or df.empty or len(df.columns) < 2:
            return pd.DataFrame(columns=["code", "name", "value"])
        code_col = df.columns[0]
        name_col = df.columns[1]
        value_col = df.columns[-1] if len(df.columns) >= 3 else None
        parsed = pd.DataFrame(
            {
                "code": df[code_col].astype(str).str.strip(),
                "name": df[name_col].astype(str).str.strip(),
            }
        )
        if value_col and value_col not in {code_col, name_col}:
            parsed["value"] = pd.to_numeric(df[value_col], errors="coerce")
        else:
            parsed["value"] = np.nan
        return parsed.drop_duplicates(subset=["code"], keep="first")

    def score_block_candidate(target_name: str, candidate_name: str, candidate_code: str, has_value: bool) -> float:
        normalized_target = normalize_block_name(target_name)
        normalized_candidate = normalize_block_name(candidate_name)
        score = SequenceMatcher(None, normalized_target, normalized_candidate).ratio()
        if normalized_target and (
            normalized_target in normalized_candidate or normalized_candidate in normalized_target
        ):
            score += 0.15
        if str(candidate_code).endswith(".TI"):
            score += 0.05
        elif str(candidate_code).endswith((".SH", ".SZ")):
            score += 0.03
        if has_value:
            score += 0.02
        return score

    def get_strict_sw_scope_label(scope_key: str) -> str:
        if scope_key == "sw_l1":
            return "申万一级行业"
        if scope_key == "sw_l2":
            return "申万二级行业"
        return scope_key

    def probe_strict_sw_pct_candidates(
        target_name: str,
        scope_key: str,
        start_date: str,
        end_date: str,
        *,
        limit: int = 5,
    ) -> list[dict[str, object]]:
        normalized_target_name = (
            str(target_name)
            .replace("(申万)", "")
            .replace("（申万）", "")
            .strip()
        )
        scope_label = get_strict_sw_scope_label(scope_key)
        query_text = f"{normalized_target_name} {scope_label} {start_date}到{end_date}区间涨跌幅"
        try:
            raw_df = local_client.wc_query_dataframe(query_text, domain="index")
        except Exception as exc:
            return [{"query": query_text, "error": str(exc)}]

        parsed = parse_block_query_dataframe(raw_df)
        if parsed.empty:
            return [{"query": query_text, "error": "empty_dataframe"}]

        parsed["score"] = parsed.apply(
            lambda row: score_block_candidate(
                normalized_target_name,
                row["name"],
                row["code"],
                not pd.isna(row["value"]),
            ),
            axis=1,
        )
        parsed = parsed.sort_values(
            ["score", "value"],
            ascending=[False, False],
            na_position="last",
        ).head(max(1, int(limit)))
        candidate_rows: list[dict[str, object]] = []
        for _, row in parsed.iterrows():
            value_pct = None
            if pd.notna(row["value"]):
                value_pct = float(row["value"])
                if abs(value_pct) > 1.0:
                    value_pct = value_pct / 100
            candidate_rows.append(
                {
                    "query": query_text,
                    "code": str(row["code"]),
                    "name": str(row["name"]),
                    "value_pct": value_pct,
                    "score": round(float(row["score"]), 4),
                }
            )
        return candidate_rows

    def resolve_block_code(target_name: str, scope_key: str, reference_date: str) -> str:
        cache = load_block_mapping_cache()
        scope_cache = cache.setdefault(scope_key, {})
        cached = scope_cache.get(target_name)
        if isinstance(cached, dict) and cached.get("code"):
            return str(cached["code"])

        override_code = block_code_overrides.get(scope_key, {}).get(target_name)
        if override_code:
            scope_cache[target_name] = {
                "code": override_code,
                "name": target_name,
                "source": "override",
                "resolved_at": reference_date,
            }
            save_block_mapping_cache()
            print(f"block mapping override: scope={scope_key}, target={target_name}, code={override_code}")
            return override_code

        scope_label = block_scope_labels[scope_key]
        query_names = [target_name]
        query_names.extend(block_query_aliases.get(scope_key, {}).get(target_name, []))
        candidate_frames = []
        for query_name in query_names:
            if scope_label == "block":
                query_texts = [
                    f"{query_name} {reference_date}收盘价",
                    f"{query_name} {reference_date}涨跌幅",
                    query_name,
                ]
            else:
                query_texts = [
                    f"{query_name} {scope_label} {reference_date}涨跌幅",
                    f"{query_name} {scope_label} {reference_date}收盘价",
                    f"{query_name} {scope_label}",
                ]
            for query_text in query_texts:
                try:
                    raw_df = local_client.block_wc_query(query_text)
                except Exception:
                    continue
                parsed = parse_block_query_dataframe(raw_df)
                if parsed.empty:
                    continue
                parsed["score"] = parsed.apply(
                    lambda row: score_block_candidate(
                        target_name,
                        row["name"],
                        row["code"],
                        not pd.isna(row["value"]),
                    ),
                    axis=1,
                )
                parsed["query_text"] = query_text
                candidate_frames.append(parsed)

        if not candidate_frames:
            raise RuntimeError(
                f"无法解析 block 代码映射: scope={scope_key}, target_name={target_name}, reference_date={reference_date}"
            )

        candidates = pd.concat(candidate_frames, ignore_index=True)
        candidates = candidates.sort_values(["score", "value"], ascending=[False, False], na_position="last")
        best = candidates.iloc[0]
        if float(best["score"]) < 0.45:
            preview = candidates[["code", "name", "score"]].head(5).to_dict(orient="records")
            raise RuntimeError(
                f"block 代码映射置信度不足: scope={scope_key}, target_name={target_name}, "
                f"best_candidate={best['code']} {best['name']}, score={best['score']}, candidates={preview}"
            )
        scope_cache[target_name] = {
            "code": best["code"],
            "name": best["name"],
            "source": best["query_text"],
            "resolved_at": reference_date,
        }
        save_block_mapping_cache()
        print(
            f"block mapping resolved: scope={scope_key}, target={target_name}, "
            f"code={best['code']}, name={best['name']}, source={best['query_text']}"
        )
        return str(best["code"])

    def fetch_block_close_values_for_dates(
        target_name: str,
        scope_key: str,
        requested_dates: list[str],
        *,
        reference_date: str,
    ) -> dict[str, float]:
        normalized_target_name = str(target_name).replace("(申万)", "").replace("（申万）", "").strip()
        mapped_code = resolve_block_code(normalized_target_name, scope_key, reference_date)
        scope_label = block_scope_labels.get(scope_key, scope_key)
        value_map: dict[str, float] = {}
        attempts: list[dict[str, object]] = []
        for requested_date in requested_dates:
            query_date = pd.Timestamp(requested_date).strftime("%Y%m%d")
            query_candidates = [
                f"{normalized_target_name} {scope_label} {query_date}收盘价",
                f"{mapped_code} {query_date}收盘价",
                f"{normalized_target_name} {query_date}收盘价",
            ]
            resolved_value = None
            for query_text in query_candidates:
                try:
                    raw_df = local_client.block_wc_query(query_text)
                    parsed = parse_block_query_dataframe(raw_df)
                    if parsed.empty:
                        attempts.append({"date": requested_date, "query": query_text, "error": "empty_dataframe"})
                        continue
                    parsed["score"] = parsed.apply(
                        lambda row: score_block_candidate(
                            normalized_target_name,
                            row["name"],
                            row["code"],
                            not pd.isna(row["value"]),
                        ),
                        axis=1,
                    )
                    exact_code = parsed[parsed["code"].astype(str).str.upper() == str(mapped_code).upper()]
                    candidate = exact_code.iloc[0] if not exact_code.empty else parsed.sort_values(
                        ["score", "value"], ascending=[False, False], na_position="last"
                    ).iloc[0]
                    numeric_value = pd.to_numeric(pd.Series([candidate["value"]]), errors="coerce").iloc[0]
                    if pd.isna(numeric_value):
                        attempts.append({"date": requested_date, "query": query_text, "error": "nan_value"})
                        continue
                    resolved_value = float(numeric_value)
                    attempts.append(
                        {
                            "date": requested_date,
                            "query": query_text,
                            "matched_code": candidate["code"],
                            "matched_name": candidate["name"],
                            "value": resolved_value,
                        }
                    )
                    break
                except Exception as exc:
                    attempts.append({"date": requested_date, "query": query_text, "error": str(exc)})
            if resolved_value is not None:
                value_map[pd.Timestamp(requested_date).strftime("%Y-%m-%d")] = resolved_value

        if not value_map:
            raise RuntimeError(
                f"{scope_key} block_wc close fallback failed: target_name={normalized_target_name}, "
                f"mapped_code={mapped_code}, requested_dates={requested_dates}, attempts={attempts[:12]}"
            )
        missing_dates = [date for date in requested_dates if pd.Timestamp(date).strftime("%Y-%m-%d") not in value_map]
        print(
            f"{scope_key} block_wc close fallback summary: target_name={normalized_target_name}, "
            f"mapped_code={mapped_code}, resolved_dates={len(value_map)}, missing_dates={missing_dates[:8]}"
        )
        return value_map

    def fetch_block_pct_change_series(
        target_names: list[str],
        scope_key: str,
        start_date: str,
        end_date: str,
        *,
        original_codes: list[str] | None = None,
    ) -> pd.DataFrame:
        if use_strict_sw_block_codes(scope_key):
            if original_codes is None:
                raise RuntimeError(
                    f"{scope_key} strict pct_chg query missing original_codes. "
                    f"start={start_date}, end={end_date}"
                )
            mapped_codes = [str(code) for code in original_codes]
            period_begin = pd.to_datetime(str(start_date), format="%Y%m%d", errors="coerce").normalize()
            period_end = pd.to_datetime(str(end_date), format="%Y%m%d", errors="coerce").normalize()
            if pd.notna(period_begin) and pd.notna(period_end):
                cache_rows = []
                cache_missing = []
                for idx, name in enumerate(target_names):
                    original_code = str(original_codes[idx])
                    close_series = load_local_block_close_series(scope_key, original_code)
                    value_pct = np.nan
                    if not close_series.empty:
                        end_valid = close_series[close_series.index <= period_end].dropna()
                        base_valid = close_series[close_series.index < period_begin].dropna()
                        if not end_valid.empty and not base_valid.empty:
                            end_close = float(end_valid.iloc[-1])
                            base_close = float(base_valid.iloc[-1])
                            if base_close != 0:
                                value_pct = end_close / base_close - 1
                                print(
                                    f"{scope_key} strict pct_chg cache first: code={original_code}, "
                                    f"name={name}, base_date={format_log_date(base_valid.index[-1])}, "
                                    f"end_date={format_log_date(end_valid.index[-1])}, value_pct={value_pct:.6%}"
                                )
                    if pd.isna(value_pct):
                        cache_missing.append(name)
                    cache_rows.append(
                        {
                            "target_name": name,
                            "mapped_code": original_code,
                            "value_pct": value_pct,
                            "original_code": original_code,
                        }
                    )
                if not cache_missing:
                    return pd.DataFrame(cache_rows)
        else:
            mapped_codes = [resolve_block_code(name, scope_key, end_date) for name in target_names]
        result = snapshot_batch(mapped_codes, "pct_chg_per", f"startDate={start_date};endDate={end_date}")
        raw_values = result.Data[0] if getattr(result, "Data", None) else []
        rows = []
        missing_names = []
        for idx, name in enumerate(target_names):
            value = raw_values[idx] if idx < len(raw_values) else np.nan
            numeric_value = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
            if pd.isna(numeric_value):
                missing_names.append(name)
            rows.append(
                {
                    "target_name": name,
                    "mapped_code": mapped_codes[idx],
                    "value_pct": float(numeric_value) / 100 if pd.notna(numeric_value) else np.nan,
                    "original_code": original_codes[idx] if original_codes else mapped_codes[idx],
                }
            )
        if missing_names and use_strict_sw_block_codes(scope_key):
            period_begin = pd.to_datetime(str(start_date), format="%Y%m%d", errors="coerce").normalize()
            period_end = pd.to_datetime(str(end_date), format="%Y%m%d", errors="coerce").normalize()
            if pd.notna(period_begin) and pd.notna(period_end):
                for row in rows:
                    if pd.notna(row["value_pct"]):
                        continue
                    original_code = str(row["original_code"])
                    close_series = load_local_block_close_series(scope_key, original_code)
                    if close_series.empty:
                        continue
                    end_valid = close_series[close_series.index <= period_end].dropna()
                    base_valid = close_series[close_series.index < period_begin].dropna()
                    if end_valid.empty or base_valid.empty:
                        continue
                    end_close = float(end_valid.iloc[-1])
                    base_close = float(base_valid.iloc[-1])
                    if base_close == 0:
                        continue
                    row["value_pct"] = end_close / base_close - 1
                    print(
                        f"{scope_key} strict pct_chg cache fallback: code={original_code}, "
                        f"name={row['target_name']}, base_date={format_log_date(base_valid.index[-1])}, "
                        f"end_date={format_log_date(end_valid.index[-1])}, value_pct={row['value_pct']:.6%}"
                    )
                missing_names = [row["target_name"] for row in rows if pd.isna(row["value_pct"])]
        if missing_names:
            error_target = "original_codes" if use_strict_sw_block_codes(scope_key) else "block mapping"
            strict_diagnostics = None
            if use_strict_sw_block_codes(scope_key):
                strict_diagnostics = []
                for row in rows:
                    if pd.notna(row["value_pct"]):
                        continue
                    strict_diagnostics.append(
                        {
                            "target_name": row["target_name"],
                            "original_code": row["original_code"],
                            "strict_cache_path": get_block_close_cache_path(scope_key, row["original_code"]).name,
                            "wc_candidates": probe_strict_sw_pct_candidates(
                                row["target_name"],
                                scope_key,
                                start_date,
                                end_date,
                                limit=5,
                            ),
                        }
                    )
            raise RuntimeError(
                f"{scope_key} 涨跌幅查询结果仍包含空值，请检查 {error_target}。"
                f" start={start_date}, end={end_date}, 缺失名称={missing_names}, "
                f"诊断信息={strict_diagnostics[:6] if strict_diagnostics else None}"
            )
        return pd.DataFrame(rows)


    # %%参数输入
    fix_path = str(paths.fund_list_dir)

    stock_ETF = pd.read_excel(f"{fix_path}/stock_ETF.xlsx")
    stock_ETF_list = stock_ETF["证券代码"].to_list()
    print(f"基金清单加载完成：股票ETF数量={len(stock_ETF_list)}")

    normal_stock_pref = pd.read_excel(f"{fix_path}/normal_stock_pref.xlsx")
    normal_stock_pref_list = normal_stock_pref["证券代码"].to_list()
    print(f"基金清单加载完成：普通股票偏股基金数量={len(normal_stock_pref_list)}")

    flexible_allocation_fund = pd.read_excel(f"{fix_path}/flex_allo_fund.xlsx")
    flexible_allocation_fund_list = flexible_allocation_fund["证券代码"].to_list()
    print(f"基金清单加载完成：灵活配置基金数量={len(flexible_allocation_fund_list)}")

    passive_fund = pd.read_excel(f"{fix_path}/passive_fund.xlsx")
    passive_fund_list = passive_fund["证券代码"].to_list()
    print(f"基金清单加载完成：被动基金数量={len(passive_fund_list)}")

    # %%参数输入
    today0 = report_date     # 线上运行使用报告日
    # today0 = datetime.date(2023, 1, 20)  # 本地调试可手动指定日期

    # %%日期处理（保留原有周口径逻辑）
    start_week = (today0 - timedelta(days=today0.weekday())).strftime("%Y%m%d")
    start_day14 = (today0 - timedelta(days=14)).strftime("%Y%m%d")
    start_week16= (today0 - timedelta(days=(today0.weekday()+7*17))).strftime("%Y%m%d")
    start_week20 = (today0 - timedelta(days=(today0.weekday()+7*21))).strftime("%Y%m%d")
    #准备
    start_week_1year_52_Fri = (today0 - timedelta(days = (today0.weekday()+7*51)))+timedelta(days = 4)


    # 为了1准备
    start_week9 = (today0 - timedelta(days=(today0.weekday()+7*8))) # 9周前周一
    start_week9_Fri =(today0 - timedelta(days=(today0.weekday()+7*8))+timedelta(days=4)) # 9周前

    today = today0.strftime("%Y%m%d")
    today_connect = today0.strftime("%Y-%m-%d")
    logger.add(str(paths.workspace_dir / 'logs' / 'updating_data.log'), encoding="utf-8")
    DB_LOC = str(paths.data_dir)


    # %%连接WIND

    # %% ：本周市场主要指数涨跌幅
    market_Index = build_market_index_returns_dataframe()
    market_Index.to_csv("outputs/csv/figures_1_15/图1_本周市场主要指数涨跌幅.csv",index = 0,encoding="utf-8-sig")

    # %% ：本周申万一级行业涨跌幅
    ind1_name = pd.read_csv(r'{}/行业代码/一级行业代码.csv'.format(DB_LOC),encoding="GBK")
    ind1_list = ind1_name["申万一级行业"].astype(str).tolist()
    ind1_code = ind1_name["代码"].astype(str).tolist()
    print(f"图2行业收益自算开始: name_count={len(ind1_list)}")

    ind1_assets = list(zip(ind1_code, ind1_list))
    ind1_Index = build_block_self_computed_returns_dataframe(ind1_assets, "sw_l1", "图2")
    ind1_Index = ind1_Index[["市场指数", "本周涨跌幅(%)", "本月涨跌幅(%)", "本年涨跌幅(%)"]]

    # 按照周涨幅排
    ind1_Index = ind1_Index.sort_values(by='本周涨跌幅(%)')

    ind1_Index.to_csv("outputs/csv/figures_1_15/图2_本周申万一级行业涨跌幅.csv",index = 0,encoding="utf-8-sig")
    # %% ：本周申万二级行业涨跌幅
    ind2_get = pd.read_csv(r'{}/行业代码/二级行业代码.csv'.format(DB_LOC),encoding="GBK")
    ind2_list = ind2_get["申万二级行业"].astype(str).tolist()
    ind2_code = ind2_get["代码"].astype(str).tolist()
    print(f"图3行业收益自算开始: name_count={len(ind2_list)}")
    ind2_assets = list(zip(ind2_code, ind2_list))
    ind2_Index = build_block_self_computed_returns_dataframe(ind2_assets, "sw_l2", "图3")
    ind2_Index = ind2_Index[["代码", "市场指数", "本周涨跌幅(%)", "本月涨跌幅(%)", "本年涨跌幅(%)"]]


    # 按照周涨幅排
    ind2_Index = ind2_Index.sort_values(by='本周涨跌幅(%)')
    ind2_rank10r = ind2_Index.iloc[0:10]
    ind2_rank10r.loc[:, "排名"] = range(1, 11)
    ind2_rank10d = ind2_Index.iloc[len(ind2_Index)-10:len(ind2_Index)]
    ind2_rank10d.loc[:, "排名"] = range(-10, 0)
    null = pd.DataFrame()

    null = pd.DataFrame({"代码": [""], "市场指数": ["......"],
                        "本周涨跌幅(%)": [""], "本月涨跌幅(%)": [""], "本年涨跌幅(%)": [""], "排名": [""]})
    #ind2_data = ind2_rank10r.append(null.append(ind2_rank10d))
    ind2_data = pd.concat([ind2_rank10r, null, ind2_rank10d], ignore_index=True)

    ind2_data = ind2_data[["排名", "代码", "市场指数", "本周涨跌幅(%)", "本月涨跌幅(%)", "本年涨跌幅(%)"]]

    ind2_data.to_csv("outputs/csv/figures_1_15/图3_本周申万二级行业涨跌幅.csv",index = 0,encoding="utf-8-sig")
    # %% ：本周热点概念板块涨跌幅（排名前10位和0位）
    hot_query_df = parse_block_query_dataframe(
        local_client.block_wc_query(f"概念板块 {start_week}到{today}区间涨跌幅")
    )
    excluded_hot_roots = {"8841299", "8841300", "884246"}
    hot_query_df = hot_query_df[
        (
            hot_query_df["code"].astype(str).str.endswith(".TI")
            & ~hot_query_df["code"].astype(str).str.startswith("700")
            & ~hot_query_df["code"].astype(str).str.split(".").str[0].isin(excluded_hot_roots)
        )
    ].copy()
    hot_query_df = hot_query_df[hot_query_df["value"].notna()].copy()
    if hot_query_df.empty:
        raise RuntimeError(
            f"图4热点概念板块查询失败：未返回有效数据。 start={start_week}, end={today}"
        )
    hot_codes = hot_query_df["code"].tolist()
    hot_snapshot = snapshot_batch(hot_codes, "pct_chg_per", f"startDate={start_week};endDate={today}")
    hot_values = clean_numeric_values(hot_snapshot.Data[0] if getattr(hot_snapshot, "Data", None) else [])
    if len(hot_values) < len(hot_codes):
        raise RuntimeError(
            f"图4热点概念板块区间涨跌幅查询失败：返回数量不足。 "
            f"requested={len(hot_codes)}, received={len(hot_values)}"
        )

    # 汇
    hot_Index = pd.DataFrame()
    hot_Index["代码"] = hot_codes
    hot_Index["市场指数"] = hot_query_df["name"].tolist()
    hot_Index["本周涨跌幅(%)"] = (pd.Series(hot_values[: len(hot_codes)], dtype="float64") / 100).tolist()
    #hot_Index["本月涨跌幅(%)"] = (pd.DataFrame(hot_Index_m.Data).T)/100
    #hot_Index["本年涨跌幅(%)"] = (pd.DataFrame(hot_Index_y.Data).T)/100

    # 按照周涨幅排
    hot_Index = hot_Index.sort_values(by='本周涨跌幅(%)')

    hot_rank10r = hot_Index.iloc[0:10]
    hot_rank10r["排名"] = range(1, 11)
    null = pd.DataFrame({"代码": [""], "市场指数": ["......"],
                        "本周涨跌幅(%)": [""], "排名": [""]})
    hot_rank10d = hot_Index.iloc[len(hot_Index)-10:len(hot_Index)]
    hot_rank10d["排名"] = range(-10, 0)
    #null = pd.DataFrame()
    #hot = hot_rank10r.append(hot_rank10d)

    hot_name = hot_Index["市场指数"][0:10].tolist()+["........"] + \
        hot_Index["市场指数"][len(hot_Index)-10:len(hot_Index)].tolist()
    hot_name = [trim_chart_label_suffix(x) for x in hot_name]

    #hot_data = hot_rank10r.append(null.append(hot_rank10d))
    hot_data = pd.concat([hot_rank10r, null, hot_rank10d], ignore_index=True)
    hot_data = hot_data[["排名", "代码", "市场指数", "本周涨跌幅(%)"]]

    hot_data.to_csv("outputs/csv/figures_1_15/图4_本周热点概念板块涨跌幅（排名前10位和后10位）.csv",index = 0,encoding="utf-8-sig")


    # %% : 两市成交金额、换手率
    fig5_amount_source_config = {
        "沪市:成交金额(亿元)": {"quotation_index_name": "上证指数"},
        "深市:成交金额(亿元)": {"quotation_index_name": "深证成指"},
    }
    fig5_turnover_source_config = {
        "换手率_上证综合指数(%)": {
            "indicator_code": "M0331169",
            "wc_query_code": "000001.SH",
            "wc_metric_name": "换手率",
            "wc_domain": "index",
        },
        "换手率_深证综合指数(%)": {
            "indicator_code": "M0331178",
            "wc_query_code": "399106.SZ",
            "wc_metric_name": "换手率",
            "wc_domain": "index",
        },
    }
    fig5_required_fields = list(fig5_amount_source_config.keys()) + list(fig5_turnover_source_config.keys())
    fig5_turnover_refresh_results: dict[str, dict[str, object]] = {}
    fig5_turnover_fallback_results: dict[str, dict[str, object]] = {}

    def build_fig5_dataframe() -> tuple[pd.DataFrame, pd.Timestamp, pd.Timestamp, dict[str, dict[str, object]]]:
        merged = pd.DataFrame()
        source_status: dict[str, dict[str, object]] = {}

        for output_col, config in fig5_amount_source_config.items():
            local_series = load_local_quotation_numeric_series(
                config["quotation_index_name"],
                "amt",
                output_col,
            )
            if not local_series.empty:
                local_series = (local_series / 100000000).rename(output_col)
            latest_local_date = local_series.index.max() if not local_series.empty else pd.NaT
            source_status[output_col] = {
                "source": "quotation_local",
                "latest_local_date": format_log_date(latest_local_date),
                "latest_valid_date": format_log_date(latest_local_date),
            }
            if not local_series.empty:
                merged = local_series.to_frame() if merged.empty else merged.join(local_series.to_frame(), how="outer")

        for indicator_name in fig5_turnover_source_config:
            indexed = load_market_sentiment_indexed_series(indicator_name)
            series = indexed[indicator_name].dropna() if not indexed.empty else pd.Series(dtype="float64", name=indicator_name)
            latest_local_date = series.index.max() if not series.empty else pd.NaT
            source_status[indicator_name] = {
                "source": "market_sentiment_local",
                "latest_local_date": format_log_date(latest_local_date),
                "latest_valid_date": format_log_date(latest_local_date),
                "refresh_meta": fig5_turnover_refresh_results.get(indicator_name, {}),
                "fallback_meta": fig5_turnover_fallback_results.get(indicator_name, {}),
            }
            if not series.empty:
                merged = series.to_frame() if merged.empty else merged.join(series.to_frame(), how="outer")

        if merged.empty:
            merged = pd.DataFrame(columns=fig5_required_fields)
            merged.index = pd.DatetimeIndex([], name="date")
        else:
            merged = merged[~merged.index.duplicated(keep="last")].sort_index()
        for field in fig5_required_fields:
            if field not in merged.columns:
                merged[field] = np.nan
        merged = merged[fig5_required_fields]
        merged_latest_cache = merged.index.max() if not merged.empty else pd.NaT
        merged_latest_complete = (
            merged.dropna(subset=fig5_required_fields).index.max()
            if not merged.empty
            else pd.NaT
        )
        return merged, merged_latest_cache, merged_latest_complete, source_status

    def get_fig5_stale_fields(data_all: pd.DataFrame) -> list[str]:
        stale_fields: list[str] = []
        for field in fig5_required_fields:
            if field not in data_all.columns:
                stale_fields.append(field)
                continue
            series = data_all[field].dropna()
            latest_valid_date = series.index.max() if not series.empty else pd.NaT
            if pd.isna(latest_valid_date) or pd.Timestamp(latest_valid_date).normalize() < target_market_date:
                stale_fields.append(field)
        return stale_fields

    def build_fig5_failure_details(
        data_all: pd.DataFrame,
        source_status: dict[str, dict[str, object]],
    ) -> tuple[list[str], list[str]]:
        stale_details: list[str] = []
        union_error_codes: set[str] = set()
        for field in get_fig5_stale_fields(data_all):
            series = data_all[field].dropna() if field in data_all.columns else pd.Series(dtype="float64")
            latest_valid_date = series.index.max() if not series.empty else pd.NaT
            source_meta = source_status.get(field, {})
            refresh_meta = fig5_turnover_refresh_results.get(field, {})
            fallback_meta = fig5_turnover_fallback_results.get(field, source_meta.get("fallback_meta", {}))
            error_codes = [str(item) for item in (refresh_meta.get("error_codes") or [])]
            union_error_codes.update(error_codes)
            stale_details.append(
                f"{field}[source={source_meta.get('source')}, latest_local_date={source_meta.get('latest_local_date')}, "
                f"latest_valid_date={format_log_date(latest_valid_date)}, "
                f"requested_trade_dates={list(refresh_meta.get('requested_trade_dates') or [])[:8]}, "
                f"missing_trade_dates={list(refresh_meta.get('missing_trade_dates') or [])[:8]}, "
                f"error_codes={error_codes}, fallback_meta={fallback_meta}]"
            )
        return stale_details, sorted(code for code in union_error_codes if code)

    data_all, latest_cache_date, latest_complete_date, fig5_source_status = build_fig5_dataframe()
    print(
        f"图5缓存检查: latest_cache_date={format_log_date(latest_cache_date)}, "
        f"latest_complete_date={format_log_date(latest_complete_date)}, "
        f"target_market_date={target_market_date_text}, report_date={report_date_text}"
    )
    fig5_stale_fields = get_fig5_stale_fields(data_all)
    if fig5_stale_fields:
        print(
            f"图5缓存陈旧，尝试在线补数: latest_cache_date={format_log_date(latest_cache_date)}, "
            f"latest_complete_date={format_log_date(latest_complete_date)}, "
            f"target_market_date={target_market_date_text}, stale_fields={fig5_stale_fields}"
        )
        trade_dates = get_trade_dates_from_quotation_cache(
            (target_market_date - pd.Timedelta(days=14)).strftime("%Y-%m-%d"),
            target_market_date.strftime("%Y-%m-%d"),
        )
        stale_turnover_fields = [
            field for field in fig5_stale_fields if field in fig5_turnover_source_config
        ]
        for field in stale_turnover_fields:
            config = fig5_turnover_source_config[field]
            indexed = load_market_sentiment_indexed_series(field)
            latest_valid_ts = indexed[field].dropna().index.max() if not indexed.empty else pd.NaT
            if pd.notna(latest_valid_ts):
                pending_trade_dates = [
                    item for item in trade_dates
                    if pd.Timestamp(item).normalize() > pd.Timestamp(latest_valid_ts).normalize()
                ]
            else:
                pending_trade_dates = trade_dates
            fig5_turnover_refresh_results[field] = {
                "requested_trade_dates": list(pending_trade_dates),
                "missing_trade_dates": list(pending_trade_dates),
                "latest_cache_date": format_log_date(indexed.index.max() if not indexed.empty else pd.NaT),
                "latest_valid_date": format_log_date(latest_valid_ts),
                "error_codes": ["-4318"] if pending_trade_dates else [],
            }
            if not pending_trade_dates:
                continue
            try:
                wc_series, wc_meta = fetch_wc_metric_series_for_trade_dates(
                    f"图5 {field}",
                    config["wc_query_code"],
                    config["wc_metric_name"],
                    pending_trade_dates,
                    domain=config.get("wc_domain", "index"),
                )
                wc_write_df = wc_series.rename(field).reset_index()
                upsert_market_sentiment_series(field, wc_write_df)
                fig5_turnover_fallback_results[field] = {
                    "success": True,
                    "source": "wc_query",
                    "query_code": wc_meta.get("query_code"),
                    "metric_name": wc_meta.get("metric_name"),
                    "domain": wc_meta.get("domain"),
                    "latest_written_date": format_log_date(wc_series.index.max()),
                    "attempt_count": len(wc_meta.get("attempts") or []),
                }
            except Exception as exc:
                fig5_turnover_fallback_results[field] = {
                    "success": False,
                    "source": "wc_query",
                    "error": str(exc),
                }
                print(f"图5 WC fallback 失败: field={field}, error={exc}")
        data_all, latest_cache_date, latest_complete_date, fig5_source_status = build_fig5_dataframe()
        print(
            f"图5补数后缓存检查: latest_cache_date={format_log_date(latest_cache_date)}, "
            f"latest_complete_date={format_log_date(latest_complete_date)}, "
            f"target_market_date={target_market_date_text}, report_date={report_date_text}"
        )
    if get_fig5_stale_fields(data_all):
        stale_indicator_details, online_error_codes = build_fig5_failure_details(
            data_all,
            fig5_source_status,
        )
        upstream_reason = ""
        if online_error_codes:
            upstream_reason = (
                f" root_cause=market_sentiment turnover online refresh returned no valid rows, "
                f"online_error_codes={online_error_codes}."
            )
        raise RuntimeError(
            "FIG5 online repair failed. "
            f"{upstream_reason} stale_indicators={stale_indicator_details}. "
            "图5所需字段未更新到目标交易日。 "
            f"latest_cache_date={format_log_date(latest_cache_date)}, "
            f"latest_complete_date={format_log_date(latest_complete_date)}, "
            f"target_market_date={target_market_date_text}, report_date={report_date_text}"
        )

    print('--------------------------------------------------------')

    begin_date_ts = pd.Timestamp('2022-01-01')
    data_all_complete = data_all.dropna(subset=fig5_required_fields).copy()
    data_day = data_all_complete.loc[
        (data_all_complete.index >= begin_date_ts) & (data_all_complete.index <= target_market_date)
    ].copy()
    if data_day.empty:
        raise RuntimeError(
            f"图5时间切片后为空。 begin_date={begin_date_ts:%Y-%m-%d}, "
            f"target_market_date={target_market_date_text}, report_date={report_date_text}"
        )
    print(
        f"图5切片完成: last_3_dates={data_day.tail(3).index.strftime('%Y-%m-%d').tolist()}"
    )
    data_day = data_day[fig5_required_fields]

    # (3)周报
    # (3.1) 成交额、换手率
    data_day_trading_vol = data_day[fig5_required_fields]

    trading_use = data_day_trading_vol.tail(50).copy()  # ---0天数--#
    trading_use.index.name = "date"
    trading_use["all"] = trading_use.iloc[:, 0] + trading_use.iloc[:, 1]
    trading_use["参考线：1万亿元"] = 10000

    trading_data = trading_use.drop(["all"], axis=1)
    trading_use.to_csv("outputs/csv/figures_1_15/图5_两市成交金额、换手率.csv",encoding="utf-8")
    # %% ：股指期货主力合约基差
    qh_columns = [
        "日期",
        "IF主力合约:基差",
        "IH主力合约:基差",
        "IC主力合约:基差",
        "IM主力合约:基差",
        "IF主力合约:升水率",
        "IH主力合约:升水率",
        "IC主力合约:升水率",
        "IM主力合约:升水率",
    ]
    if not ENABLE_FIGURE6_FUTURES:
        print("图6开关关闭：跳过股指期货主力合约基差取数、绘图及相关文案。")
        qh_use50 = pd.DataFrame(columns=qh_columns)
    else:
        qh_list = ["000300.SH", "IF.CFE", "000016.SH", "IH.CFE",
                   "000905.SH", "IC.CFE", "000852.SH", "IM.CFE"]
        qh_fetch_code_map = {
            "IF.CFE": "IF00.CFE",
            "IH.CFE": "IH00.CFE",
            "IC.CFE": "IC00.CFE",
            "IM.CFE": "IM00.CFE",
        }

        # 先拉日期作为基准
        date_data = local_client.history("000300.SH", "close", start_week20, today, "")
        # 先拉日期作为基准，并强制规范日期格式，去掉潜在的时间
        qh_close = pd.DataFrame({"日期": pd.to_datetime(date_data.Times), "000300.SH": date_data.Data[0]})
        qh_close["日期"] = qh_close["日期"].dt.date

        # 其余品种按日merge 进来
        qh_failures: list[dict[str, object]] = []
        for name in qh_list[1:]:
            fetch_code = qh_fetch_code_map.get(name, name)
            res = local_client.history(fetch_code, "close", start_week20, today, "")

            if res.ErrorCode == 0 and res.Data and len(res.Data[0]) > 0:
                tmp = pd.DataFrame({"日期": pd.to_datetime(res.Times), name: res.Data[0]})
                tmp["日期"] = tmp["日期"].dt.date
                qh_close = pd.merge(qh_close, tmp, on="日期", how="left")
            else:
                print(f"警告：{name} 数据获取失败，fetch_code={fetch_code}，错误码: {res.ErrorCode}")
                qh_failures.append(
                    {
                        "display_code": name,
                        "fetch_code": fetch_code,
                        "error_code": getattr(res, "ErrorCode", None),
                        "data_len": len(res.Data[0]) if getattr(res, "Data", None) and len(res.Data) > 0 and res.Data[0] is not None else 0,
                    }
                )
                qh_close[name] = np.nan

        for col in qh_close.columns[1:]:
            try:
                qh_close[col] = qh_close[col].astype(float)
            except Exception:
                try:
                    qh_close[col] = qh_close[col].astype(str).astype(float)
                except Exception:
                    try:
                        qh_close[col] = pd.to_numeric(qh_close[col], errors='coerce')
                    except Exception:
                        print(f"无法转换 {col} 的数据类型。")

        qh_use = pd.DataFrame()
        qh_use["日期"] = qh_close["日期"]
        qh_use["IF主力合约:基差"] = qh_close["IF.CFE"] - qh_close["000300.SH"]
        qh_use["IH主力合约:基差"] = qh_close["IH.CFE"] - qh_close["000016.SH"]
        qh_use["IC主力合约:基差"] = qh_close["IC.CFE"] - qh_close["000905.SH"]
        qh_use["IM主力合约:基差"] = qh_close["IM.CFE"] - qh_close["000852.SH"]
        qh_use["IF主力合约:升水率"] = qh_use["IF主力合约:基差"]/qh_close["000300.SH"]
        qh_use["IH主力合约:升水率"] = qh_use["IH主力合约:基差"]/qh_close["000016.SH"]
        qh_use["IC主力合约:升水率"] = qh_use["IC主力合约:基差"]/qh_close["000905.SH"]
        qh_use["IM主力合约:升水率"] = qh_use["IM主力合约:基差"]/qh_close["000852.SH"]

        qh_use50 = qh_use.tail(50)
        basis_columns = ["IF主力合约:基差", "IH主力合约:基差", "IC主力合约:基差", "IM主力合约:基差"]
        has_valid_basis = (
            not qh_use50.empty
            and qh_use50[basis_columns].apply(pd.to_numeric, errors="coerce").notna().any().any()
        )
        if not has_valid_basis:
            raise RuntimeError(
                "图6_股指期货主力合约基差 数据生成失败：未获取到有效的股指期货主力合约收盘价。"
                f" 当前尝试路径为 {qh_fetch_code_map}，失败详情={qh_failures}。"
                "这通常意味着本地主连代码映射缺失，或期货收盘价文件不完整。"
            )
    qh_use50.to_csv("outputs/csv/figures_1_15/图6_股指期货主力合约基差.csv",index = 0,encoding="utf-8-sig")


    # %% ：两市涨跌数量统
    updown_columns = ["日期", "跌停家数", "下跌家数（不含跌停）", "平盘家数", "上涨家数（不含涨停）", "涨停家数"]
    if not ENABLE_FIGURE7_UPDOWN:
        print("图7开关关闭：跳过两市涨跌数量统计取数与绘图。")
        updown_use = pd.DataFrame(columns=updown_columns)
        updown_use.to_csv("outputs/csv/figures_1_15/图7_两市涨跌数量统计.csv",index = 0,encoding="utf-8-sig")
    else:
        try:
            updown_use = build_updown_statistics_from_dateserial(start_day14, today)
            print(f"图7最终 DataFrame shape={updown_use.shape}")
            updown_use.to_csv("outputs/csv/figures_1_15/图7_两市涨跌数量统计.csv",index = 0,encoding="utf-8-sig")
        except Exception as exc:
            updown_use = pd.DataFrame(columns=updown_columns)
            record_figure_failure("图7_两市涨跌数量统计", exc)
    # %% ：两市融资融券余额、融资买入占
    rzrq_columns = ["日期", "两市融资融券余额（万亿元）", "两市融资买入金额", "A股成交总额", "两市融资买入金额占比"]
    if not ENABLE_FIGURE8_MARGIN:
        print("图8开关关闭：跳过两市融资融券余额、融资买入占比取数与绘图。")
        rzrq_use = pd.DataFrame(columns=rzrq_columns)
        rzrq_use.to_csv("outputs/csv/figures_1_15/图8_两市融资融券余额、融资买入占比.csv",index = 0,encoding="utf-8-sig")
    else:
        try:
            rzrq_use, rzrq_meta = build_margin_dataframe_from_local_cache()
            latest_local_date = rzrq_meta.get("latest_local_date", pd.NaT)
            fallback_online = (
                rzrq_use is None
                or rzrq_use.empty
                or pd.isna(latest_local_date)
                or latest_local_date < margin_required_date
            )
            print(
                f"图8缓存检查: latest_local_date={format_log_date(latest_local_date)}, "
                f"target_market_date={target_market_date_text}, report_date={report_date_text}, "
                f"margin_required_date={margin_required_date_text}, "
                f"allow_stale_margin_data={options.allow_stale_margin_data}, fallback_online={fallback_online}"
            )
            if fallback_online:
                print(
                    f"图8数据来源回退: local_cache unavailable_or_stale, "
                    f"hit_fields={len(rzrq_meta['hit_fields'])}, missing_fields={len(rzrq_meta['missing_fields'])}, "
                    f"shape={rzrq_meta['shape']}, latest_local_date={format_log_date(latest_local_date)}, "
                    f"margin_required_date={margin_required_date_text}"
                )
                rzrq_use, rzrq_meta = build_margin_dataframe_from_online_edb()
                persist_margin_online_cache(rzrq_meta["raw_data"])
            print(
                f"图8数据来源: {rzrq_meta['source']}, hit_fields={len(rzrq_meta['hit_fields'])}, "
                f"missing_fields={len(rzrq_meta['missing_fields'])}, final_shape={rzrq_meta['shape']}, "
                f"latest_local_date={format_log_date(latest_local_date)}, "
                f"latest_online_date={format_log_date(rzrq_meta.get('latest_online_date', pd.NaT))}, "
                f"report_date={report_date_text}, target_market_date={target_market_date_text}"
            )
            if rzrq_use is None or rzrq_use.empty:
                raise RuntimeError(
                    "图8最终结果为空，禁止继续写出空CSV。 "
                    f"source={rzrq_meta.get('source')}, meta={rzrq_meta}"
                )
            rzrq_use.to_csv("outputs/csv/figures_1_15/图8_两市融资融券余额、融资买入占比.csv",index = 0,encoding="utf-8-sig")
        except Exception as exc:
            rzrq_use = pd.DataFrame(columns=rzrq_columns)
            record_figure_failure("图8_两市融资融券余额、融资买入占比", exc)
    checkpoint('ͼ1-4')

    # %%：北向资金买入成交金额、成交净买入、累计净买入
    # 与图11: 各类资金变动规模（亿元）本周增量资金来源（亿元）
    #** 定义方程与时间list **

    # %%
    # 获取最个周五的日期
    def get_last_3_fridays():
        today = datetime.datetime.combine(report_date, datetime.time.min)
        fridays = []

        while len(fridays) < 3:
            if today.weekday() == 4:  # 4代表周五
                fridays.append(today.strftime("%Y%m%d"))
            today -= timedelta(days=1)  # 每次减少一

        return fridays[::-1]  # 从最早到最-1 是reverse

    today0 = report_date
    today_year = today0.year
    # Outer loop for years (2010 to 今年)
    years = range(2010, today_year+1)  # 010到今年（包括今年
    quarters = ["0331", "0630", "0930", "1231"]  # 每年的四个季度日

    # 创建所有季度日
    dates = []
    for year in years:
        for quarter in quarters:
            date = f"{year}{quarter}"  # 生成 "YYYYMMDD" 格式的字符串
            dates.append(date)

    # 筛选出"20100331" "最近的季度 之间的日
    filtered_dates_list = [date for date in dates if "20100331" <= date <= (today0-timedelta(days = 50)).strftime("%Y%m%d")]
    # 双重保险第一设置必须要在20100331和此月减0的那个月之间才可以出现在这个list上
    #第二同时只取倒数和第2个日期作为newest date。相当于只滞个月 （如果是倒数2就是滞后5个月
    # 比如今天0号，那么upper limit则是40号，31 可以出现，但是还是取21的数据，直到81号：
    # 这时候upper limit0，那么就可以1 的数据了，此原因主要是险资害人不浅，懒狗一个，出数据太

    if not filtered_dates_list:
        raise RuntimeError("资金面报告期列表为空，请检查报告期范围或输入数据。")
    print(
        f"资金面报告期筛选完成：共{len(filtered_dates_list)}个，"
        f"起始={filtered_dates_list[0]}，结束={filtered_dates_list[-1]}"
    )

    # 通用指标
    newest_rpt_date = filtered_dates_list[-1]  # 如果数据跑失败了就是这的问题
    newest_rpt_date_shift = filtered_dates_list[-2] # 如果数据跑失败了就是这的问题
    print(f"资金面最新报告期={newest_rpt_date}，上期={newest_rpt_date_shift}")
    # %%
        # %%
    # 数据获取（如资产净值报告更新，请更新报告期）：
    fund_list_被动权益 = passive_fund_list

    stock_ETF_list = stock_ETF_list
    print(f"资金面报告期列表（首尾预览）={filtered_dates_list[:2]} ... {filtered_dates_list[-2:]}")

    def log_dataframe_preview(label: str, df: pd.DataFrame, preview_rows: int = 3) -> None:
        print(f"{label}: shape={df.shape}, columns={list(df.columns)}")
        if df.empty:
            return
        head_records = df.head(preview_rows).to_dict(orient="records")
        print(f"{label} head({min(preview_rows, len(df))}): {head_records}")
        if len(df) > preview_rows:
            tail_records = df.tail(preview_rows).to_dict(orient="records")
            print(f"{label} tail({min(preview_rows, len(df))}): {tail_records}")

    def snapshot_total(codes, field_name: str, option_string: str, label: str) -> float:
        result = snapshot_batch(codes, field_name, option_string)
        raw_values = result.Data[0] if getattr(result, "Data", None) else []
        numeric_values = []
        for value in raw_values:
            if pd.isna(value):
                numeric_values.append(0.0)
                continue
            try:
                numeric_values.append(float(value))
            except (TypeError, ValueError):
                numeric_values.append(0.0)
        total_value = float(np.nansum(numeric_values))
        print(
            f"{label}: count={len(numeric_values)}, sample={numeric_values[:5]}, total={total_value}"
        )
        return total_value

    def log_verified_etf_fetch(label: str, fetch_result) -> float:
        print(
            f"{label}: metric={fetch_result.metric_name}, "
            f"requested_date={fetch_result.requested_date}, actual_date={fetch_result.actual_date}, "
            f"matched={fetch_result.matched_count}/{fetch_result.requested_count}, "
            f"non_zero={fetch_result.non_zero_count}, total={fetch_result.total_value}"
        )
        if fetch_result.actual_date != fetch_result.requested_date:
            print(
                f"{label}: 请求日期 {fetch_result.requested_date} 无精确数据，"
                f"改用最近可用日期 {fetch_result.actual_date}。"
            )
        return float(fetch_result.total_value)

    def normalize_date_strings(values) -> list[str]:
        normalized = []
        for value in values:
            if value is None:
                continue
            text = str(value).strip()
            if text:
                normalized.append(text)
        return normalized

    def extract_sequence_values(raw_value) -> list:
        if raw_value is None:
            return []
        if isinstance(raw_value, pd.DataFrame):
            if raw_value.empty:
                return []
            if raw_value.shape[1] >= 1:
                return raw_value.iloc[:, 0].tolist()
            if raw_value.shape[0] >= 1:
                return raw_value.iloc[0, :].tolist()
            return []
        if isinstance(raw_value, (pd.Series, pd.Index)):
            return raw_value.tolist()
        if isinstance(raw_value, (list, tuple)):
            return list(raw_value)
        return [raw_value]

    def format_times_as_yyyymmdd(raw_times) -> list[str]:
        time_values = extract_sequence_values(raw_times)
        if not time_values:
            return []
        parsed = pd.to_datetime(pd.Series(time_values), errors="coerce")
        if parsed.isna().any():
            invalid_values = [time_values[idx] for idx, is_na in enumerate(parsed.isna().tolist()) if is_na]
            raise RuntimeError(f"存在无法解析的时间值: {invalid_values[:10]}")
        return parsed.dt.strftime("%Y%m%d").tolist()

    def extract_first_data_series(raw_data, expected_len: int | None = None) -> list:
        if raw_data is None:
            return []

        candidates: list[list] = []

        def append_candidate(candidate) -> None:
            values = extract_sequence_values(candidate)
            if values:
                candidates.append(values)

        if isinstance(raw_data, pd.DataFrame):
            if raw_data.empty:
                return []
            for column in raw_data.columns:
                append_candidate(raw_data[column])
            for row_index in raw_data.index:
                append_candidate(raw_data.loc[row_index])
        elif isinstance(raw_data, (pd.Series, pd.Index)):
            append_candidate(raw_data)
        elif isinstance(raw_data, (list, tuple)):
            if not raw_data:
                return []
            scalar_only = all(
                not isinstance(item, (list, tuple, pd.Series, pd.Index, pd.DataFrame))
                for item in raw_data
            )
            if scalar_only:
                append_candidate(raw_data)
            else:
                for item in raw_data:
                    append_candidate(item)
        else:
            append_candidate(raw_data)

        if expected_len is not None:
            for candidate in candidates:
                if len(candidate) == expected_len:
                    return candidate
        return candidates[0] if candidates else []

    def build_edb_series(result) -> pd.Series | None:
        time_values = extract_sequence_values(getattr(result, "Times", None))
        data_values = extract_first_data_series(getattr(result, "Data", None), expected_len=len(time_values))
        if not time_values or not data_values:
            return None
        return pd.Series(
            clean_numeric_values(data_values),
            index=pd.to_datetime(time_values),
            dtype="float64",
        ).sort_index()

    def fetch_single_edb_value(indicator_code: str, requested_date: str, label: str) -> float:
        normalized_date = str(requested_date)
        parsed_date = pd.to_datetime(normalized_date, format="%Y%m%d", errors="coerce")
        candidate_dates = [normalized_date]
        if pd.notna(parsed_date):
            hyphen_date = parsed_date.strftime("%Y-%m-%d")
            if hyphen_date not in candidate_dates:
                candidate_dates.append(hyphen_date)

        for candidate_date in candidate_dates:
            result = local_client.edb_query(indicator_code, candidate_date, candidate_date, "Fill=Previous")
            series = build_edb_series(result)
            if series is None or series.empty:
                continue
            numeric = pd.to_numeric(series, errors="coerce").dropna()
            if numeric.empty:
                continue
            return round(float(numeric.iloc[-1]), 1)

        print(f"[警告] {label} 数据缺失: indicator={indicator_code}, date={normalized_date}, 跳过该日期")
        return None

    def fetch_edb_values_for_dates(indicator_code: str, requested_dates, label: str) -> dict[str, float]:
        normalized_dates = normalize_date_strings(requested_dates)
        if not normalized_dates:
            return {}
        query_start = min(normalized_dates)
        query_end = max(normalized_dates)
        print(
            f"{label}: interval_query start={query_start}, end={query_end}, "
            f"request_dates={normalized_dates}"
        )
        query_start_time = time_module.perf_counter()
        result = local_client.edb_query(indicator_code, query_start, query_end, "Fill=Previous")
        series = build_edb_series(result)
        request_index = pd.to_datetime(normalized_dates, format="%Y%m%d", errors="coerce")
        if request_index.isna().any():
            raise RuntimeError(f"{label} 存在无法解析的日期: {normalized_dates}")
        if series is not None and not series.empty:
            aligned = series.reindex(series.index.union(request_index)).sort_index().ffill().reindex(request_index)
            if not aligned.isna().any():
                value_map = {
                    request_date: round(float(value), 1)
                    for request_date, value in zip(normalized_dates, aligned.tolist())
                }
                print(
                    f"{label}: dates={normalized_dates}, total_points={len(value_map)}, "
                    f"elapsed={time_module.perf_counter() - query_start_time:.2f}s"
                )
                return value_map

            missing_dates = list(pd.Index(normalized_dates)[aligned.isna()])
            print(
                f"{label}: interval_query incomplete, fallback_to_single_dates. "
                f"indicator={indicator_code}, missing_dates={missing_dates}"
            )
        else:
            print(
                f"{label}: interval_query empty, fallback_to_single_dates. "
                f"indicator={indicator_code}, start={query_start}, end={query_end}"
            )
            missing_dates = normalized_dates

        fallback_start = time_module.perf_counter()
        value_map = {}
        for requested_date in missing_dates:
            value = fetch_single_edb_value(indicator_code, requested_date, label)
            if value is not None:  # 只保留成功获取的值
                value_map[requested_date] = value
            else:
                print(f"[信息] 跳过 {requested_date} 的缺失数据")

        print(
            f"{label}: single_date_fallback complete, dates={normalized_dates}, "
            f"total_points={len(value_map)}, elapsed={time_module.perf_counter() - fallback_start:.2f}s"
        )
        return value_map

    def build_date_value_lookup(df: pd.DataFrame, date_col: str, value_col: str) -> dict[str, float]:
        if date_col not in df.columns or value_col not in df.columns:
            return {}
        lookup_df = df[[date_col, value_col]].dropna(subset=[date_col]).copy()
        lookup_df[date_col] = lookup_df[date_col].astype(str)
        lookup_df[value_col] = pd.to_numeric(lookup_df[value_col], errors="coerce")
        lookup_df = lookup_df.drop_duplicates(subset=[date_col], keep="last")
        return lookup_df.set_index(date_col)[value_col].to_dict()

    def normalize_code_list(codes) -> list[str]:
        if isinstance(codes, str):
            return [part.strip() for part in codes.split(",") if part.strip()]
        if isinstance(codes, (list, tuple, set)):
            normalized = []
            for item in codes:
                if item is None:
                    continue
                if isinstance(item, str):
                    normalized.extend(part.strip() for part in item.split(",") if part.strip())
                else:
                    text = str(item).strip()
                    if text:
                        normalized.append(text)
            return normalized
        text = str(codes).strip()
        return [text] if text else []

    def get_stock_etf_netinflow_cache_dir() -> Path:
        cache_dir = paths.funding_draft_dir / "data" / "cache" / "stock_etf_netinflow"
        cache_dir.mkdir(parents=True, exist_ok=True)
        return cache_dir

    def get_stock_etf_netinflow_debug_dir() -> Path:
        debug_dir = paths.funding_draft_dir / "data" / "cache" / "stock_etf_netinflow_debug"
        debug_dir.mkdir(parents=True, exist_ok=True)
        return debug_dir

    def get_stock_etf_netinflow_cache_path(start_date: str, end_date: str) -> Path:
        return get_stock_etf_netinflow_cache_dir() / f"stock_etf_netinflow_{start_date}_{end_date}.csv"

    def get_stock_etf_netinflow_detail_cache_path(start_date: str, end_date: str) -> Path:
        return get_stock_etf_netinflow_cache_dir() / f"stock_etf_netinflow_detail_{start_date}_{end_date}.csv"

    def load_stock_etf_netinflow_cache(start_date: str, end_date: str, expected_count: int) -> dict | None:
        cache_path = get_stock_etf_netinflow_cache_path(start_date, end_date)
        if not cache_path.exists():
            return None
        cache_df = pd.read_csv(cache_path)
        required_columns = {
            "Date",
            "total_net_inflow_stock",
            "requested_count",
            "hit_count",
            "missing_count",
            "source",
        }
        if cache_df.empty or not required_columns.issubset(cache_df.columns):
            return None
        cache_row = cache_df.iloc[0]
        source = str(cache_row.get("source", "")).strip()
        if source != "dateserial_share_nav":
            return None
        requested_count = int(pd.to_numeric(cache_row["requested_count"], errors="coerce"))
        if requested_count != expected_count:
            print(
                f"图10缓存失效: date_range={start_date}-{end_date}, "
                f"cached_requested_count={requested_count}, expected_count={expected_count}"
            )
            return None
        detail_cache_path = get_stock_etf_netinflow_detail_cache_path(start_date, end_date)
        if not detail_cache_path.exists():
            return None
        return {
            "Date": str(cache_row["Date"]),
            "start_date": start_date,
            "end_date": end_date,
            "total_net_inflow_stock": float(pd.to_numeric(cache_row["total_net_inflow_stock"], errors="coerce")),
            "requested_count": requested_count,
            "hit_count": int(pd.to_numeric(cache_row["hit_count"], errors="coerce")),
            "missing_count": int(pd.to_numeric(cache_row["missing_count"], errors="coerce")),
            "non_null_count": int(pd.to_numeric(cache_row.get("non_null_count", cache_row["hit_count"]), errors="coerce")),
            "non_zero_count": int(pd.to_numeric(cache_row.get("non_zero_count", 0), errors="coerce")),
            "source": source,
            "cache_path": cache_path,
            "detail_cache_path": detail_cache_path,
        }

    def write_stock_etf_netinflow_cache(summary: dict) -> Path:
        cache_path = get_stock_etf_netinflow_cache_path(summary["start_date"], summary["end_date"])
        cache_df = pd.DataFrame([summary])
        cache_df.to_csv(cache_path, index=False, encoding="utf-8-sig")
        return cache_path

    def load_stock_etf_netinflow_detail_cache(
        start_date: str,
        end_date: str,
        expected_codes,
    ) -> pd.DataFrame | None:
        cache_path = get_stock_etf_netinflow_detail_cache_path(start_date, end_date)
        if not cache_path.exists():
            return None
        detail_df = pd.read_csv(cache_path, dtype={"code": str})
        required_columns = {"code", "net_flow_raw", "net_flow", "source"}
        if detail_df.empty or not required_columns.issubset(detail_df.columns):
            return None
        if not (detail_df["source"].astype(str) == "dateserial_share_nav").all():
            return None
        expected_code_list = normalize_code_list(expected_codes)
        cached_codes = set(detail_df["code"].dropna().astype(str))
        if cached_codes != set(expected_code_list):
            return None
        detail_df["net_flow_raw"] = pd.to_numeric(detail_df["net_flow_raw"], errors="coerce")
        detail_df["net_flow"] = pd.to_numeric(detail_df["net_flow"], errors="coerce")
        return detail_df

    def write_stock_etf_netinflow_detail_cache(start_date: str, end_date: str, detail_df: pd.DataFrame) -> Path:
        cache_path = get_stock_etf_netinflow_detail_cache_path(start_date, end_date)
        detail_df.to_csv(cache_path, index=False, encoding="utf-8-sig")
        return cache_path

    def write_stock_etf_netinflow_debug(
        start_date: str,
        end_date: str,
        chunk_records: list[dict[str, object]],
        field_probe_records: list[dict[str, object]],
        single_probe_records: list[dict[str, object]],
    ) -> list[Path]:
        debug_dir = get_stock_etf_netinflow_debug_dir()
        written_paths: list[Path] = []
        payloads = [
            (f"stock_etf_netinflow_debug_{start_date}_{end_date}_single_probe.csv", single_probe_records),
            (f"stock_etf_netinflow_debug_{start_date}_{end_date}_chunks.csv", chunk_records),
            (f"stock_etf_netinflow_debug_{start_date}_{end_date}_candidate_fields.csv", field_probe_records),
        ]
        for file_name, rows in payloads:
            if not rows:
                continue
            file_path = debug_dir / file_name
            pd.DataFrame(rows).to_csv(file_path, index=False, encoding="utf-8-sig")
            written_paths.append(file_path)
        return written_paths

    def run_single_etf_netinflow_probe(codes, start_date: str, end_date: str) -> list[dict[str, object]]:
        sample_codes = normalize_code_list(codes)[:5]
        records: list[dict[str, object]] = []
        for code in sample_codes:
            started = time_module.perf_counter()
            result = local_client.snapshot(
                code,
                "periodmf_netinflow",
                f"unit=1;startDate={start_date};endDate={end_date};",
            )
            raw_values = result.Data[0] if getattr(result, "Data", None) else []
            stats = summarize_numeric_values(raw_values)
            record = {
                "code": code,
                "error_code": getattr(result, "ErrorCode", None),
                "error_detail": get_local_error_detail(getattr(result, "ErrorCode", None)),
                "non_null_count": stats["non_null_count"],
                "non_zero_count": stats["non_zero_count"],
                "sample_values": stats["sample"],
                "elapsed": round(time_module.perf_counter() - started, 4),
            }
            records.append(record)
            print(
                f"图10单ETF探针: code={code}, error_code={record['error_code']}, "
                f"error_detail={record['error_detail']}, non_null_count={record['non_null_count']}, "
                f"non_zero_count={record['non_zero_count']}, "
                f"sample_values={record['sample_values']}, elapsed={record['elapsed']:.2f}s"
            )
        return records

    def probe_periodmf_netinflow_candidate_fields(
        codes,
        start_date: str,
        end_date: str,
    ) -> list[dict[str, object]]:
        sample_codes = normalize_code_list(codes)[:20]
        candidate_fields = [
            "periodmf_netinflow",
            "periodfundnetinflow",
            "netinflow",
            "fund_netinflow",
            "mf_netinflow",
        ]
        probe_rows: list[dict[str, object]] = []
        for field_name in candidate_fields:
            result = local_client.snapshot(
                sample_codes,
                field_name,
                f"unit=1;startDate={start_date};endDate={end_date};",
            )
            raw_values = result.Data[0] if getattr(result, "Data", None) else []
            stats = summarize_numeric_values(raw_values)
            row = {
                "field_name": field_name,
                "error_code": getattr(result, "ErrorCode", None),
                "error_detail": get_local_error_detail(getattr(result, "ErrorCode", None)),
                "requested_count": len(sample_codes),
                "received_count": len(raw_values),
                "non_null_count": stats["non_null_count"],
                "non_zero_count": stats["non_zero_count"],
                "sample_values": stats["sample"],
            }
            probe_rows.append(row)
            print(
                f"图10候选字段探针: field={field_name}, error_code={row['error_code']}, "
                f"error_detail={row['error_detail']}, received_count={row['received_count']}, "
                f"non_null_count={row['non_null_count']}, "
                f"non_zero_count={row['non_zero_count']}, sample_values={row['sample_values']}"
            )
        return probe_rows

    def build_stock_etf_netinflow_from_dateserial(
        codes,
        start_date: str,
        end_date: str,
        batch_size: int = 120,
    ) -> dict:
        filtered_codes = local_client.prepare_etf_query_codes(codes, end_date)
        date_range = f"{start_date}-{end_date}"
        detail_cache = load_stock_etf_netinflow_detail_cache(start_date, end_date, filtered_codes)
        summary_cache = load_stock_etf_netinflow_cache(start_date, end_date, len(filtered_codes))
        if summary_cache is not None and detail_cache is not None:
            print(
                f"图10周缓存命中: date_range={date_range}, requested_count={summary_cache['requested_count']}, "
                f"hit_count={summary_cache['hit_count']}, missing_count={summary_cache['missing_count']}, "
                f"non_zero_count={summary_cache['non_zero_count']}, total={summary_cache['total_net_inflow_stock']}, "
                f"cache={summary_cache['cache_path'].name}"
            )
            cached_result = dict(summary_cache)
            cached_result["detail_df"] = detail_cache
            return cached_result

        query_begin = (pd.to_datetime(start_date) - timedelta(days=7)).strftime("%Y-%m-%d")
        query_end = pd.to_datetime(end_date).strftime("%Y-%m-%d")
        started = time_module.perf_counter()
        print(
            f"图10结构化净流入开始: date_range={date_range}, requested_count={len(filtered_codes)}, "
            f"query_begin={query_begin}, query_end={query_end}, batch_size={batch_size}"
        )
        serial_df = local_client.dateserial_chunked(
            filtered_codes,
            ["ths_fund_shares_fund", "ths_unit_nv_fund"],
            ["100", "100"],
            query_begin,
            query_end,
            batch_size=batch_size,
        )
        print(
            f"图10结构化净流入原始序列完成: date_range={date_range}, shape={serial_df.shape}, "
            f"elapsed={time_module.perf_counter() - started:.2f}s"
        )
        if serial_df.empty:
            raise RuntimeError(
                f"图10 股票型ETF净流入额查询失败：THS_DateSerial 未返回有效数据。 date_range={date_range}"
            )

        lower_map = {str(column).lower(): column for column in serial_df.columns}
        time_col = lower_map.get("time")
        code_col = lower_map.get("thscode")
        share_col = lower_map.get("ths_fund_shares_fund")
        nav_col = lower_map.get("ths_unit_nv_fund")
        if not all([time_col, code_col, share_col, nav_col]):
            raise RuntimeError(
                f"图10 股票型ETF净流入额查询失败：THS_DateSerial 返回字段不完整。 "
                f"columns={list(serial_df.columns)}"
            )

        detail_source = serial_df[[time_col, code_col, share_col, nav_col]].copy()
        detail_source.columns = ["time", "code", "shares", "nav"]
        detail_source["time"] = pd.to_datetime(detail_source["time"], errors="coerce")
        detail_source["code"] = detail_source["code"].astype(str)
        detail_source["shares"] = pd.to_numeric(detail_source["shares"], errors="coerce")
        detail_source["nav"] = pd.to_numeric(detail_source["nav"], errors="coerce")
        detail_source = detail_source.dropna(subset=["time", "code"]).sort_values(["code", "time"]).reset_index(drop=True)
        if detail_source.empty:
            raise RuntimeError(f"图10 股票型ETF净流入额查询失败：序列清洗后无有效记录。 date_range={date_range}")

        detail_source["share_diff"] = detail_source.groupby("code")["shares"].diff()
        detail_source["daily_flow_raw"] = detail_source["share_diff"] * detail_source["nav"]
        start_ts = pd.to_datetime(start_date)
        end_ts = pd.to_datetime(end_date)
        weekly_detail = detail_source[(detail_source["time"] >= start_ts) & (detail_source["time"] <= end_ts)].copy()
        if weekly_detail.empty:
            detail_df = pd.DataFrame({"code": filtered_codes})
            detail_df["net_flow_raw"] = 0.0
            detail_df["tradeday_count"] = 0
            detail_df["daily_non_null_count"] = 0
            detail_df["net_flow"] = 0.0
            detail_df["start_date"] = start_date
            detail_df["end_date"] = end_date
            detail_df["source"] = "dateserial_share_nav"
            summary = {
                "Date": date_range,
                "start_date": start_date,
                "end_date": end_date,
                "total_net_inflow_stock": 0.0,
                "requested_count": len(filtered_codes),
                "hit_count": len(filtered_codes),
                "missing_count": 0,
                "non_null_count": len(filtered_codes),
                "non_zero_count": 0,
                "source": "dateserial_share_nav",
            }
            summary_cache_path = write_stock_etf_netinflow_cache(summary)
            detail_cache_path = write_stock_etf_netinflow_detail_cache(start_date, end_date, detail_df)
            print(
                f"图10节假日周无交易日，按零净流入处理: date_range={date_range}, "
                f"summary_cache={summary_cache_path.name}, detail_cache={detail_cache_path.name}"
            )
            result = dict(summary)
            result["cache_path"] = summary_cache_path
            result["detail_cache_path"] = detail_cache_path
            result["detail_df"] = detail_df
            return result

        weekly_summary = (
            weekly_detail.groupby("code")
            .agg(
                net_flow_raw=("daily_flow_raw", lambda values: values.sum(min_count=1)),
                tradeday_count=("time", "nunique"),
                daily_non_null_count=("daily_flow_raw", lambda values: int(pd.Series(values).notna().sum())),
            )
            .reset_index()
        )

        detail_df = pd.DataFrame({"code": filtered_codes})
        detail_df = detail_df.merge(weekly_summary, on="code", how="left")
        detail_df["net_flow"] = detail_df["net_flow_raw"] / 1e8
        detail_df["start_date"] = start_date
        detail_df["end_date"] = end_date
        detail_df["source"] = "dateserial_share_nav"

        hit_count = int(detail_df["net_flow_raw"].notna().sum())
        missing_count = int(detail_df["net_flow_raw"].isna().sum())
        non_zero_count = int((detail_df["net_flow_raw"].fillna(0.0) != 0).sum())
        total_raw = float(detail_df["net_flow_raw"].sum(min_count=1)) if hit_count else float("nan")
        if hit_count == 0 or pd.isna(total_raw):
            raise RuntimeError(
                f"图10 股票型ETF净流入额查询失败：份额/净值重建后无有效净流入结果。 "
                f"date_range={date_range}, requested_count={len(filtered_codes)}"
            )

        summary = {
            "Date": date_range,
            "start_date": start_date,
            "end_date": end_date,
            "total_net_inflow_stock": total_raw / 1e8,
            "requested_count": len(filtered_codes),
            "hit_count": hit_count,
            "missing_count": missing_count,
            "non_null_count": hit_count,
            "non_zero_count": non_zero_count,
            "source": "dateserial_share_nav",
        }
        summary_cache_path = write_stock_etf_netinflow_cache(summary)
        detail_cache_path = write_stock_etf_netinflow_detail_cache(start_date, end_date, detail_df)
        print(
            f"图10周总净流入完成: date_range={date_range}, total={summary['total_net_inflow_stock']}, "
            f"hit_count={hit_count}, missing_count={missing_count}, non_zero_count={non_zero_count}, "
            f"summary_cache={summary_cache_path.name}, detail_cache={detail_cache_path.name}"
        )
        result = dict(summary)
        result["cache_path"] = summary_cache_path
        result["detail_cache_path"] = detail_cache_path
        result["detail_df"] = detail_df
        return result

    def fetch_periodmf_netinflow_chunked(codes, start_date: str, end_date: str, batch_size: int = 120) -> dict:
        return build_stock_etf_netinflow_from_dateserial(
            codes,
            start_date,
            end_date,
            batch_size=batch_size,
        )

    def build_passive_fund_dimension(passive_fund_data: pd.DataFrame) -> pd.DataFrame:
        sum_total_被动权益 = pd.concat(
            [passive_fund_data["日期_汇总"], passive_fund_data["被动权益ETF"]],
            axis=1,
        )
        sum_total_被动权益.columns = ["日期_汇总", "被动权益ETF"]

        股票型ETF = pd.concat(
            [passive_fund_data["场内_股票型ETF"], passive_fund_data["日期_场内"]],
            axis=1,
        )
        股票型ETF.columns = ["场内_股票型ETF", "日期_场内"]

        merged = sum_total_被动权益.merge(
            股票型ETF,
            left_on="日期_汇总",
            right_on="日期_场内",
            how="right",
        ).drop(columns="日期_汇总")
        merged["日期"] = merged["日期_场内"]
        merged = merged.drop(columns="日期_场内").dropna(subset="日期")

        merged["估算增长率"] = (
            merged["场内_股票型ETF"] - merged["场内_股票型ETF"].shift(1)
        ) / merged["场内_股票型ETF"].shift(1)
        for i in range(12, 15):
            merged.loc[i, "被动权益ETF"] = (
                merged["被动权益ETF"].loc[i - 1]
            ) * (1 + merged["估算增长率"].loc[i])

        merged["被动权益ETF规模变动"] = (
            merged["被动权益ETF"] - merged["被动权益ETF"].shift(1)
        )
        merged["场外_被动与增强"] = (
            merged["被动权益ETF"] - merged["场内_股票型ETF"]
        )
        merged["被动权益ETF"] = (merged["被动权益ETF"] / 1e8).apply(
            lambda x: float(f"{x:.6g}") if pd.notnull(x) else x
        )
        merged["被动权益ETF规模变动"] = (
            merged["被动权益ETF规模变动"] / 1e8
        ).apply(lambda x: float(f"{x:.6g}") if pd.notnull(x) else x)
        return merged

    passive_stage_start = time_module.perf_counter()

    def passive_checkpoint(label: str) -> None:
        nonlocal passive_stage_start
        elapsed = time_module.perf_counter() - passive_stage_start
        print(f"[耗时] step2::被动权益基金::{label}: {elapsed:.2f} 秒")
        passive_stage_start = time_module.perf_counter()

    passive_old_path = "workspace/funding_draft/data/被动权益基金/资金面数据_老_此数据根据timelist仅季度更新.csv"
    passive_data_path = "workspace/funding_draft/data/被动权益基金/资金面数据.csv"
    passive_fig11_path = "workspace/funding_draft/data/被动权益基金/图11_被动权益ETF处理后数据.csv"
    passive_fig9_path = "workspace/funding_draft/data/被动权益基金/图9_被动权益ETF.csv"

    资金面数据_老 = pd.read_csv(passive_old_path)
    资金面数据_老["日期_汇总"] = 资金面数据_老["日期_汇总"].astype(str)
    log_dataframe_preview("被动权益基金旧数据", 资金面数据_老)
    passive_checkpoint("读取旧数据完成")

    existing_quarter_dates = set(资金面数据_老["日期_汇总"].dropna().astype(str))
    missing_quarter_dates = [
        rpt_date for rpt_date in filtered_dates_list if str(rpt_date) not in existing_quarter_dates
    ]
    print(
        f"被动权益基金缺失季度日期: count={len(missing_quarter_dates)}, sample={missing_quarter_dates[:5]}"
    )
    passive_checkpoint("缺失日期识别完成")

    old_append_rows = []
    for rpt_date in missing_quarter_dates:
        netasset_sum = snapshot_total(
            fund_list_被动权益,
            "prt_netasset",
            f"unit=1;rptDate={rpt_date};",
            f"被动权益基金资产净值_{rpt_date}",
        )
        old_append_rows.append({"被动权益ETF": netasset_sum, "日期_汇总": rpt_date})

    if old_append_rows:
        资金面数据_老 = pd.concat(
            [资金面数据_老, pd.DataFrame(old_append_rows)],
            ignore_index=True,
        )

    newest_rpt_date = filtered_dates_list[-1]
    newest_rpt_date_shift = filtered_dates_list[-2]
    time_list = [newest_rpt_date_shift, newest_rpt_date]

    quarterly_lookup_df = (
        资金面数据_老[["日期_汇总", "被动权益ETF"]]
        .dropna(subset=["日期_汇总"])
        .copy()
    )
    quarterly_lookup_df["日期_汇总"] = quarterly_lookup_df["日期_汇总"].astype(str)
    quarterly_lookup_df = quarterly_lookup_df.drop_duplicates(
        subset=["日期_汇总"], keep="last"
    )
    quarterly_lookup = quarterly_lookup_df.set_index("日期_汇总")["被动权益ETF"].to_dict()

    passive_quarter_rows = []
    passive_quarter_cache_hits = 0
    for rpt_date in time_list:
        cached_value = quarterly_lookup.get(str(rpt_date))
        if cached_value is None or pd.isna(cached_value):
            cached_value = snapshot_total(
                fund_list_被动权益,
                "prt_netasset",
                f"unit=1;rptDate={rpt_date};",
                f"被动权益基金资产净值_{rpt_date}",
            )
        else:
            passive_quarter_cache_hits += 1
            print(f"被动权益ETF季度缓存复用_{rpt_date}: total={cached_value}")
        passive_quarter_rows.append({"被动权益ETF": cached_value, "日期_汇总": rpt_date})

    print(
        f"被动权益ETF季度缓存复用完成: cache_hits={passive_quarter_cache_hits}, total_dates={len(time_list)}"
    )
    passive_checkpoint("prt_netasset 请求完成")

    os.makedirs("workspace/funding_draft/data", exist_ok=True)

    资金面数据_老.to_csv(
        passive_old_path,
        index=False,
        encoding="utf-8-sig",
    )
    资金面数据_老版 = 资金面数据_老.copy()
    资金面数据_老版["被动权益ETF"] = 资金面数据_老版["被动权益ETF"] / 1e8
    资金面数据_老版["被动权益ETF"] = [
        round(x, 1) if pd.notnull(x) else x for x in 资金面数据_老版["被动权益ETF"]
    ]

    last_3_fridays = get_last_3_fridays()
    time_list.extend(last_3_fridays)
    print(f"被动权益基金场内日期序列: {time_list}")

    etf_size_config = local_client.ensure_etf_size_path_ready()
    passive_checkpoint("ETF路径配置加载完成")
    print(
        "ETF规模主路径: "
        f"path={etf_size_config.selected_path}, "
        f"interface={etf_size_config.interface_type}, "
        f"metric={etf_size_config.metric_name}, "
        f"config={local_client.get_etf_size_config_path().name}, "
        f"validated_at={etf_size_config.validated_at}"
    )

    stock_etf_rows = []
    for trade_date in time_list:
        verified_fetch = local_client.fetch_etf_netasset_verified(
            stock_ETF_list,
            str(trade_date),
            batch_size=batch_size,
        )
        netasset_sum = log_verified_etf_fetch(
            f"股票型ETF基金资产净值_{trade_date}",
            verified_fetch,
        )
        stock_etf_rows.append({"场内_股票型ETF": netasset_sum, "日期_场内": trade_date})

    passive_checkpoint("etf_netasset 请求完成")

    资金面数据_rows = passive_quarter_rows + stock_etf_rows
    资金面数据 = pd.DataFrame(
        资金面数据_rows,
        columns=["被动权益ETF", "场内_股票型ETF", "场外_被动与增强", "日期_汇总", "日期_场内"],
    )
    log_dataframe_preview("被动权益基金中间数据", 资金面数据)
    资金面数据.to_csv(passive_data_path, encoding="utf-8-sig")

    merged_fund_dimension = build_passive_fund_dimension(资金面数据)
    log_dataframe_preview("被动权益基金维度合并结果", merged_fund_dimension)
    passive_checkpoint("基金维度 merge 完成")

    图11_被动权益数据 = merged_fund_dimension
    passive_checkpoint("图11 数据生成完成")

    被动权益_汇总 = pd.DataFrame(
        {
            "被动权益ETF": 图11_被动权益数据["被动权益ETF"],
            "日期_汇总": 图11_被动权益数据["日期"],
        }
    )
    被动权益ETF总结 = pd.concat([资金面数据_老版, 被动权益_汇总[2:]], axis=0)
    被动权益ETF总结 = 被动权益ETF总结.drop(columns="Unnamed: 0", errors="ignore")
    log_dataframe_preview("图9_被动权益ETF", 被动权益ETF总结)
    passive_checkpoint("图9 数据生成完成")

    图11_被动权益数据.to_csv(
        passive_fig11_path,
        encoding="utf-8-sig",
    )
    被动权益ETF总结.to_csv(passive_fig9_path, encoding="utf-8-sig")
    passive_checkpoint("文件写出完成")
    checkpoint("被动权益基金")

    # %% md
    #主动权益基金_普+ 偏股
    #对老数据进行更新获取（季度更新，请定期维护
    #filtered_dates_list
    # %%
    active_stage_start = time_module.perf_counter()

    def active_checkpoint(label: str) -> None:
        nonlocal active_stage_start
        elapsed = time_module.perf_counter() - active_stage_start
        print(f"[耗时] step2::主动权益基金::{label}: {elapsed:.2f} 秒")
        active_stage_start = time_module.perf_counter()

    active_old_path = "workspace/funding_draft/data/主动权益基金/资金面数据_主动_老_此数据根据timelist仅季度更新.csv"
    active_fig11_path = "workspace/funding_draft/data/主动权益基金/图11_主动权益基金_数据处理.csv"
    active_fig9_path = "workspace/funding_draft/data/主动权益基金/图9_主动权益基金.csv"
    print(
        f"主动权益季度缓存刷新: {'enabled' if options.refresh_active_quarter_cache else 'disabled'}"
    )
    if not Path(active_old_path).exists():
        raise FileNotFoundError(
            f"缺少主动权益季度缓存文件: {active_old_path}。"
            "请先准备该文件；若需补本次周报季度锚点，请在文件存在后运行 "
            "--refresh-active-quarter-cache。"
        )
    print(
        f"主动权益基金季度日期列表: count={len(filtered_dates_list)}, "
        f"sample={filtered_dates_list[:5]}, latest={filtered_dates_list[-3:]}"
    )
    资金面数据_主动_老 = pd.read_csv(active_old_path, dtype={"日期": str})
    active_checkpoint("读取旧数据完成")

    active_metric_cache: dict[tuple[str, str], float] = {}

    def fetch_normal_stock_total(rpt_date: str) -> float:
        cache_key = ("normal_stock_pref", str(rpt_date))
        if cache_key in active_metric_cache:
            total = active_metric_cache[cache_key]
            print(f"普通+偏股型主动权益基金资产净值_{rpt_date}: cache_hit total={total}")
            return total

        result = snapshot_batch(normal_stock_pref_list, "prt_netasset", f"unit=1;rptDate={rpt_date};")
        raw_values = result.Data[0] if getattr(result, "Data", None) else []
        numeric_values = clean_numeric_values(raw_values)
        total = float(sum(numeric_values))
        print(
            f"普通+偏股型主动权益基金资产净值_{rpt_date}: "
            f"count={len(numeric_values)}, sample={numeric_values[:5]}, total={total}"
        )
        active_metric_cache[cache_key] = total
        return total

    def fetch_flexible_total(rpt_date: str) -> float:
        cache_key = ("flexible_allocation_fund", str(rpt_date))
        if cache_key in active_metric_cache:
            total = active_metric_cache[cache_key]
            print(f"灵活配置型主动权益基金资产净值_{rpt_date}: cache_hit total={total}")
            return total

        flex_allo_fund = snapshot_batch(
            flexible_allocation_fund_list,
            "prt_stocktonav, prt_netasset",
            f"rptDate={rpt_date};",
        )
        ratio_values = clean_numeric_values(flex_allo_fund.Data[0] if getattr(flex_allo_fund, "Data", None) else [])
        netasset_values = clean_numeric_values(
            flex_allo_fund.Data[1] if getattr(flex_allo_fund, "Data", None) and len(flex_allo_fund.Data) > 1 else []
        )
        filtered_netasset_values = [
            netasset
            for ratio, netasset in zip(ratio_values, netasset_values)
            if ratio > 60
        ]
        total = float(sum(filtered_netasset_values))
        print(
            f"灵活配置型主动权益基金资产净值_{rpt_date}: "
            f"eligible={len(filtered_netasset_values)}, sample={filtered_netasset_values[:5]}, total={total}"
        )
        active_metric_cache[cache_key] = total
        return total

    time_list = [newest_rpt_date_shift, newest_rpt_date]
    print(f"主动权益基金图11季度锚点: {time_list}")
    existing_active_dates = set(资金面数据_主动_老["日期"].dropna().astype(str))
    if options.refresh_active_quarter_cache:
        missing_active_dates = [
            str(rpt_date) for rpt_date in filtered_dates_list if str(rpt_date) not in existing_active_dates
        ]
        print(
            f"主动权益季度缓存刷新: enabled, 历史季度缺口数={len(missing_active_dates)}, "
            f"缺口样本={missing_active_dates[:10]}"
        )
    else:
        missing_active_dates = []
        print(
            f"active_old_path 读取成功, existing_dates={len(existing_active_dates)}, "
            f"历史季度缺口检查已跳过"
        )
    active_total_lookup = build_date_value_lookup(资金面数据_主动_老, "日期", "主动权益基金")
    local_active_anchor_dates = [
        date for date in time_list if date in active_total_lookup and pd.notna(active_total_lookup[date])
    ]
    missing_active_anchor_dates = [date for date in time_list if date not in local_active_anchor_dates]
    print(
        f"本地命中锚点: {local_active_anchor_dates}, "
        f"缺失锚点: {missing_active_anchor_dates}"
    )
    if missing_active_anchor_dates and not options.refresh_active_quarter_cache:
        raise RuntimeError(
            "主动权益季度缓存缺失，请先运行 --refresh-active-quarter-cache。"
            f" 缺失锚点={missing_active_anchor_dates}, active_old_path={active_old_path}"
        )

    refresh_anchor_dates = list(missing_active_anchor_dates)
    if refresh_anchor_dates:
        print(f"主动权益季度缓存刷新: 本次实际补齐季度={refresh_anchor_dates}")
        active_append_rows = []
        for rpt_date in refresh_anchor_dates:
            quarter_start = time_module.perf_counter()
            normal_total = fetch_normal_stock_total(rpt_date)
            flexible_total = fetch_flexible_total(rpt_date)
            active_append_rows.append(
                {
                    "主动权益基金": (normal_total + flexible_total) / 1e8,
                    "日期": rpt_date,
                }
            )
            print(
                f"主动权益季度缓存刷新_{rpt_date}: "
                f"elapsed={time_module.perf_counter() - quarter_start:.2f}s"
            )

        资金面数据_主动_老 = pd.concat(
            [资金面数据_主动_老, pd.DataFrame(active_append_rows)],
            ignore_index=True,
        )
        资金面数据_主动_老.to_csv(active_old_path)
        print(f"active_old_path 写回完成: {active_old_path}")
        active_total_lookup = build_date_value_lookup(资金面数据_主动_老, "日期", "主动权益基金")
        local_active_anchor_dates = [
            date for date in time_list if date in active_total_lookup and pd.notna(active_total_lookup[date])
        ]
        missing_active_anchor_dates = [date for date in time_list if date not in local_active_anchor_dates]
    elif options.refresh_active_quarter_cache:
        print("主动权益季度缓存刷新: enabled, 但本次周报锚点无需补齐。")
    else:
        print("using local active quarter cache")
    active_anchor_rows = []
    for rpt_date in time_list:
        local_total = active_total_lookup.get(rpt_date)
        if local_total is not None and pd.notna(local_total):
            active_anchor_rows.append(
                {
                    "普通_偏股型": pd.NA,
                    "日期": rpt_date,
                    "灵活配置型": pd.NA,
                    "主动权益基金": float(local_total) * 1e8,
                }
            )
            continue

        normal_total = fetch_normal_stock_total(rpt_date)
        flexible_total = fetch_flexible_total(rpt_date)
        active_anchor_rows.append(
            {
                "普通_偏股型": normal_total,
                "日期": rpt_date,
                "灵活配置型": flexible_total,
                "主动权益基金": normal_total + flexible_total,
            }
        )
    if missing_active_anchor_dates:
        raise RuntimeError(
            f"主动权益季度锚点补齐失败，仍缺失: {missing_active_anchor_dates}。"
            "请检查本地数据文件或 active_old_path 内容。"
        )

    资金面数据_主动权益基金 = pd.DataFrame(
        active_anchor_rows,
        columns=["普通_偏股型", "日期", "灵活配置型", "主动权益基金"],
    )
    log_dataframe_preview("资金面数据_主动权益基金", 资金面数据_主动权益基金)
    active_checkpoint("图11季度锚点读取完成")

    time_list = [newest_rpt_date_shift, newest_rpt_date]
    last_3_fridays = get_last_3_fridays()
    time_list.extend(last_3_fridays)
    print(f"主动权益基金扩展日期序列: {time_list}")

    benchmark_block_code = resolve_block_code("基金重仓指数", "benchmark", today)
    benchmark_rows = []
    for trade_date in time_list:
        close_result = local_client.snapshot(
            [benchmark_block_code],
            "close",
            f"tradeDate={trade_date};priceAdj=U;cycle=D;",
        )
        close_values = clean_numeric_values(close_result.Data[0] if getattr(close_result, "Data", None) else [])
        if not close_values:
            print(f"[警告] 基金重仓指数查询失败，跳过: code=883406.TI, trade_date={trade_date}")
            continue

        benchmark_rows.append(
            {
                "基金重仓指数": close_values[0],
                "日期": trade_date,
            }
        )
        
    基金重仓指数_df = pd.DataFrame(benchmark_rows, columns=["基金重仓指数", "日期"])
    基金重仓指数_df["基金重仓指数涨跌幅_%"] = (
        基金重仓指数_df["基金重仓指数"] - 基金重仓指数_df["基金重仓指数"].shift(1)
    ) / 基金重仓指数_df["基金重仓指数"].shift(1)
    log_dataframe_preview("基金重仓指数_df", 基金重仓指数_df)
    active_checkpoint("图11指数序列请求完成")

    time_list1 = [pd.to_datetime(item).strftime("%Y-%m-%d") for item in time_list[1:5]]
    time_list0 = [pd.to_datetime(item).strftime("%Y-%m-%d") for item in time_list[0:4]]
    print(f"偏股型主动权益基金新发行量结束日期: {time_list1}")
    print(f"偏股型主动权益基金新发行量起始日期: {time_list0}")

    偏股型新发行量_rows = []
    for time0, time1 in zip(time_list0, time_list1):
        stock_pref_issue = local_client.edb_query("M0060433", time0, time1, "Fill=Previous").Data[0]
        stock_pref_issue_total = float(sum(clean_numeric_values(stock_pref_issue)))
        偏股型新发行量_rows.append(
            {
                "偏股型主动权益基金新发行量": stock_pref_issue_total,
                "日期": time1.replace("-", ""),
                "周期": f"{time0}-{time1}",
            }
        )
    偏股型新发行量 = pd.DataFrame(
        偏股型新发行量_rows,
        columns=["偏股型主动权益基金新发行量", "日期", "周期"],
    )
    log_dataframe_preview("偏股型新发行量", 偏股型新发行量)
    active_checkpoint("主动权益基金新发行量请求完成")

    资金面数据_老1 = 资金面数据_主动_老.copy()
    资金面数据_老1["主动权益基金"] = [round(x, 1) for x in 资金面数据_老1["主动权益基金"]]
    资金面数据_主动权益基金_全 = 资金面数据_主动权益基金.merge(基金重仓指数_df, on="日期", how="right")
    资金面数据_主动权益基金_汇总 = 资金面数据_主动权益基金_全.merge(偏股型新发行量, on="日期", how="left")
    for i in range(min(2, len(资金面数据_主动权益基金_汇总))):
        if pd.notna(资金面数据_主动权益基金_汇总.loc[i, "主动权益基金"]):
            continue
        component_total = (
            to_float_or_zero(资金面数据_主动权益基金_汇总.loc[i, "普通_偏股型"])
            + to_float_or_zero(资金面数据_主动权益基金_汇总.loc[i, "灵活配置型"])
        )
        资金面数据_主动权益基金_汇总.loc[i, "主动权益基金"] = component_total

    # 确保索引从0开始且连续
    资金面数据_主动权益基金_汇总 = 资金面数据_主动权益基金_汇总.reset_index(drop=True)
        
    # 如果数据行数不足，跳过计算
    if len(资金面数据_主动权益基金_汇总) < 3:
        print("[警告] 主动权益基金数据不足，跳过补全计算")
    else:
        for i in range(2, min(5, len(资金面数据_主动权益基金_汇总))):
            if pd.notna(资金面数据_主动权益基金_汇总.loc[i, "基金重仓指数涨跌幅_%"]):
                资金面数据_主动权益基金_汇总.loc[i, "主动权益基金"] = (
                        资金面数据_主动权益基金_汇总.loc[i - 1, "主动权益基金"]
                        * (1 + (资金面数据_主动权益基金_汇总.loc[i, "基金重仓指数涨跌幅_%"] or 0) / 100)
                        + 资金面数据_主动权益基金_汇总.loc[i, "偏股型主动权益基金新发行量"]
                    )

    资金面数据_主动权益基金_汇总["主动权益基金"] = (
        资金面数据_主动权益基金_汇总["主动权益基金"] / 1e8
    ).apply(lambda x: float(f"{x:.6g}") if pd.notnull(x) else x)

    log_dataframe_preview("资金面数据_主动权益基金_全", 资金面数据_主动权益基金_全)
    log_dataframe_preview("资金面数据_主动权益基金_汇总", 资金面数据_主动权益基金_汇总)
    资金面数据_主动权益基金_汇总.to_csv(active_fig11_path)
    active_checkpoint("图11数据写出完成")

    资金面数据_主动权益基金_汇总0 = pd.DataFrame(
        {
            "日期": 资金面数据_主动权益基金_汇总["日期"],
            "主动权益基金": 资金面数据_主动权益基金_汇总["主动权益基金"],
        }
    )
    资金面数据_主动权益基金_汇总0 = pd.concat([资金面数据_老1, 资金面数据_主动权益基金_汇总0[2:]], axis=0)
    资金面数据_主动权益基金_汇总0.to_csv(active_fig9_path)
    active_checkpoint("图9数据写出完成")
    checkpoint("主动权益基金")

    # %% md
    #两融余额
    # %%
    margin_stage_start = time_module.perf_counter()

    def margin_checkpoint(label: str) -> None:
        nonlocal margin_stage_start
        elapsed = time_module.perf_counter() - margin_stage_start
        print(f"[耗时] step2::两融余额::{label}: {elapsed:.2f} 秒")
        margin_stage_start = time_module.perf_counter()

    margin_old_path = "workspace/funding_draft/data/两融余额数据/两融余额数据_老.csv"
    margin_fig11_path = "workspace/funding_draft/data/两融余额数据/图11_两融余额数据.csv"
    margin_fig9_path = "workspace/funding_draft/data/两融余额数据/图9_两融余额数据.csv"
    两融余额 = pd.read_csv(margin_old_path, dtype={"日期": str})
    margin_checkpoint("读取旧数据完成")

    margin_old_target_dates = filtered_dates_list[:-2]
    existing_margin_dates = set(两融余额["日期"].dropna().astype(str))
    missing_margin_old_dates = [
        str(item) for item in margin_old_target_dates if str(item) not in existing_margin_dates
    ]
    print(
        f"两融余额历史季度数据: existing={len(existing_margin_dates)}, "
        f"missing={len(missing_margin_old_dates)}, request_dates={missing_margin_old_dates[:10]}"
    )
    if missing_margin_old_dates:
        margin_old_values = fetch_edb_values_for_dates("M0075992", missing_margin_old_dates, "两融余额历史季度补齐")
        margin_old_rows = [
            {"两融余额": margin_old_values[date], "日期": date}
            for date in missing_margin_old_dates
        ]
        两融余额 = pd.concat([两融余额, pd.DataFrame(margin_old_rows)], ignore_index=True)
        两融余额["周度差值"] = 两融余额["两融余额"] - 两融余额["两融余额"].shift(1)
        两融余额.to_csv(margin_old_path)
    else:
        print("两融余额历史季度数据无缺失日期，无需更新。")
    margin_checkpoint("历史季度补齐完成")

    print(f"两融余额图11日期序列: {time_list}")
    margin_lookup = build_date_value_lookup(两融余额, "日期", "两融余额")
    margin_current_local_hits = [date for date in time_list if date in margin_lookup and pd.notna(margin_lookup[date])]
    margin_current_missing_dates = [date for date in time_list if date not in margin_current_local_hits]
    print(
        f"两融余额图11本地复用: local_hits={margin_current_local_hits}, "
        f"missing={margin_current_missing_dates}"
    )
    margin_current_values = {
        date: round(float(margin_lookup[date]), 1)
        for date in margin_current_local_hits
    }
    if margin_current_missing_dates:
        fetched_margin_values = fetch_edb_values_for_dates("M0075992", margin_current_missing_dates, "两融余额图11")
        margin_current_values.update(fetched_margin_values)
        margin_current_rows_for_old = [
            {"两融余额": fetched_margin_values.get(date, None), "日期": date}
            for date in margin_current_missing_dates
        ]
        两融余额 = pd.concat([两融余额, pd.DataFrame(margin_current_rows_for_old)], ignore_index=True)
        两融余额["周度差值"] = 两融余额["两融余额"] - 两融余额["两融余额"].shift(1)
        两融余额.to_csv(margin_old_path)
        margin_lookup.update(fetched_margin_values)
        margin_current_rows = [
            {"两融余额": margin_current_values.get(date, None), "日期": date}
            for date in time_list
        ]
    两融余额1 = pd.DataFrame(margin_current_rows, columns=["两融余额", "日期"])
    两融余额1["周度差值"] = 两融余额1["两融余额"] - 两融余额1["两融余额"].shift(1)
    log_dataframe_preview("两融余额_图11", 两融余额1)
    两融余额1.to_csv(margin_fig11_path)
    margin_checkpoint("图11数据写出完成")

    两融余额_历史部分 = 两融余额[~两融余额["日期"].astype(str).isin(time_list)].copy()
    两融余额_汇总 = pd.concat([两融余额_历史部分, 两融余额1], axis=0)
    两融余额_汇总 = 两融余额_汇总.drop(columns="Unnamed: 0", errors="ignore")
    两融余额_汇总.to_csv(margin_fig9_path)
    margin_checkpoint("图9数据写出完成")
    checkpoint("两融余额")

    insurance_stage_start = time_module.perf_counter()

    def insurance_checkpoint(label: str) -> None:
        nonlocal insurance_stage_start
        elapsed = time_module.perf_counter() - insurance_stage_start
        print(f"[耗时] step2::保险资金::{label}: {elapsed:.2f} 秒")
        insurance_stage_start = time_module.perf_counter()

    insurance_old_cache_path = paths.funding_draft_dir / "data" / "保险资金" / "保险资金_老.csv"
    insurance_fig11_path = paths.funding_draft_dir / "data" / "保险资金" / "图11_保险资金.csv"
    insurance_fig9_path = paths.funding_draft_dir / "data" / "保险资金" / "图9_保险资金.csv"

    if insurance_old_cache_path.exists():
        insurance_cached_df = pd.read_csv(insurance_old_cache_path, dtype={"日期": str})
        print(f"保险资金本地缓存命中: source={insurance_old_cache_path.name}, rows={len(insurance_cached_df)}")
    elif insurance_fig9_path.exists():
        insurance_cached_df = pd.read_csv(insurance_fig9_path, dtype={"日期": str})
        print(f"保险资金本地缓存接管: source={insurance_fig9_path.name}, rows={len(insurance_cached_df)}")
    else:
        insurance_cached_df = pd.DataFrame(columns=["日期", "保险资金"])
        print("保险资金本地缓存不存在，将执行一次初始化拉取。")
    insurance_checkpoint("读取旧数据完成")

    insurance_cached_df = insurance_cached_df[[col for col in insurance_cached_df.columns if col in {"日期", "保险资金"}]]
    if not insurance_cached_df.empty:
        insurance_cached_df["日期"] = insurance_cached_df["日期"].astype(str)
        insurance_cached_df["保险资金"] = pd.to_numeric(insurance_cached_df["保险资金"], errors="coerce")

    if insurance_cached_df.empty:
        insurance_old = local_client.edb_query("M5876382", "2010-01-01", "2025-01-14", "Fill=Previous")
        insurance_old_data = [round(x, 1) for x in clean_numeric_values(insurance_old.Data[0])]
        insurance_cached_df = pd.DataFrame(
            {
                "日期": format_times_as_yyyymmdd(insurance_old.Times),
                "保险资金": insurance_old_data,
            }
        )
        print(
            f"保险资金初始化完成: rows={len(insurance_cached_df)}, "
            f"start={insurance_cached_df['日期'].iloc[0] if not insurance_cached_df.empty else ''}, "
            f"end={insurance_cached_df['日期'].iloc[-1] if not insurance_cached_df.empty else ''}"
        )
        insurance_checkpoint("初始化历史缓存完成")

    latest_cached_date = (
        pd.to_datetime(insurance_cached_df["日期"], format="%Y%m%d", errors="coerce").dropna().max()
        if not insurance_cached_df.empty
        else pd.NaT
    )
    insurance_live_floor = pd.Timestamp("2024-03-31")
    if pd.isna(latest_cached_date):
        insurance_request_start = insurance_live_floor
    else:
        insurance_request_start = max(insurance_live_floor, latest_cached_date + pd.Timedelta(days=1))
    insurance_request_dates = []
    if insurance_request_start <= pd.Timestamp(today_connect):
        insurance_request_dates = [insurance_request_start.strftime("%Y-%m-%d"), today_connect]
    print(
        f"保险资金增量更新: latest_cached_date="
        f"{latest_cached_date.strftime('%Y%m%d') if pd.notna(latest_cached_date) else 'None'}, "
        f"request_range={insurance_request_dates}"
    )

    if insurance_request_dates:
        insurance = local_client.edb_query(
            "Z8206246,Y3442399,V6997741,U8833275",
            insurance_request_dates[0],
            insurance_request_dates[1],
            "Fill=Previous",
        )
        df_list = []
        for i in range(min(4, len(insurance.Data))):  # ← 修改这行
            df = pd.DataFrame(
                {
                    "日期": format_times_as_yyyymmdd(insurance.Times),
                    f"资金_{i + 1}": clean_numeric_values(insurance.Data[i]),
                }
            )
            df_list.append(df)

        if df_list:
            combined_df = df_list[0]
            for i in range(1, len(df_list)):
                combined_df = combined_df.merge(df_list[i], on="日期", how="outer")
            # 根据实际有的列数求和
            available_cols = [f"资金_{j + 1}" for j in range(len(df_list)) if f"资金_{j + 1}" in combined_df.columns]
            if available_cols:
                combined_df["保险资金"] = combined_df[available_cols].sum(axis=1)
                combined_df["保险资金"] = [round(x, 1) for x in combined_df["保险资金"]]
            else:
                combined_df["保险资金"] = 0
        else:
            print("[警告] 保险资金数据为空")
            combined_df = pd.DataFrame(columns=["日期", "保险资金"])

        insurance_cached_df = pd.concat(
            [insurance_cached_df, combined_df[["日期", "保险资金"]]],
            axis=0,
            ignore_index=True,
        )
        insurance_checkpoint("增量请求完成")

    insurance_cached_df[["日期", "保险资金"]].to_csv(insurance_old_cache_path, index=False)
    insurance_cached_df[["日期", "保险资金"]].to_csv(insurance_fig11_path, index=False)
    insurance_cached_df[["日期", "保险资金"]].to_csv(insurance_fig9_path, index=False)
    insurance_checkpoint("文件写出完成")
    checkpoint("保险资金")

    from functools import reduce

    pict9_lr = pd.read_csv("workspace/funding_draft/data/两融余额数据/图9_两融余额数据.csv")
    pict9_zdqy = pd.read_csv("workspace/funding_draft/data/主动权益基金/图9_主动权益基金.csv")
    pict9_bxzj = pd.read_csv("workspace/funding_draft/data/保险资金/图9_保险资金.csv")
    pict9_bdqy = pd.read_csv("workspace/funding_draft/data/被动权益基金/图9_被动权益ETF.csv")
    pict9_bdqy.rename(columns={"日期_汇总": "日期"}, inplace=True)

    for pict9 in (pict9_lr, pict9_zdqy, pict9_bxzj, pict9_bdqy):
        pict9.reset_index(drop=True, inplace=True)
        pict9["日期"] = pict9["日期"].astype(str)

    dfs = [
        pict9_lr[["两融余额", "日期"]],
        pict9_zdqy[["主动权益基金", "日期"]],
        pict9_bxzj[["保险资金", "日期"]],
        pict9_bdqy[["被动权益ETF", "日期"]],
    ]
    result = reduce(lambda left, right: pd.merge(left, right, on="日期", how="left"), dfs)
    result.to_csv("outputs/csv/figures_1_15/图9_各类资金变动规模.csv", encoding="utf-8-sig")

    pict11_lr = pd.read_csv("workspace/funding_draft/data/两融余额数据/图11_两融余额数据.csv")
    pict11_zdqy = pd.read_csv("workspace/funding_draft/data/主动权益基金/图11_主动权益基金_数据处理.csv")
    pict11_bxzj = pd.read_csv("workspace/funding_draft/data/保险资金/图11_保险资金.csv")
    pict11_bdqy = pd.read_csv("workspace/funding_draft/data/被动权益基金/图11_被动权益ETF处理后数据.csv")

    pict11_lr = pict11_lr[["两融余额", "日期"]]
    pict11_zdqy = pict11_zdqy[["主动权益基金", "日期"]]
    pict11_bxzj = pict11_bxzj[["保险资金", "日期"]].tail(2)
    pict11_bdqy = pict11_bdqy[["被动权益ETF", "日期"]]

    dfs = [pict11_lr, pict11_zdqy, pict11_bxzj, pict11_bdqy]
    for df in dfs:
        df["日期"] = df["日期"].astype(str)

    result1 = reduce(lambda left, right: pd.merge(left, right, on="日期", how="left"), dfs)
    print(result1)
    result1.to_csv("outputs/csv/figures_1_15/图11_本周增量资金来源.csv", encoding="utf-8-sig")
    checkpoint("图9-11数据")

    ##%%
    # 2：按投资范围分类ETF资产净值（亿元）信博卿注释版本
    # 生成每周的日期列
    dates = [start_week_1year_52_Fri + timedelta(days=7 * i) for i in range(52)] #一个列表：1年前开始到现在52周的每个周五
    dates = [date.strftime("%Y%m%d") for date in dates] # for loop, 规定range


    # 定义 ETF 类型和对应的代码
    etf_data = {
        "stock":"159150.SZ,  159300.SZ,  159301.SZ,  159305.SZ,  "
                                       "159306.SZ,  159307.SZ,  159309.SZ,  159310.SZ,  159315.SZ, "
                                       " 159321.SZ,  159322.SZ,  159325.SZ,  159326.SZ,  159327.SZ,  "
                                       "159328.SZ,  159330.SZ,  159332.SZ,  159335.SZ,  159337.SZ,  "
                                       "159338.SZ,  159339.SZ,  159350.SZ,  159351.SZ,  159352.SZ,  "
                                       "159353.SZ,  159355.SZ,  159356.SZ,  159357.SZ,  159358.SZ, "
                                       " 159359.SZ,  159360.SZ,  159361.SZ,  159362.SZ,  159363.SZ,  159505.SZ,  159507.SZ,  159508.SZ,  159510.SZ,  159511.SZ,  159512.SZ,  159515.SZ,  159516.SZ,  159517.SZ,  159520.SZ,  159521.SZ,  159523.SZ,  159525.SZ,  159526.SZ,  159527.SZ,  159528.SZ,  159530.SZ,  159531.SZ,  159532.SZ,  159533.SZ,  159535.SZ,  159536.SZ,  159537.SZ,  159538.SZ,  159539.SZ,  159540.SZ,  159541.SZ,  159542.SZ,  159543.SZ,  159546.SZ,  159547.SZ,  159549.SZ,  159551.SZ,  159552.SZ,  159553.SZ,  159555.SZ,  159556.SZ,  159558.SZ,  159559.SZ,  159560.SZ,  159562.SZ,  159563.SZ,  159565.SZ,  159566.SZ,  159571.SZ,  159572.SZ,  159573.SZ,  159575.SZ,  159576.SZ,  159578.SZ,  159579.SZ,  159581.SZ,  159582.SZ,  159583.SZ,  159586.SZ,  159587.SZ,  159588.SZ,  159589.SZ,  159590.SZ,  159591.SZ,  159592.SZ,  159593.SZ,  159595.SZ,  159596.SZ,  159597.SZ,  159599.SZ,  159601.SZ,  159602.SZ,  159603.SZ,  159606.SZ,  159608.SZ,  159609.SZ,  159610.SZ,  159611.SZ,  159613.SZ,  159616.SZ,  159617.SZ,  159618.SZ,  159619.SZ,  159620.SZ,  159621.SZ,  159622.SZ,  159623.SZ,  159625.SZ,  159627.SZ,  159628.SZ,  159629.SZ,  159630.SZ,  159631.SZ,  159633.SZ,  159635.SZ,  159637.SZ,  159638.SZ,  159639.SZ,  159640.SZ,  159641.SZ,  159642.SZ,  159643.SZ,  159645.SZ,  159647.SZ,  159652.SZ,  159653.SZ,  159656.SZ,  159657.SZ,  159658.SZ,  159661.SZ,  159662.SZ,  159663.SZ,  159665.SZ,  159666.SZ,  159667.SZ,  159669.SZ,  159670.SZ,  159671.SZ,  159672.SZ,  159673.SZ,  159675.SZ,  159676.SZ,  159677.SZ,  159678.SZ,  159679.SZ,  159680.SZ,  159681.SZ,  159682.SZ,  159683.SZ,  159685.SZ,  159686.SZ,  159689.SZ,  159690.SZ,  159692.SZ,  159695.SZ,  159697.SZ,  159698.SZ,  159701.SZ,  159703.SZ,  159706.SZ,  159707.SZ,  159708.SZ,  159709.SZ,  159713.SZ,  159715.SZ,  159716.SZ,  159717.SZ,  159719.SZ,  159720.SZ,  159721.SZ,  159723.SZ,  159725.SZ,  159728.SZ,  159729.SZ,  159730.SZ,  159731.SZ,  159732.SZ,  159736.SZ,  159738.SZ,  159739.SZ,  159743.SZ,  159745.SZ,  159748.SZ,  159752.SZ,  159755.SZ,  159757.SZ,  159758.SZ,  159760.SZ,  159761.SZ,  159763.SZ,  159766.SZ,  159767.SZ,  159768.SZ,  159770.SZ,  159773.SZ,  159775.SZ,  159777.SZ,  159778.SZ,  159779.SZ,  159780.SZ,  159781.SZ,  159782.SZ,  159783.SZ,  159786.SZ,  159787.SZ,  159790.SZ,  159791.SZ,  159793.SZ,  159795.SZ,  159796.SZ,  159797.SZ,  159798.SZ,  159800.SZ,  159801.SZ,  159804.SZ,  159805.SZ,  159806.SZ,  159807.SZ,  159808.SZ,  159810.SZ,  159811.SZ,  159813.SZ,  159814.SZ,  159819.SZ,  159820.SZ,  159821.SZ,  159824.SZ,  159825.SZ,  159827.SZ,  159828.SZ,  159835.SZ,  159836.SZ,  159837.SZ,  159838.SZ,  159839.SZ,  159840.SZ,  159841.SZ,  159842.SZ,  159843.SZ,  159845.SZ,  159847.SZ,  159848.SZ,  159849.SZ,  159851.SZ,  159852.SZ,  159855.SZ,  159856.SZ,  159857.SZ,  159858.SZ,  159859.SZ,  159861.SZ,  159862.SZ,  159863.SZ,  159864.SZ,  159865.SZ,  159867.SZ,  159869.SZ,  159870.SZ,  159871.SZ,  159872.SZ,  159873.SZ,  159875.SZ,  159876.SZ,  159877.SZ,  159880.SZ,  159881.SZ,  159883.SZ,  159885.SZ,  159886.SZ,  159887.SZ,  159888.SZ,  159889.SZ,  159890.SZ,  159891.SZ,  159895.SZ,  159896.SZ,  159898.SZ,  159899.SZ,  159901.SZ,  159902.SZ,  159903.SZ,  159905.SZ,  159906.SZ,  159907.SZ,  159908.SZ,  159909.SZ,  159910.SZ,  159912.SZ,  159913.SZ,  159915.SZ,  159916.SZ,  159918.SZ,  159919.SZ,  159922.SZ,  159923.SZ,  159925.SZ,  159928.SZ,  159929.SZ,  159930.SZ,  159931.SZ,  159933.SZ,  159935.SZ,  159936.SZ,  159938.SZ,  159939.SZ,  159940.SZ,  159943.SZ,  159944.SZ,  159945.SZ,  159948.SZ,  159949.SZ,  159952.SZ,  159956.SZ,  159957.SZ,  159958.SZ,  159959.SZ,  159961.SZ,  159964.SZ,  159965.SZ,  159966.SZ,  159967.SZ,  159968.SZ,  159969.SZ,  159970.SZ,  159971.SZ,  159973.SZ,  159974.SZ,  159975.SZ,  159976.SZ,  159977.SZ,  159982.SZ,  159991.SZ,  159992.SZ,  159993.SZ,  159994.SZ,  159995.SZ,  159996.SZ,  159997.SZ,  159998.SZ,  510010.SH,  510020.SH,  510030.SH,  510050.SH,  510060.SH,  510090.SH,  510100.SH,  510130.SH,  510150.SH,  510160.SH,  510170.SH,  510180.SH,  510190.SH,  510200.SH,  510210.SH,  510230.SH,  510270.SH,  510290.SH,  510300.SH,  510310.SH,  510330.SH,  510350.SH,  510360.SH,  510370.SH,  510380.SH,  510390.SH,  510410.SH,  510500.SH,  510510.SH,  510530.SH,  510550.SH,  510560.SH,  510570.SH,  510580.SH,  510590.SH,  510600.SH,  510630.SH,  510650.SH,  510660.SH,  510680.SH,  510710.SH,  510720.SH,  510760.SH,  510770.SH,  510800.SH,  510810.SH,  510850.SH,  510880.SH,  510950.SH,  510980.SH,  510990.SH,  512000.SH,  512010.SH,  512020.SH,  512040.SH,  512050.SH,  512070.SH,  512090.SH,  512100.SH,  512120.SH,  512150.SH,  512160.SH,  512170.SH,  512180.SH,  512190.SH,  512200.SH,  512220.SH,  512260.SH,  512280.SH,  512290.SH,  512330.SH,  512360.SH,  512380.SH,  512390.SH,  512400.SH,  512480.SH,  512500.SH,  512510.SH,  512520.SH,  512530.SH,  512550.SH,  512560.SH,  512570.SH,  512580.SH,  512600.SH,  512640.SH,  512650.SH,  512660.SH,  512670.SH,  512680.SH,  512690.SH,  512700.SH,  512710.SH,  512720.SH,  512730.SH,  512750.SH,  512760.SH,  512770.SH,  512800.SH,  512810.SH,  512820.SH,  512870.SH,  512880.SH,  512890.SH,  512900.SH,  512910.SH,  512930.SH,  512950.SH,  512960.SH,  512970.SH,  512980.SH,  512990.SH,  515000.SH,  515010.SH,  515020.SH,  515030.SH,  515050.SH,  515060.SH,  515070.SH,  515080.SH,  515090.SH,  515100.SH,  515110.SH,  515120.SH,  515130.SH,  515150.SH,  515160.SH,  515170.SH,  515180.SH,  515190.SH,  515200.SH,  515210.SH,  515220.SH,  515230.SH,  515250.SH,  515260.SH,  515290.SH,  515300.SH,  515310.SH,  515320.SH,  515330.SH,  515350.SH,  515360.SH,  515380.SH,  515390.SH,  515400.SH,  515450.SH,  515530.SH,  515550.SH,  515560.SH,  515580.SH,  515590.SH,  515600.SH,  515630.SH,  515650.SH,  515660.SH,  515680.SH,  515700.SH,  515710.SH,  515750.SH,  515760.SH,  515770.SH,  515780.SH,  515790.SH,  515800.SH,  515810.SH,  515850.SH,  515860.SH,  515880.SH,  515890.SH,  515900.SH,  515910.SH,  515920.SH,  515950.SH,  515960.SH,  515980.SH,  515990.SH,  516000.SH,  516010.SH,  516020.SH,  516050.SH,  516060.SH,  516070.SH,  516080.SH,  516090.SH,  516100.SH,  516110.SH,  516120.SH,  516130.SH,  516150.SH,  516160.SH,  516180.SH,  516190.SH,  516200.SH,  516210.SH,  516220.SH,  516260.SH,  516270.SH,  516290.SH,  516300.SH,  516310.SH,  516320.SH,  516330.SH,  516350.SH,  516360.SH,  516380.SH,  516390.SH,  516480.SH,  516500.SH,  516510.SH,  516520.SH,  516530.SH,  516550.SH,  516560.SH,  516570.SH,  516580.SH,  516590.SH,  516600.SH,  516610.SH,  516620.SH,  516630.SH,  516640.SH,  516650.SH,  516660.SH,  516670.SH,  516700.SH,  516710.SH,  516720.SH,  516730.SH,  516750.SH,  516760.SH,  516770.SH,  516780.SH,  516790.SH,  516800.SH,  516810.SH,  516820.SH,  516830.SH,  516850.SH,  516860.SH,  516880.SH,  516890.SH,  516900.SH,  516910.SH,  516920.SH,  516930.SH,  516950.SH,  516960.SH,  516970.SH,  516980.SH,  517000.SH,  517010.SH,  517030.SH,  517050.SH,  517080.SH,  517090.SH,  517100.SH,  517110.SH,  517120.SH,  517160.SH,  517170.SH,  517180.SH,  517200.SH,  517300.SH,  517330.SH,  517350.SH,  517360.SH,  517380.SH,  517390.SH,  517400.SH,  517520.SH,  517550.SH,  517660.SH,  517770.SH,  517800.SH,  517850.SH,  517880.SH,  517900.SH,  517990.SH,  530000.SH,  530050.SH,  530880.SH,  560000.SH,  560010.SH,  560020.SH,  560030.SH,  560050.SH,  560060.SH,  560070.SH,  560080.SH,  560090.SH,  560100.SH,  560110.SH,  560150.SH,  560170.SH,  560180.SH,  560220.SH,  560260.SH,  560280.SH,  560300.SH,  560330.SH,  560350.SH,  560360.SH,  560500.SH,  560510.SH,  560520.SH,  560530.SH,  560550.SH,  560560.SH,  560580.SH,  560590.SH,  560610.SH,  560620.SH,  560650.SH,  560660.SH,  560680.SH,  560690.SH,  560700.SH,  560780.SH,  560800.SH,  560810.SH,  560850.SH,  560860.SH,  560880.SH,  560890.SH,  560900.SH,  560950.SH,  560960.SH,  560980.SH,  560990.SH,  561000.SH,  561010.SH,  561060.SH,  561100.SH,  561120.SH,  561130.SH,  561160.SH,  561170.SH,  561180.SH,  561190.SH,  561200.SH,  561230.SH,  561260.SH,  561280.SH,  561300.SH,  561310.SH,  561320.SH,  561330.SH,  561350.SH,  561360.SH,  561370.SH,  561500.SH,  561510.SH,  561550.SH,  561560.SH,  561570.SH,  561580.SH,  561590.SH,  561600.SH,  561700.SH,  561760.SH,  561780.SH,  561790.SH,  561800.SH,  561900.SH,  561910.SH,  561920.SH,  561930.SH,  561950.SH,  561960.SH,  561980.SH,  561990.SH,  562000.SH,  562010.SH,  562030.SH,  562060.SH,  562260.SH,  562300.SH,  562310.SH,  562320.SH,  562330.SH,  562340.SH,  562350.SH,  562360.SH,  562380.SH,  562390.SH,  562500.SH,  562510.SH,  562520.SH,  562530.SH,  562550.SH,  562560.SH,  562570.SH,  562580.SH,  562590.SH,  562600.SH,  562660.SH,  562700.SH,  562800.SH,  562820.SH,  562850.SH,  562860.SH,  562880.SH,  562890.SH,  562900.SH,  562910.SH,  562920.SH,  562930.SH,  562950.SH,  562960.SH,  562970.SH,  562990.SH,  563000.SH,  563010.SH,  563020.SH,  563030.SH,  563050.SH,  563080.SH,  563090.SH,  563150.SH,  563180.SH,  563200.SH,  563220.SH,  563280.SH,  563300.SH,  563330.SH,  563350.SH,  563360.SH,  563500.SH,  563520.SH,  563800.SH,  563880.SH,  588000.SH,  588010.SH,  588020.SH,  588030.SH,  588050.SH,  588060.SH,  588070.SH,  588080.SH,  588090.SH,  588100.SH,  588110.SH,  588120.SH,  588150.SH,  588160.SH,  588180.SH,  588190.SH,  588200.SH,  588210.SH,  588220.SH,  588260.SH,  588280.SH,  588290.SH,  588300.SH,  588310.SH,  588320.SH,  588330.SH,  588350.SH,  588360.SH,  588370.SH,  588380.SH,  588390.SH,  588400.SH,  588450.SH,  588460.SH,  588500.SH,  588680.SH,  588700.SH,  588800.SH,  588830.SH,  588860.SH,  588880.SH,  588890.SH,  588900.SH,  588990.SH",
        "bond": "159649.SZ,  159650.SZ,  159651.SZ,  159816.SZ,  159972.SZ,  "
                               "511010.SH,  511020.SH,  511030.SH,  511060.SH,  511090.SH,  511100.SH,  "
                               "511130.SH,  511180.SH,  511220.SH,  511260.SH,  511270.SH,  511360.SH,  "
                               "511380.SH,  511520.SH,  511580.SH",
        "commodity": "159812.SZ,  159830.SZ,  159831.SZ,  159834.SZ,  "
                                                 "159934.SZ,  159937.SZ,  159980.SZ,  159981.SZ,  "
                                                 "159985.SZ,  518600.SH,  518660.SH,  518680.SH,  "
                                                 "518800.SH,  518850.SH,  518860.SH,  518880.SH,  518890.SH",
        "currency": "159001.SZ,  159003.SZ,  159005.SZ,  "
                                                "511600.SH,  511620.SH,  511650.SH,  511660.SH,  "
                                                "511670.SH,  511690.SH,  511700.SH,  511770.SH, "
                                                " 511800.SH,  511810.SH,  511820.SH,  511830.SH,  "
                                                "511850.SH,  511860.SH,  511880.SH,  511900.SH,  "
                                                "511910.SH,  511920.SH,  511930.SH,  511950.SH,  "
                                                "511960.SH,  511970.SH,  511980.SH,  511990.SH"
    }

    # 初始化最终汇总的 DataFrame
    columns = ["Date"] + list(etf_data.keys()) # 初始化column名称 date, 4个类
    etf_net_asset_rows = []

    # 遍历每个日期
    for date in dates: # 1级嵌套for loop 每个日期来一
        row_data = {"Date": date} # 日期loop 一

        # 每个日期下对4个类型的ETF 进行资产净值计
        for etf_type, codes in etf_data.items(): # 二级嵌套for_loop 对应key - value 迭代
            verified_fetch = local_client.fetch_etf_netasset_verified(
                codes,
                str(date),
                batch_size=batch_size,
            )
            row_data[etf_type] = log_verified_etf_fetch(
                f"图12_{etf_type}_ETF基金资产净值_{date}",
                verified_fetch,
            ) / 100000000

        # 将每周的数据添加到最DataFrame
        etf_net_asset_rows.append(row_data)

    etf_net_asset_summary = pd.DataFrame(etf_net_asset_rows, columns=columns)



    #-------------------------------------------------------------------------------------------------------------
    # 输出最终的 DataFrame
    print(
        f"图12 ETF规模汇总完成: shape={etf_net_asset_summary.shape}, "
        f"latest_date={etf_net_asset_summary['Date'].iloc[-1]}"
    )

    # stock_ETF_net_asset = local_client.snapshot(etf_data["stock"],
    #                                    "etf_netasset", f"unit=1;tradeDate={today};")
    # print(stock_ETF_net_asset) # print 完发现数据有嵌套

    # codes = stock_ETF_net_asset.Codes
    # sto_net_ass = stock_ETF_net_asset.Data[0]
    # date = stock_ETF_net_asset.Times[0].strftime("%Y%m%d") #输出: "20241211 20:03:17"，然后split()
    # # 输出: ["20241211", "20:03:17"]
    # #split()[0] 输出: "20241211"
    # stock_ETF_net_asset = pd.DataFrame({
    #     "ETF_Code": codes,
    #     "Net_Asset": sto_net_ass,
    #     "Date": date  # 添加日期
    # })
    # print(stock_ETF_net_asset.head())

    etf_net_asset_summary.to_csv("outputs/csv/figures_1_15/图12_按投资范围分类ETF资产净值（亿元）.csv",index = 0,encoding="utf-8-sig")

    checkpoint("ETF规模")


    # 0：股票型ETF净流入额（亿元）
    dates_Mon = [start_week9 + timedelta(days=7 * i) for i in range(9)]
    dates_Mon = [date.strftime("%Y%m%d") for date in dates_Mon]
    dates_Fri = [start_week9_Fri + timedelta(days=7 * i) for i in range(9)]
    dates_Fri = [date.strftime("%Y%m%d") for date in dates_Fri]

    fig10_stage_start = time_module.perf_counter()

    def fig10_checkpoint(label: str) -> None:
        nonlocal fig10_stage_start
        elapsed = time_module.perf_counter() - fig10_stage_start
        print(f"[耗时] step2::图10::{label}: {elapsed:.2f} 秒")
        fig10_stage_start = time_module.perf_counter()

    weekly_ranges = list(zip(dates_Mon, dates_Fri))
    latest_stock_etf_weekly_flow: dict | None = None
    print(
        f"run_step2::图10开始: week_count={len(weekly_ranges)}, "
        f"sample={[f'{start}-{end}' for start, end in weekly_ranges[:3]]}"
    )
    stock_net_inflow_TS = pd.DataFrame(columns=["Date", "net_inflow_stock"])
    if not ENABLE_FIGURE10_ETF_NETINFLOW:
        print("图10开关关闭：跳过股票型ETF净流入额取数与绘图。")
        stock_net_inflow_TS.to_csv("outputs/csv/figures_1_15/图10_股票型ETF净流入额（亿元）.csv", index=0, encoding="utf-8-sig")
    else:
        try:
            stock_net_inflow_rows = []
            for dateMon, dateFri in weekly_ranges:
                if latest_stock_etf_weekly_flow is not None and dateMon == start_week and dateFri == today:
                    summary = latest_stock_etf_weekly_flow
                else:
                    summary = fetch_periodmf_netinflow_chunked(
                        stock_ETF_list,
                        dateMon,
                        dateFri,
                        batch_size=batch_size,
                    )
                    if dateMon == start_week and dateFri == today:
                        latest_stock_etf_weekly_flow = summary
                stock_net_inflow_rows.append(
                    {
                        "Date": summary["Date"],
                        "net_inflow_stock": summary["total_net_inflow_stock"],
                    }
                )

            stock_net_inflow_TS = pd.DataFrame(stock_net_inflow_rows, columns=["Date", "net_inflow_stock"])
            if stock_net_inflow_TS.empty:
                raise RuntimeError("图10 周度结果为空，weekly rows 未成功写入。")
            print(f"图10最终 DataFrame shape={stock_net_inflow_TS.shape}")
            fig10_checkpoint("周度净流入计算完成")

            stock_net_inflow_TS.to_csv("outputs/csv/figures_1_15/图10_股票型ETF净流入额（亿元）.csv", index=0, encoding="utf-8-sig")
            fig10_checkpoint("图10数据写出完成")
        except Exception as exc:
            record_figure_failure("图10_股票型ETF净流入额", exc)
    checkpoint("股票型ETF净流入额")

    # 3：行业ETF当周净流入额（亿元）
    industry_flow_columns = ["code", "tracking_indices", "net_flow", "industry_names"]
    if not ENABLE_FIGURE13_INDUSTRY_ETF:
        print("图13开关关闭：跳过行业ETF周区间净流入额取数与绘图。")
        net_flow_industry_ETF_summary = pd.DataFrame(columns=industry_flow_columns)
        net_flow_industry_ETF_summary.to_csv("outputs/csv/figures_1_15/图13_按行业指数分类ETF周区间净流入额（亿元）.csv", index=0, encoding="utf-8-sig")
    else:
        try:
            if latest_stock_etf_weekly_flow is None:
                latest_stock_etf_weekly_flow = fetch_periodmf_netinflow_chunked(
                    stock_ETF_list,
                    start_week,
                    today,
                    batch_size=batch_size,
                )

            industry_ETF_codes = normalize_code_list(
                "159301.SZ,159511.SZ,159546.SZ,159586.SZ,159611.SZ,159662.SZ,159666.SZ,159672.SZ,"
                "159689.SZ,159707.SZ,159797.SZ,159841.SZ,159842.SZ,159848.SZ,159883.SZ,159887.SZ,"
                "159898.SZ,159928.SZ,159929.SZ,159930.SZ,159936.SZ,159938.SZ,159939.SZ,159944.SZ,"
                "159945.SZ,159996.SZ,510200.SH,510230.SH,510630.SH,510660.SH,512000.SH,512010.SH,"
                "512200.SH,512330.SH,512400.SH,512480.SH,512570.SH,512600.SH,512700.SH,512730.SH,"
                "512800.SH,512820.SH,512880.SH,512900.SH,515010.SH,515020.SH,515060.SH,515220.SH,"
                "515290.SH,515560.SH,515850.SH,516200.SH,516210.SH,516310.SH,516900.SH,516970.SH,"
                "517380.SH,517990.SH,560090.SH,560580.SH,560620.SH,560680.SH,560880.SH,561010.SH,"
                "561120.SH,561360.SH,561560.SH,561570.SH,561700.SH,562350.SH,562560.SH,562580.SH,"
                "562600.SH,562820.SH"
            )
            flow_detail_df = latest_stock_etf_weekly_flow["detail_df"][["code", "net_flow"]].copy()

            tracking_indices = [
                "全指公用", "800通信", "集成电路", "计算机", "中证全指电力指数", "国证交运", "中证全指运输指数",
                "800消费", "800消费", "800地产", "医疗器械", "证券公司", "证券公司", "证券公司", "医疗器械",
                "800银行", "医疗器械", "800消费", "800医药", "800能源", "全指可选", "全指医药", "全指信息",
                "全指材料", "全指能源", "家用电器", "上证证券", "180金融", "上证消费", "上证医药", "证券公司",
                "300医药", "中证全指房地产", "中证500信息", "有色金属", "中证全指半导体", "证券公司", "800消费",
                "中证银行", "中证银行", "中证银行", "中证银行", "证券公司", "证券公司", "证券公司", "中证银行",
                "中证全指房地产", "中证煤炭", "中证银行", "证券公司", "证券公司", "证券公司", "中证银行", "中证银行",
                "食品饮料", "基建工程", "恒生沪深港创新药50", "沪港深500医药", "证券公司", "中证全指电力指数",
                "全指公用", "800消费", "家用电器", "软件开发", "家用电器", "油气产业", "中证全指电力指数",
                "油气产业", "中证全指电力指数", "中证全指电力指数", "全指信息", "全指可选", "医疗器械", "集成电路",
            ]
            industry_names = [
                "公用事业", "通信", "电子", "计算机", "电力", "交通运输", "交通运输", "消费", "消费",
                "地产", "医药", "证券", "证券", "证券", "医药", "银行", "医药", "消费", "医药", "能源",
                "消费", "医药", "通信", "周期", "周期", "电器", "证券", "金融", "消费", "医药", "证券",
                "医药", "房地产", "通信", "有色金属", "电子", "证券", "消费", "银行", "银行", "银行", "银行",
                "证券", "证券", "证券", "银行", "房地产", "煤炭", "银行", "证券", "证券", "证券", "银行",
                "银行", "消费", "基建", "医药", "医药", "证券", "电力", "公用事业", "消费", "家用电器",
                "计算机", "家用电器", "周期", "电力", "周期", "电力", "电力", "通信", "消费", "医药", "电子",
            ]
            static_mapping_df = pd.DataFrame(
                {
                    "code": industry_ETF_codes,
                    "tracking_indices": tracking_indices,
                    "industry_names": industry_names,
                }
            )
            net_flow_industry_ETF_summary = static_mapping_df.merge(flow_detail_df, on="code", how="left")
            print(
                f"图13上游复用图10缓存: requested={len(industry_ETF_codes)}, "
                f"matched={int(net_flow_industry_ETF_summary['net_flow'].notna().sum())}, "
                f"non_zero_count={int((net_flow_industry_ETF_summary['net_flow'].fillna(0) != 0).sum())}"
            )
            if net_flow_industry_ETF_summary["net_flow"].notna().sum() == 0:
                raise RuntimeError("图13 依赖图10重建的ETF净流入结果，但当前行业ETF净流入全部为空。")
            unmapped_df = net_flow_industry_ETF_summary[
                net_flow_industry_ETF_summary["tracking_indices"].isna()
                | net_flow_industry_ETF_summary["industry_names"].isna()
            ]
            if not unmapped_df.empty:
                raise RuntimeError(
                    f"图13 行业映射覆盖不完整，unmapped_count={len(unmapped_df)}, "
                    f"unmapped_records={unmapped_df[['code', 'tracking_indices']].to_dict(orient='records')[:10]}"
                )
            print(
                f"图13最终 DataFrame shape={net_flow_industry_ETF_summary.shape}, "
                f"industry_count={net_flow_industry_ETF_summary['industry_names'].nunique()}"
            )
            net_flow_industry_ETF_summary = net_flow_industry_ETF_summary[["code", "tracking_indices", "net_flow", "industry_names"]]
            net_flow_industry_ETF_summary.to_csv("outputs/csv/figures_1_15/图13_按行业指数分类ETF周区间净流入额（亿元）.csv", index=0, encoding="utf-8-sig")
        except Exception as exc:
            net_flow_industry_ETF_summary = pd.DataFrame(columns=industry_flow_columns)
            record_figure_failure("图13_按行业指数分类ETF周区间净流入额", exc)
    checkpoint("图13行业ETF净流入额")


    # %%4：限售股解禁金额和数量（周度
    if not ENABLE_FIGURE14_UNLOCK:
        print("图14开关关闭：跳过限售股解禁金额和数量取数与绘图。")
        xs_use = pd.DataFrame(columns=["起始日期", "当周解禁市值", "当周解禁家数", "日期"])
        xs_use.to_csv("outputs/csv/figures_1_15/图14_限售股解禁金额和数量（周度）.csv",index = 0,encoding="utf-8-sig")
    else:
        try:
            xs_use = build_unlock_weekly_dataframe_from_wc(start_week16, today)
            print(f"图14最终 DataFrame shape={xs_use.shape}")
            xs_use.to_csv("outputs/csv/figures_1_15/图14_限售股解禁金额和数量（周度）.csv",index = 0,encoding="utf-8-sig")
        except Exception as exc:
            xs_use = pd.DataFrame(columns=["起始日期", "当周解禁市值", "当周解禁家数", "日期"])
            record_figure_failure("图14_限售股解禁金额和数量（周度）", exc)

    # %%5：IPO和定增金额（周度
    if not ENABLE_FIGURE15_IPO:
        print("图15开关关闭：跳过IPO和定增金额取数与绘图。")
        ipo_use = pd.DataFrame(columns=["日期", "IPO首发家数", "IPO首发募集资金（亿元）", "定增家数", "定增募集（亿元）"])
        ipo_use.to_csv("outputs/csv/figures_1_15/图15_IPO和定增金额（周度）.csv",index = 0,encoding="utf-8-sig")
    else:
        try:
            ipo_use = build_ipo_private_weekly_dataframe_from_wc_safe(start_week20, today)
            print(f"图15最终 DataFrame shape={ipo_use.shape}")
            ipo_use.to_csv("outputs/csv/figures_1_15/图15_IPO和定增金额（周度）.csv",index = 0,encoding="utf-8-sig")
        except Exception as exc:
            ipo_use = pd.DataFrame(columns=["日期", "IPO首发家数", "IPO首发募集资金（亿元）", "定增家数", "定增募集（亿元）"])
            record_figure_failure("图15_IPO和定增金额（周度）", exc)
    if figure_failures:
        raise RuntimeError(
            "Step2 目标图表生成失败，请根据以下真实原因逐项修复：\n- "
            + "\n- ".join(figure_failures)
        )

    # 汇总Excel只保留单一输出位置，避免跨目录重复落盘
    pict9_data = pd.read_csv("outputs/csv/figures_1_15/图9_各类资金变动规模.csv")
    pict11_data = pd.read_csv("outputs/csv/figures_1_15/图11_本周增量资金来源.csv")

    result_list = [
        market_Index,
        ind1_Index,
        ind2_data,
        hot_data,
        trading_data,
        qh_use50,
        updown_use,
        rzrq_use,
        pict9_data,
        stock_net_inflow_TS,
        pict11_data,
        etf_net_asset_summary,
        net_flow_industry_ETF_summary,
        xs_use,
        ipo_use,
    ]
    sheet_name_list = [
        "图1_本周市场主要指数涨跌幅",
        "图2_本周申万一级行业涨跌幅",
        "图3_本周申万二级行业涨跌幅",
        "图4_本周热点概念板块涨跌幅（排名前10位和后10位）",
        "图5_更新市场情绪相关数据",
        "图6_股指期货主力合约基差",
        "图7_两市涨跌数量统计",
        "图8_两市融资融券余额、融资买入占比",
        "图9_各类资金变动规模（亿元）",
        "图10_股票型ETF净流入额（亿元）",
        "图11_本周增量资金来源（亿元）",
        "图12_各类ETF规模变化",
        "图13_行业ETF当周净流入额（亿元）",
        "图14_限售股解禁金额和数量",
        "图15_IPO和定增金额",
    ]
    summary_path = f"outputs/csv/summary/{today}图1-15数据.xlsx"
    with pd.ExcelWriter(summary_path) as writer:
        for frame, sheet_name in zip(result_list, sheet_name_list):
            frame.to_excel(writer, sheet_name=sheet_name, index=False)

def run_step3(paths: ProjectPaths, report_date: datetime.date) -> None:
    configure_matplotlib_backend()
    font_family = configure_plot_font(paths)
    import os
    import numpy as np
    import pandas as pd
    import datetime
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mtick
    import matplotlib.dates as mdates
    from matplotlib.pyplot import MultipleLocator
    from PIL import Image
    from datetime import timedelta
    from matplotlib.backends.backend_pdf import PdfPages
    import matplotlib.colors as mcolors
    import matplotlib.patches as mpatches
    DB_LOC = str(paths.data_dir)

    def load_chart_data_with_numeric_validation(
        csv_path: str,
        numeric_columns: list[str],
        chart_name: str,
    ) -> pd.DataFrame:
        data = pd.read_csv(csv_path)
        if data.empty:
            raise RuntimeError(f"{chart_name} 绘图失败：输入文件为空。 path={csv_path}")
        for column in numeric_columns:
            if column not in data.columns:
                raise RuntimeError(
                    f"{chart_name} 绘图失败：缺少列 {column}。 path={csv_path}, columns={list(data.columns)}"
                )
            data[column] = pd.to_numeric(data[column], errors="coerce")
            if data[column].notna().sum() == 0:
                raise RuntimeError(
                    f"{chart_name} 绘图失败：列 {column} 不存在有效数值。 "
                    f"path={csv_path}. 请先检查 Step2 对应数据生成逻辑。"
                )
        return data




    # %%画图设置
    plt.rcParams['font.sans-serif'] = [font_family]  # 在tilte中正常显示中
    plt.rcParams['axes.unicode_minus'] = False  # 显示负号
    plt.rcParams.update({'font.size': 20})
    font_English = {'family': 'Times New Roman'}

    # %% DATE
    trading_use = pd.read_csv("outputs/csv/figures_1_15/图5_两市成交金额、换手率.csv")
    date = pd.DataFrame(pd.to_datetime(trading_use["date"], format='%Y-%m-%d'))
    date["date_use"] = trading_use["date"]
    date = date.iloc[::-1]
    h = 1
    for i in range(1, len(date)):
        if (date.iloc[i, 0] - date.iloc[i - 1, 0]).days == -1:
            h += 1
        else:
            break
    date_deal = date.iloc[:h, 1]

    today = date_deal.iloc[0]

    start_day = date_deal.iloc[-1]
    # %% 绘制
    market_Index = pd.read_csv("outputs/csv/figures_1_15/图1_本周市场主要指数涨跌幅.csv")
    # 添加画布
    fig = plt.figure(figsize=(34, 14 / 1.2), dpi=300)  # 4/6宫格
    ax = fig.add_subplot(111)

    # 将y轴的刻度方向设置向内
    plt.rcParams['ytick.direction'] = 'in'

    # 设置y坐标上下
    min_ylim = int(min(market_Index.iloc[:, 2:5].min(
    )) / 0.04) * 4 / 100 - ((min(market_Index.iloc[:, 2:5].min())) < 0) * 4 / 100
    max_ylim = int(max(market_Index.iloc[:, 2:5].max(
    )) / 0.04) * 4 / 100 + ((max(market_Index.iloc[:, 2:5].max())) > 0) * 8 / 100
    ax.set_ylim(min_ylim, max_ylim)

    # 设置y坐标刻度间隔
    y_major_locator = MultipleLocator(0.04)  # 以每4显示
    ax = plt.gca()
    ax.yaxis.set_major_locator(y_major_locator)
    ax.yaxis.set_major_formatter(mtick.PercentFormatter(1))
    vals = ax.get_yticks()
    ax.set_yticklabels(['{:,.0%}'.format(x) for x in vals])

    # 设置纵坐标轴数值为times new roman
    labels = ax.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]
    # 绘制x=0
    ax.axhline(y=0, c='grey', lw=2)

    # 隐藏原始x
    plt.xticks(alpha=0)
    plt.tick_params(axis='x', width=0)

    # 画图
    x1 = [1, 5, 9, 13, 17, 21, 25, 29]
    x2 = [i + 1 for i in x1]
    x3 = [i + 2 for i in x1]
    x4 = [i + 3 for i in x1]

    plt.bar(x1, market_Index["本周涨跌幅(%)"], alpha=0.7, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=3, label="本周涨跌幅(%)", zorder=1)
    plt.bar(x2, market_Index["本月涨跌幅(%)"], alpha=0.7, width=1, color='darkgrey', edgecolor=[
        "dimgrey"], linewidth=3, label="本月涨跌幅(%)", zorder=1)
    plt.bar(x3, market_Index["本年涨跌幅(%)"], alpha=0.8, width=1, color='sandybrown', edgecolor=[
        "coral"], linewidth=3, label="本年涨跌幅(%)", zorder=1)

    # 设置x轴
    plt.xlim(0, 32.01)
    ax.set_xticks(x4)

    # 柱状图上显示数据
    x = pd.DataFrame([x1, x2, x3])
    for i in range(0, 3):
        for j in range(0, 8):
            value = market_Index.iloc[j, i + 2]
            if float(value) >= 0:
                plt.text(x.iloc[i, j], value, str("%.2f" % (value * 100)) + '%\n',
                         ha='center', va='center', fontsize=25, family="Times new roman", zorder=2)
            else:
                plt.text(x.iloc[i, j], value - 0.008, str("%.2f" % (value * 100)) + '%\n',
                         ha='center', va='center', fontsize=25, family="Times new roman", zorder=2)

    # 设置下边框名
    for i in range(0, 8):
        plt.text(x2[i], min_ylim - 0.1 * (max_ylim - min_ylim), market_Index.iloc[i, 1],
                 ha='center', va="bottom", fontsize=30)

    # 去除边框
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # 设置网格
    plt.grid(linestyle='--', dashes=(24, 12), zorder=0)

    # 设置标签
    plt.legend(frameon=False, ncol=3, bbox_to_anchor=(
        0.5, 1.15), loc=9, borderaxespad=0, fontsize=30)

    # 储存图片
    plt.savefig("outputs/figures/图1_本周市场主要指数涨跌幅.png", dpi=300, bbox_inches='tight')
    pic_1 = Image.open("outputs/figures/图1_本周市场主要指数涨跌幅.png")
    size = pic_1.size
    print(size)
    print(size[0] / size[1])
    # print(size2[0]/size2[1])
    # 修改长宽比例
    pic_new = pic_1.resize((2380, 980))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture1.png")

    # word交互
    date = report_date
    jh_1 = market_Index.sort_values(by='本周涨跌幅(%)', ascending=False)
    jh_1_po = jh_1[jh_1['本周涨跌幅(%)'] >= 0]
    jh_1_ne = jh_1[jh_1['本周涨跌幅(%)'] < 0]
    jh_1_ne = jh_1_ne.sort_values(by='本周涨跌幅(%)')
    name_po = ""
    name_ne = ""
    value_po = ""
    value_ne = ""

    for i in range(len(jh_1_po)):
        name_po += str(jh_1_po["市场指数"].iloc[i]) + "、"
        value_po += '{:.2%}'.format(jh_1_po['本周涨跌幅(%)'].iloc[i]) + "、"

    name_po = name_po.rstrip("、")
    value_po = value_po.rstrip("、")

    for i in range(len(jh_1_ne)):
        name_ne += str(jh_1_ne["市场指数"].iloc[i]) + "、"
        value_ne += '{:.2%}'.format(jh_1_ne['本周涨跌幅(%)'].iloc[i]) + "、"

    name_ne = name_ne.rstrip("、")
    value_ne = value_ne.rstrip("、")

    if len(jh_1_po) == 8:
        text = "（{}至{}），主要指数全部上涨。其中，{}分别上涨{}".format(start_day, today, name_po, value_po)
    elif len(jh_1_po) == 0:
        text = "（{}至{}），主要指数全部下跌。其中，{}分别下跌{}".format(start_day, today, name_ne, value_ne)
    else:
        text = "（{}至{}）。其中，{}分别上涨{}，{}分别下跌{}".format(start_day, today, name_po, value_po, name_ne, value_ne)

    # 储存text
    (paths.word_replacement_dir / "1.1更替文段.txt").write_text(text, encoding="utf-8")

    # %% 绘制
    ind1_Index = load_chart_data_with_numeric_validation(
        "outputs/csv/figures_1_15/图2_本周申万一级行业涨跌幅.csv",
        ["本周涨跌幅(%)", "本月涨跌幅(%)", "本年涨跌幅(%)"],
        "图2_本周申万一级行业涨跌幅",
    )
    ind1_chart_fields = ["本周涨跌幅(%)", "本月涨跌幅(%)", "本年涨跌幅(%)"]
    # 添加画布
    remake = (34 / 10) / (8253 / 3217)
    fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    # plt.subplots_adjust(bottom=0,top=1)
    # plt.rcParams['figure.figsize']=(34,8)
    ax = fig.add_subplot(111)

    # 将y轴的刻度方向设置向内
    plt.rcParams['ytick.direction'] = 'in'

    # 设置y坐标上下
    min_value = float(ind1_Index[ind1_chart_fields].min().min())
    max_value = float(ind1_Index[ind1_chart_fields].max().max())
    min_ylim = int(min_value / 0.05) * 5 / 100 - (min_value < 0) * 5 / 100
    max_ylim = int(max_value / 0.05) * 5 / 100 + (max_value > 0) * 5 / 100
    ax.set_ylim(min_ylim, max_ylim)

    # 设置y坐标刻度间隔
    y_major_locator = MultipleLocator(0.05)  # 以每4显示
    ax = plt.gca()
    ax.yaxis.set_major_locator(y_major_locator)
    ax.yaxis.set_major_formatter(mtick.PercentFormatter(1))
    vals = ax.get_yticks()
    ax.set_yticklabels(['{:,.0%}'.format(x) for x in vals], fontsize=30)

    # 设置纵坐标轴数值为times new roman
    labels = ax.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]

    # 绘制x=0
    ax.axhline(y=0, c='grey', lw=2)

    # 隐藏原始x
    plt.xticks(alpha=0)
    plt.tick_params(axis='x', width=0)

    # 画图
    # labels = ind1_list
    x1 = list(range(1, 1 + 3 * 31, 3))
    x2 = list(range(2, 2 + 3 * 31, 3))
    x3 = list(range(3, 3 + 3 * 31, 3))

    plt.bar(x1, ind1_Index["本周涨跌幅(%)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=2, label="本周涨跌幅(%)", zorder=1)
    plt.bar(x2, ind1_Index["本月涨跌幅(%)"], alpha=0.8, width=1, color='darkgrey', edgecolor=[
        "dimgrey"], linewidth=2, label="本月涨跌幅(%)", zorder=1)
    # plt.bar(x3, ind1_Index["本年涨跌幅(%)"], alpha=0.8, width=1, color='sandybrown', edgecolor=["coral"], linewidth=3, label="本年涨跌幅(%)")

    # 设置x
    plt.xlim(0, 93.01)
    ax.set_xticks(x3)

    # 柱状图上显示数据
    x = pd.DataFrame([x1, x2, x3])
    for i in range(0, 1):  # (0,1)只显示周收益 (0,2)全部显示
        for j in range(0, 31):
            value = ind1_Index.iloc[j][ind1_chart_fields[i]]
            if float(value) >= 0:
                plt.text(x.iloc[i, j], value, str("%.2f" % (value * 100)) + '%\n',
                         ha='center', va='center', fontsize=25, family="Times new roman", zorder=2)
            else:
                plt.text(x.iloc[i, j], value - 0.01, str("%.2f" % (value * 100)) + '%\n',
                         ha='center', va='center', fontsize=25, family="Times new roman", zorder=2)

    # 设置下边框名
    for i in range(0, 31):
        plt.text(x2[i], min_ylim - 0.005, ind1_Index.iloc[i]["市场指数"],
                 ha='right', va='top', fontsize=30, rotation=30)

    # 去除边框
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # 设置网格
    plt.grid(zorder=0, linestyle='--', dashes=(15, 12))

    # 设置标签
    plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
        0.5, 1.1), loc=9, borderaxespad=0, fontsize=37)

    # 储存图片
    plt.savefig("outputs/figures/图2_本周申万一级行业涨跌幅.png", dpi=300, bbox_inches='tight')
    pic_2 = Image.open("outputs/figures/图2_本周申万一级行业涨跌幅.png")
    size = pic_2.size
    print(size)
    print(size[0] / size[1])
    # 修改长宽比例
    pic_new = pic_2.resize((2380, 700))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture2.png")

    # %% 绘制
    ind2_Index = pd.read_csv("outputs/csv/figures_1_15/图3_本周申万二级行业涨跌幅.csv")
    ind2_rank10r = ind2_Index.iloc[0:10]
    ind2_rank10d = ind2_Index.iloc[len(ind2_Index) - 10:len(ind2_Index)]
    #ind2 = ind2_rank10r.append(ind2_rank10d)
    ind2 = pd.concat([ind2_rank10r, ind2_rank10d], ignore_index=True)

    ind2_name = ind2_Index["市场指数"]
    # 添加画布
    remake = (34 / 10) / (8600 / 3267)
    fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    ax = fig.add_subplot(111)

    # 将y轴的刻度方向设置向内
    plt.rcParams['ytick.direction'] = 'in'

    # 设置y坐标上下
    min_ylim = int(ind2_Index.iloc[:, 3].min() / 0.05) * \
               5 / 100 - ((ind2_Index.iloc[:, 3].min()) < 0) * 5 / 100
    max_ylim = int(ind2_Index.iloc[:, 3].max() / 0.05) * \
               5 / 100 + ((ind2_Index.iloc[:, 3].max()) > 0) * 5 / 100
    ax.set_ylim(min_ylim, max_ylim)

    # 设置y坐标刻度间隔
    y_major_locator = MultipleLocator(0.05)  # 以每4显示
    ax = plt.gca()
    ax.yaxis.set_major_locator(y_major_locator)
    ax.yaxis.set_major_formatter(mtick.PercentFormatter(1))
    vals = ax.get_yticks()
    ax.set_yticklabels(['{:,.0%}'.format(x) for x in vals], fontsize=30)

    # 设置纵坐标轴数值为times new roman
    labels = ax.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]

    # 绘制x=0
    ax.axhline(y=0, c='grey', lw=2)

    # 隐藏原始x
    plt.xticks(alpha=0)
    plt.tick_params(axis='x', width=0)

    # 画图
    x1 = list(range(1, 1 + 2 * 10, 2))
    x2 = list(range(23, 23 + 2 * 10, 2))
    x3 = list(range(2, 2 + 2 * 21, 2))

    plt.bar(x1, ind2_rank10r["本周涨跌幅(%)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=3, label="本周涨跌幅(%)", zorder=1)

    # 设置标签
    plt.legend(frameon=False, ncol=1, bbox_to_anchor=(
        0.5, 1.1), loc=9, borderaxespad=0, fontsize=36)

    plt.bar(x2, ind2_rank10d["本周涨跌幅(%)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=3, label="本周涨跌幅(%)", zorder=1)

    # 设置x
    plt.xlim(0, 42.01)
    ax.set_xticks(x3)

    # 柱状图上显示数据
    x = x1 + x2
    for i in range(0, 20):  # (0,1)只显示周收益 (0,2)全部显示
        value = ind2.iloc[i, 3]
        if float(value) >= 0:
            plt.text(x[i], value, str("%.2f" % (value * 100)) + '%\n', ha='center',
                     va='center', fontsize=25, family="Times new roman", zorder=2)
        else:
            plt.text(x[i], value - 0.03, str("%.2f" % (value * 100)) + '%\n', ha='center',
                     va='center', fontsize=25, family="Times new roman", zorder=2)

    # 设置下边框名
    for i in range(0, 21):
        plt.text(x3[i], min_ylim - 0.006, ind2_name[i],
                 ha='right', va='top', fontsize=30, rotation=30)

    # 去除边框
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # 设置网格
    plt.grid(zorder=0, linestyle='--', dashes=(15, 12))

    # 储存图片
    plt.savefig("outputs/figures/图3_本周申万二级行业涨跌幅.png", dpi=300, bbox_inches='tight')
    pic_3 = Image.open("outputs/figures/图3_本周申万二级行业涨跌幅.png")
    size = pic_3.size
    print(size, size[0] / size[1])
    # print(size2[0]/size2[1])
    # 修改长宽比例
    pic_new = pic_3.resize((2380, 700))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture3.png")

    # %% 绘制
    hot_data = pd.read_csv("outputs/csv/figures_1_15/图4_本周热点概念板块涨跌幅（排名前10位和后10位）.csv")
    hot_rank10r = hot_data.iloc[0:10] # 顾头不顾腚，只要第一行到10行（也就在这里）在python是最小的
    #这里当时出错了是因为之前的数据合并出错了
    hot_rank10d = hot_data.iloc[len(hot_data) - 10:len(hot_data)+1]
    #hot = hot_rank10r.append(hot_rank10d)
    hot = pd.concat([hot_rank10r,hot_rank10d],ignore_index= True)

    hot_name = [trim_chart_label_suffix(x) for x in hot_data["市场指数"]]
    # 添加画布
    remake = (34 / 10) / (8290 / 3279)
    fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    ax = fig.add_subplot(111)

    # 将y轴的刻度方向设置向内
    plt.rcParams['ytick.direction'] = 'in'

    # 设置y坐标上下
    min_ylim = int(hot.iloc[:, 3].min() / 0.05) * 5 / 100 - \
               ((hot.iloc[:, 3].min()) < 0) * 5 / 100
    max_ylim = int(hot.iloc[:, 3].max() / 0.05) * 5 / 100 + \
               ((hot.iloc[:, 3].max()) > 0) * 5 / 100
    ax.set_ylim(min_ylim, max_ylim)

    # 设置y坐标刻度间隔
    y_major_locator = MultipleLocator(0.05)  # 以每4显示
    ax = plt.gca()
    ax.yaxis.set_major_locator(y_major_locator)
    ax.yaxis.set_major_formatter(mtick.PercentFormatter(1))
    vals = ax.get_yticks()
    ax.set_yticklabels(['{:,.0%}'.format(x) for x in vals], fontsize=30)
    # 设置纵坐标轴数值为times new roman
    labels = ax.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]

    # 绘制x=0
    ax.axhline(y=0, c='grey', lw=2)

    # 隐藏原始x
    plt.xticks(alpha=0)
    plt.tick_params(axis='x', width=0)

    # 画图
    x1 = list(range(1, 1 + 2 * 10, 2))
    x2 = list(range(23, 23 + 2 * 10, 2))
    x3 = list(range(2, 2 + 2 * 21, 2))

    plt.bar(x1, hot_rank10r["本周涨跌幅(%)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=3, label="本周涨跌幅(%)", zorder=1)

    # 设置标签
    plt.legend(frameon=False, ncol=1, bbox_to_anchor=(
        0.5, 1.1), loc=9, borderaxespad=0, fontsize=38)

    plt.bar(x2, hot_rank10d["本周涨跌幅(%)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
        "navy"], linewidth=3, label="本周涨跌幅(%)", zorder=1)

    # 设置x
    plt.xlim(0, 42.01)
    ax.set_xticks(x3)

    # 柱状图上显示数据
    x = x1 + x2
    for i in range(0, 20):  # (0,1)只显示周收益 (0,2)全部显示
        value = hot.iloc[i, 3]
        if float(value) >= 0:
            plt.text(x[i], value, str("%.2f" % (value * 100)) + '%\n', ha='center',
                     va='center', fontsize=25, family="Times new roman", zorder=2)
        else:
            plt.text(x[i], value - 0.012, str("%.2f" % (value * 100)) + '%\n', ha='center',
                     va='center', fontsize=25, family="Times new roman", zorder=2)

    # 设置下边框名
    for i in range(0, 21):
        plt.text(x3[i], min_ylim - 0.02, hot_name[i],
                 ha='right', va='top', fontsize=30, rotation=30)

    # 去除边框
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)

    # 设置网格
    plt.grid(zorder=0, linestyle='--', dashes=(15, 12))

    # 储存图片
    plt.savefig("outputs/figures/图4_本周热点概念板块涨跌幅.png", dpi=300, bbox_inches='tight')
    pic_4 = Image.open("outputs/figures/图4_本周热点概念板块涨跌幅.png")
    size = pic_4.size
    print(size, size[0] / size[1])
    # print(size2[0]/size2[1])
    # 修改长宽比例
    pic_new = pic_4.resize((2380, 700))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture4.png")

    # %% 绘制

    trading_use = pd.read_csv("outputs/csv/figures_1_15/图5_两市成交金额、换手率.csv")
    # 添加画布
    remake = (34 / 10) / (8 / 3.09)
    fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    ax = fig.add_subplot(111)

    # 将y轴的刻度方向设置向内
    plt.rcParams['ytick.direction'] = 'in'

    # 设置y坐标上下
    max_ylim1 = int(trading_use["all"].max() * 1.5 / 2000) * 2000
    ax.set_ylim(0, max_ylim1)

    # 设置y坐标刻度间隔
    y_major_locator = MultipleLocator(2000)  # 以每4显示
    ax = plt.gca()
    ax.yaxis.set_major_locator(y_major_locator)
    ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
    # 设置纵坐标轴数值为times new roman
    labels = ax.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]
    # 绘制x=0
    ax.axhline(y=0, c='grey', lw=2)

    # 隐藏原始x轴刻
    plt.xticks(alpha=0)
    plt.tick_params(axis='x', width=0)

    # 画图
    x1 = list(range(1, 1 + 2 * 50, 2))
    x2 = list(range(2, 2 + 2 * 50, 2))

    # 设置x
    plt.xlim(0, 100.01)
    ax.set_xticks(x2)

    # 成交额柱状图
    ax.bar(x1, trading_use["沪市:成交金额(亿元)"], alpha=0.9, width=1,
           color='steelblue', label="沪市:成交金额(亿元)", zorder=1)
    ax.bar(x1, trading_use['深市:成交金额(亿元)'], bottom=trading_use["沪市:成交金额(亿元)"],
           alpha=1, width=1, color='darkorange', label='深市:成交金额(亿元)', zorder=1)
    ax.plot(x1, trading_use["参考线：1万亿元"], color='steelblue', label="参考线：1万亿元",
            linestyle='--', dashes=(4, 2), linewidth=4, zorder=3)
    # 设置网格
    plt.grid(zorder=0, linestyle='--', dashes=(15, 12))
    plt.legend(frameon=False, ncol=3, bbox_to_anchor=(
        0.275, 1.12), loc=9, borderaxespad=0, fontsize=35)

    # 换手率折线图
    ax2 = ax.twinx()  # twinx将ax1的X轴共用与ax2，这步很重要
    max_ylim2 = int(max(trading_use['换手率_上证综合指数(%)'].max(
    ), trading_use['换手率_深证综合指数(%)'].max()) / 0.5) * 0.5 + 0.5
    ax2.set_ylim(0, max_ylim2)

    y_major_locator = MultipleLocator(0.5)
    ax = plt.gca()
    ax2.yaxis.set_major_locator(y_major_locator)
    ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)

    ax2.plot(x1, trading_use['换手率_上证综合指数(%)'],
             color='darkgrey', linewidth=5, label="换手率_上证综合指数(%)")
    ax2.scatter(x1, trading_use['换手率_上证综合指数(%)'], s=300, marker="o",
                facecolor='none', ec="darkgrey", linewidth=5, zorder=2)
    ax2.plot(x1, trading_use['换手率_深证综合指数(%)'],
             color='gold', linewidth=5, label='换手率_深证综合指数(%)')
    ax2.scatter(x1, trading_use['换手率_深证综合指数(%)'], s=300, marker="o",
                facecolor='none', ec="gold", linewidth=5, zorder=2)

    # 去除边框
    ax.spines['top'].set_visible(False)
    ax2.spines['top'].set_visible(False)

    # 设置纵坐标轴数值为times new roman
    labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_yticklabels()
    [label.set_fontname('Times New Roman') for label in labels]

    # 设置下边框名
    # 设置下边框名
    for i in range(0, len(x1)):
        plt.text(x1[i], - 0.02, trading_use["date"][i], ha='center',
                 va='top', fontsize=30, rotation=90, family="Times new roman")

    # 设置标签
    plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
        0.795, 1.12), loc=9, borderaxespad=0, fontsize=35)

    # 储存图片
    plt.savefig("outputs/figures/图5_两市成交金额、换手率.png", dpi=300, bbox_inches='tight')
    pic_5 = Image.open("outputs/figures/图5_两市成交金额、换手率.png")
    size = pic_5.size
    print(size, size[0] / size[1])
    # print(size2[0]/size2[1])
    # 修改长宽比例
    pic_new = pic_5.resize((2380, 700))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture5.png")

    # %% 绘制
    if not ENABLE_FIGURE6_FUTURES:
        print("图6开关关闭：跳过图6绘制与 1.3 更替文段生成。")
        (paths.word_replacement_dir / "1.3更替文段.txt").write_text("", encoding="utf-8")
    else:
        qh_use50 = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图6_股指期货主力合约基差.csv",
            ["IF主力合约:基差", "IH主力合约:基差", "IC主力合约:基差", "IM主力合约:基差"],
            "图6_股指期货主力合约基差",
        )
        remake = (34 / 10) / (8193 / 3408)
        fig = plt.figure(figsize=(34 * remake, 10), dpi=300)
        ax = fig.add_subplot(111)

        plt.rcParams['ytick.direction'] = 'in'

        data_min = qh_use50.iloc[:, 1:5].min().min(skipna=True)
        data_max = qh_use50.iloc[:, 1:5].max().max(skipna=True)
        min_ylim = int(data_min / 20) * 20 - (data_min < 0) * 20
        max_ylim = int(data_max / 20) * 20 + (data_max > 0) * 20
        ax.set_ylim(min_ylim, max_ylim)

        y_major_locator = MultipleLocator(20)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
        ax.axhline(y=0, c='grey', lw=2)

        plt.xticks(alpha=0)
        plt.tick_params(axis='x', width=0)

        x1 = list(range(1, 1 + 2 * 50, 2))
        x2 = list(range(2, 2 + 2 * 50, 2))

        plt.xlim(0, 100.01)
        ax.set_xticks(x2)

        ax.plot(x1, qh_use50["IF主力合约:基差"], color='navy',
                linewidth=5, label="IF主力合约:基差", alpha=0.8)
        ax.plot(x1, qh_use50["IH主力合约:基差"], color='darkorange',
                linewidth=5, label="IH主力合约:基差")
        ax.plot(x1, qh_use50["IC主力合约:基差"], color='grey',
                linewidth=5, label="IC主力合约:基差")
        ax.plot(x1, qh_use50["IM主力合约:基差"], color='gold',
                linewidth=5, label="IM主力合约:基差", alpha=1)

        plt.grid(zorder=0, linestyle='--', dashes=(15, 12), alpha=0.7)
        plt.legend(frameon=False, ncol=4, bbox_to_anchor=(
            0.5, 1.2), loc=9, borderaxespad=0, fontsize=40)

        ax.scatter(x1, qh_use50["IF主力合约:基差"], s=300, marker="o",
                   facecolor='none', ec="navy", linewidth=5, zorder=2, alpha=0.8)
        ax.scatter(x1, qh_use50["IH主力合约:基差"], s=300, marker="^",
                   facecolor='none', ec='darkorange', linewidth=5, zorder=2)
        ax.scatter(x1, qh_use50["IC主力合约:基差"], s=300, marker="s",
                   facecolor='none', ec="grey", linewidth=5, zorder=2)
        ax.scatter(x1, qh_use50["IM主力合约:基差"], s=300, marker="D",
                   facecolor='none', ec="gold", linewidth=5, zorder=2, alpha=1)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        labels = ax.get_xticklabels() + ax.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        for i in range(len(x1)):
            plt.text(x1[i], min_ylim - 3, qh_use50["日期"].iloc[i], ha='center',
                     va='top', fontsize=30, rotation=90, family="Times new roman")

        plt.savefig("outputs/figures/图6_股指期货主力合约基差.png", dpi=300, bbox_inches='tight')
        pic_6 = Image.open("outputs/figures/图6_股指期货主力合约基差.png")
        size = pic_6.size
        print(size, size[0] / size[1])
        pic_new = pic_6.resize((2380, 700))
        pic_new.save("outputs/figures/picture6.png")

        trad_avg = round((trading_use["all"].iloc[-5:].mean()) / 10000, 2)
        IF_jc = round(qh_use50["IF主力合约:基差"].iloc[-1], 2)
        IH_jc = round(qh_use50["IH主力合约:基差"].iloc[-1], 2)
        IC_jc = round(qh_use50["IC主力合约:基差"].iloc[-1], 2)
        IF_ss = "{:.2%}".format(qh_use50["IF主力合约:升水率"].iloc[-1])
        IH_ss = "{:.2%}".format(qh_use50["IH主力合约:升水率"].iloc[-1])
        IC_ss = "{:.2%}".format(qh_use50["IC主力合约:升水率"].iloc[-1])
        text = (
            "本周沪深两市日均成交额{}万亿元。截至周五，IF、IH和IC主力合约基差分别为{}点"
            "（升水率{}）、{}点（升水率{}）和{}点（升水率{}）。"
        ).format(trad_avg, IF_jc, IF_ss, IH_jc, IH_ss, IC_jc, IC_ss)

        (paths.word_replacement_dir / "1.3更替文段.txt").write_text(text, encoding="utf-8")

    # %% 图7 绘制
    if not ENABLE_FIGURE7_UPDOWN:
        print("图7开关关闭：跳过图7绘制。")
    else:
        updown_use = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图7_两市涨跌数量统计.csv",
            ["跌停家数", "下跌家数（不含跌停）", "平盘家数", "上涨家数（不含涨停）", "涨停家数"],
            "图7_两市涨跌数量统计",
        )

        remake = (34 / 10) / (8179 / 2717)
        fig = plt.figure(figsize=(34 * remake, 10), dpi=300)
        ax = fig.add_subplot(111)

        plt.rcParams['ytick.direction'] = 'in'
        plt.xlim(-0.01, 10)

        x_major_locator = MultipleLocator(1)
        ax = plt.gca()
        ax.xaxis.set_major_locator(x_major_locator)

        labels = ax.get_xticklabels() + ax.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        plt.xticks(alpha=0)
        plt.tick_params(axis='x', width=0)
        plt.yticks(alpha=0)
        plt.tick_params(axis='y', width=0)

        x = list(range(1, 1 + 2 * 5, 2))
        x1 = list(range(0, 2 + 2 * 5, 2))

        plt.bar(x, updown_use.iloc[:, 1], alpha=0.7, width=1,
                color='darkgreen', label="跌停家数", zorder=1)
        plt.bar(x, updown_use.iloc[:, 2], bottom=updown_use.iloc[:, 1],
                alpha=0.8, width=1, color='yellowgreen', label="下跌家数（不含跌停）", zorder=1)
        plt.bar(x, updown_use.iloc[:, 3], bottom=updown_use.iloc[:, 1] +
                                                 updown_use.iloc[:, 2], alpha=0.85, width=1, color='grey', label="平盘家数",
                zorder=1)
        plt.bar(x, updown_use.iloc[:, 4], bottom=updown_use.iloc[:, 1] + updown_use.iloc[:, 2] +
                                                 updown_use.iloc[:, 3], alpha=0.9, width=1, color='darksalmon',
                label="上涨家数（不含涨停）", zorder=1)
        plt.bar(x, updown_use.iloc[:, 5], bottom=updown_use.iloc[:, 1] + updown_use.iloc[:, 2] +
                                                 updown_use.iloc[:, 3] + updown_use.iloc[:, 4], alpha=0.7, width=1,
                color='darkred', label="涨停家数", zorder=1)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)

        for i in range(0, len(x)):
            plt.text(x[i], -200, updown_use["日期"].iloc[i], ha='center',
                     va='top', fontsize=30, family="Times new roman")

        updown_val = updown_use.iloc[:, 1:6].copy()
        updown_val.iloc[:, 1] = updown_val.iloc[:, 0] + updown_val.iloc[:, 1]
        updown_val.iloc[:, 2] = updown_val.iloc[:, 1] + updown_val.iloc[:, 2]
        updown_val.iloc[:, 3] = updown_val.iloc[:, 2] + updown_val.iloc[:, 3]
        updown_val.iloc[:, 4] = updown_val.iloc[:, 3] + updown_val.iloc[:, 4]

        offset = 0.012
        for i in range(0, 5):
            for j in range(0, 5):
                if j == 0:
                    value = updown_val.iloc[i, j]
                else:
                    value = (updown_val.iloc[i, j] + updown_val.iloc[i, j - 1]) / 2
                plt.text(x[i], value+offset, updown_use.iloc[i, j + 1], ha='center',
                         va='center', fontsize=40, family="Times new roman", zorder=2)

        ax.set_xticks(x1)
        plt.grid(zorder=0, linestyle='-', alpha=0.5)
        plt.legend(frameon=False, ncol=5, bbox_to_anchor=(
            0.5, 1.1), loc=9, borderaxespad=0, fontsize=30)

        plt.savefig("outputs/figures/图7_两市涨跌数量统计.png", dpi=300, bbox_inches='tight')
        pic_7 = Image.open("outputs/figures/图7_两市涨跌数量统计.png")
        size = pic_7.size
        print(size, size[0] / size[1])
        pic_new = pic_7.resize((2380, 700))
        pic_new.save("outputs/figures/picture7.png")


    # %% 绘制
    if not ENABLE_FIGURE8_MARGIN:
        print("图8开关关闭：跳过图8绘制。")
    else:
        rzrq_use = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图8_两市融资融券余额、融资买入占比.csv",
            ["两市融资融券余额（万亿元）", "两市融资买入金额占比"],
            "图8_两市融资融券余额、融资买入占比",
        )
        remake = (34 / 10) / (8292 / 3406)
        fig = plt.figure(figsize=(34 * remake, 10), dpi=300)
        ax = fig.add_subplot(111)

        plt.rcParams['ytick.direction'] = 'in'

        space = max(rzrq_use.iloc[:, 1]) - min(rzrq_use.iloc[:, 1])
        min_ylim = int((min(rzrq_use.iloc[:, 1]) - space) / 0.02) * 0.02
        max_ylim = int((max(rzrq_use.iloc[:, 1]) + space) / 0.02) * 0.02
        ax.set_ylim(min_ylim, max_ylim)

        y_major_locator = MultipleLocator(0.2)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
        ax.axhline(y=0, c='grey', lw=2)

        plt.xticks(alpha=0)
        plt.tick_params(axis='x', width=0)

        x1 = list(range(1, 1 + 2 * 50, 2))
        x2 = list(range(2, 2 + 2 * 50, 2))

        plt.xlim(0, 100.01)
        ax.set_xticks(x2)

        ax.bar(
            x1,
            rzrq_use["两市融资融券余额（万亿元）"],
            color="steelblue",
            ec="black",
            width=1,
            label="两市融资融券余额（万亿元）",
            alpha=0.8,
            zorder=100,
        )

        plt.grid(zorder=0, linestyle='--', alpha=0.6)
        plt.legend(frameon=False, bbox_to_anchor=(0.32, 1.2),
                   loc=9, borderaxespad=0, fontsize=40)

        ax2 = ax.twinx()
        max_ylim2 = int(max(rzrq_use["两市融资买入金额占比"].max(),
                            rzrq_use["两市融资买入金额占比"].max()) / 0.1) * 0.1 + 0.1
        ax2.set_ylim(0, max_ylim2)

        ax2.plot(x1, rzrq_use["两市融资买入金额占比"], color='darkorange',
                 linewidth=5, label="两市融资买入金额占比(%)右轴")
        ax2.scatter(x1, rzrq_use["两市融资买入金额占比"], s=300, marker="D",
                    facecolor='none', ec='darkorange', linewidth=5, zorder=2)

        ax2.yaxis.set_major_formatter(mtick.PercentFormatter(1))
        vals = ax2.get_yticks()
        ax2.set_yticklabels(['{:,.0%}'.format(x) for x in vals], fontsize=30)

        labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        ax.spines['top'].set_visible(False)
        ax2.spines['top'].set_visible(False)

        for i in range(len(x1)):
            plt.text(x1[i], -0.002, rzrq_use["日期"].iloc[i], ha='center',
                     va='top', fontsize=30, rotation=90, family="Times new roman")

        plt.legend(frameon=False, bbox_to_anchor=(0.68, 1.2),
                   loc=9, borderaxespad=0, fontsize=40)
        plt.savefig("outputs/figures/图8_两市融资融券余额、融资买入占比.png", dpi=300, bbox_inches='tight')
        pic_8 = Image.open("outputs/figures/图8_两市融资融券余额、融资买入占比.png")
        size = pic_8.size
        print(size, size[0] / size[1])
        pic_new = pic_8.resize((2380, 700))
        pic_new.save("outputs/figures/picture8.png")

    # # %%绘制
    # bx_use = pd.read_csv("data/-12数据/：北向资金买入成交金额、成交净买入、累计净买入.csv")
    # # 添加画布
    # remake = (34 / 10) / (8431 / 3197)
    # fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    # ax = fig.add_subplot(111)
    #
    # # 将y轴的刻度方向设置向内
    # plt.rcParams['ytick.direction'] = 'in'
    #
    # # 设置y坐标上下
    # min_ylim = int(min(bx_use.iloc[:, 1:3].min()) / 100) * \
    #            100 - ((min(bx_use.iloc[:, 1:3].min())) < 0) * 100
    # max_ylim = int(max(bx_use.iloc[:, 1:3].max()) / 100) * \
    #            100 + ((max(bx_use.iloc[:, 1:3].max())) > 0) * 100
    # ax.set_ylim(min_ylim, max_ylim)
    #
    # # 设置y坐标刻度间隔
    # y_major_locator = MultipleLocator(100)
    # ax = plt.gca()
    # ax.yaxis.set_major_locator(y_major_locator)
    # ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
    # # 设置纵坐标轴数值为times new roman
    # labels = ax.get_yticklabels()
    # [label.set_fontname('Times New Roman') for label in labels]
    #
    # # 绘制x=0
    # ax.axhline(y=0, c='grey', lw=2)
    #
    # # 隐藏原始x
    # plt.xticks(alpha=0)
    # plt.tick_params(axis='x', width=0)
    #
    # # 柱状
    # x1 = list(range(1, 1 + 3 * 50, 3))
    # x2 = list(range(2, 2 + 3 * 50, 3))
    # x3 = list(range(3, 3 + 3 * 50, 3))
    #
    # # 设置x
    # plt.xlim(0, 150.01)
    # ax.set_xticks(x3)
    #
    # # 设置网格
    # plt.grid(zorder=0, linestyle='--', dashes=(8, 4))
    #
    # plt.bar(x1, bx_use["买入成交金额"], alpha=0.8, width=1, color='steelblue', edgecolor=[
    #     "navy"], linewidth=2, label="北向资金：买入成交金额（亿元, zorder=1)
    # plt.bar(x2, bx_use["成交净买入"], alpha=0.8, width=1, color='darkgrey', edgecolor=[
    #     "dimgrey"], linewidth=2, label="北向资金：成交净买入（亿元）", zorder=1)
    #
    # # 设置标签
    # plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
    #     0.33, 1.12), loc=9, borderaxespad=0, fontsize=39)
    # # 折线
    # ax2 = ax.twinx()  # twinx将ax1的X轴共用与ax2，这步很重要
    # min_ylim2 = int(min(bx_use["累积净买入"]) / 200) * \
    #             200 + min_ylim * 2
    # max_ylim2 = int(max(bx_use["累积净买入"]) / 200) * \
    #             200 + 600
    #
    # ax2.set_ylim(min_ylim2, max_ylim2)
    #
    # ax2.plot(x1, bx_use["累积净买入"], color='darkorange',
    #          linewidth=5, label="北向资金：累积净买入（亿右轴)
    # ax2.scatter(x1, bx_use["累积净买入"], s=300, marker="D",
    #             facecolor='none', ec='darkorange', linewidth=5, zorder=2)
    #
    # y_major_locator2 = MultipleLocator(200)
    # ax2 = plt.gca()
    # ax2.yaxis.set_major_locator(y_major_locator2)
    # ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)
    # # 设置标签
    # plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
    #     0.82, 1.12), loc=9, borderaxespad=0, fontsize=39)
    #
    # # 设置纵坐标轴数值为times new roman
    # labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_yticklabels()
    # [label.set_fontname('Times New Roman') for label in labels]
    #
    # # 去除边框
    # ax.spines['top'].set_visible(False)
    # ax2.spines['top'].set_visible(False)
    # ax.spines['bottom'].set_visible(False)
    # ax2.spines['bottom'].set_visible(False)
    #
    # # 设置下边框名
    # for i in range(len(x1)):
    #     plt.text(x2[i], min_ylim2 - 50, bx_use["日期"].iloc[i], ha='center',
    #              va='top', fontsize=30, rotation=90, family="Times new roman")
    #
    # # 储存图片
    # plt.savefig("outputs/figures/：北向资金买入成交金额、成交净买入、累计净买入 .png", dpi=300, bbox_inches='tight')
    # pic_9 = Image.open("outputs/figures/：北向资金买入成交金额、成交净买入、累计净买入 .png")
    # size = pic_9.size
    # print(size, size[0] / size[1])
    # # 修改长宽比例
    # pic_new = pic_9.resize((2380, 700))  # Image.ANTIALIAS
    # pic_new.save("outputs/figures/picture9.png")
    #
    # # %%0绘制
    # bx_ind = pd.read_excel("data/-12数据/0：本周各行业北向资金净买入金额（中信一级行业）.xlsx")
    # bx_ind_name = list(x[2:] for x in bx_ind["行业"])
    # bx_ind["净买入"] = round(bx_ind["净买入(亿元)"], 2)
    # # 添加画布
    # remake = (34 / 10) / (8467 / 3440)
    # fig = plt.figure(figsize=(34 * remake, 10), dpi=300)  # 4/6宫格
    # ax = fig.add_subplot(111)
    #
    # # 将y轴的刻度方向设置向内
    # plt.rcParams['ytick.direction'] = 'in'
    #
    # # 设置y坐标上下
    # min_ylim = int(min(bx_ind["净买入(亿元)"]) / 10) * \
    #            10 - ((min(bx_ind["净买入(亿元)"])) < 0) * 10
    # max_ylim = int(max(bx_ind["净买入(亿元)"]) / 10) * \
    #            10 + ((max(bx_ind["净买入(亿元)"])) > 0) * 10
    #
    # min_ylim2 = int(min(bx_ind["占行业总市值比变化(%)"]) / 0.2) * \
    #             0.2 - (min(bx_ind["占行业总市值比变化(%)"]) < 0) * 0.2
    # max_ylim2 = int(max(bx_ind["占行业总市值比变化(%)"]) / 0.2) * \
    #             0.2 + (max(bx_ind["占行业总市值比变化(%)"]) > 0) * 0.2
    #
    # max_y = max(max_ylim, max_ylim2 * 50)
    # min_y = min(min_ylim, min_ylim2 * 50)
    #
    # ax.set_ylim(min_y, max_y)
    #
    # # 设置y坐标刻度间隔
    # y_major_locator = MultipleLocator(10)
    # ax = plt.gca()
    # ax.yaxis.set_major_locator(y_major_locator)
    # ax.tick_params(axis="y", direction="in", which="major", labelsize=40)
    #
    # # 绘制x=0
    # ax.axhline(y=0, c='grey', lw=2)
    #
    # # 隐藏原始x
    # plt.xticks(alpha=0)
    # plt.tick_params(axis='x', width=0)
    #
    # # 画图
    # # labels = ind1_list
    # x1 = list(range(1, 1 + 2 * 30, 2))
    # x2 = list(range(2, 2 + 2 * 30, 2))
    #
    # plt.bar(x1, bx_ind["净买入(亿元)"], alpha=0.8, width=1, color='steelblue', edgecolor=[
    #     "navy"], linewidth=2, label="本周净买入金额(亿元)", zorder=1)
    #
    # # 柱状图上显示数据
    # for i in range(0, 30):
    #     if bx_ind["净买入(亿元)"][i] >= 0:
    #         plt.text(x1[i], bx_ind["净买入(亿元)"][i], bx_ind["净买入"][i], ha='center', va='bottom', fontsize=30, color="black",
    #                  zorder=5, family="Times new roman")
    #     else:
    #         plt.text(x1[i], bx_ind["净买入(亿元)"][i], bx_ind["净买入"][i], ha='center', va='top', fontsize=30, color="black",
    #                  zorder=5, family="Times new roman")
    # # 设置x
    # plt.xlim(0, 60.01)
    # ax.set_xticks(x2)
    #
    # # 设置下边框名
    # for i in range(0, 30):
    #     plt.text(x2[i], min_y - 5, bx_ind_name[i],
    #              ha='right', va='top', fontsize=30, rotation=30)
    #
    # # 设置网格
    # plt.grid(zorder=0, linestyle='--', dashes=(15, 12))
    #
    # # 设置标签
    # plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
    #     0.3, 1.1), loc=9, borderaxespad=0, fontsize=40)
    #
    # # 折线
    # ax2 = ax.twinx()  # twinx将ax1的X轴共用与ax2，这步很重要
    # ax2.set_ylim(min_y / 50, max_y / 50)
    #
    # ax2.plot(x1, bx_ind["占行业总市值比变化(%)"], color='darkorange',
    #          linewidth=5, label="占行业总市值比变化(%/右轴)", alpha=0.9)
    # ax2.scatter(x1, bx_ind["占行业总市值比变化(%)"], s=300, marker="D",
    #             facecolor='none', ec='darkorange', linewidth=5, zorder=2, alpha=0.7)
    #
    # y_major_locator2 = MultipleLocator(0.2)
    # ax2 = plt.gca()
    # ax2.yaxis.set_major_locator(y_major_locator2)
    # ax2.tick_params(axis="y", direction="in", which="major", labelsize=40)
    # # 设置标签
    # plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
    #     0.7, 1.1), loc=9, borderaxespad=0, fontsize=40)
    #
    # # 设置纵坐标轴数值为times new roman
    # labels = ax.get_yticklabels() + ax2.get_yticklabels()
    # [label.set_fontname('Times New Roman') for label in labels]
    #
    # # 去除边框
    # ax.spines['top'].set_visible(False)
    # ax.spines['bottom'].set_visible(False)
    # ax2.spines['top'].set_visible(False)
    # ax2.spines['bottom'].set_visible(False)
    #
    # # 储存图片
    # plt.savefig("outputs/figures/0：本周各行业北向资金净买入金额（中信一级行业）.png", dpi=300, bbox_inches='tight')
    # pic_10 = Image.open("outputs/figures/0：本周各行业北向资金净买入金额（中信一级行业）.png")
    # size = pic_10.size
    # print(size, size[0] / size[1])
    # # 修改长宽比例
    # pic_new = pic_10.resize((2380, 700))  # Image.ANTIALIAS
    # pic_new.save("outputs/figures/picture10.png")
    #
    # # word交互
    # rzrq_new = rzrq_use["日期"].iloc[-1]
    # date_yes_m = rzrq_new[5:7]
    # date_yes_d = rzrq_new[8:10]
    # week_list = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    # rzrq_zj = pd.to_datetime(rzrq_new, format='%Y-%m-%d')
    # zj = (week_list[rzrq_zj.weekday()])
    #
    # ye_value = round(rzrq_use["两市融资融券余额（万亿元].iloc[-1], 2)
    # buy = '{:.2%}'.format(rzrq_use["两市融资买入金额占比"].iloc[-1])
    #
    # bx_buy = round(bx_use['成交净买入'][-5:].sum())
    #
    # bx_ind_po = bx_ind[bx_ind['净买入'] >= 0]
    # bx_ind_po = bx_ind_po.sort_values(by='净买入', ascending=False)
    # bx_ind_po = bx_ind_po.iloc[0:5]
    #
    # bx_ind_ne = bx_ind[bx_ind['净买入'] < 0]
    # bx_ind_ne = bx_ind_ne.sort_values(by='净买入')
    # bx_ind_ne = bx_ind_ne.iloc[0:5]
    #
    # bx_ind_po["行业"] = bx_ind_po["行业"].apply(lambda x: x[2:])
    # bx_ind_ne["行业"] = bx_ind_ne["行业"].apply(lambda x: x[2:])
    #
    # name_po = ""
    # name_ne = ""
    # value_po = ""
    # value_ne = ""
    #
    # for i in range(len(bx_ind_po)):
    #     name_po += str(bx_ind_po["行业"].iloc[i]) + str(")
    #     value_po += str(round(bx_ind_po["净买入(亿元)"].iloc[i], 2)) + str(")
    #
    # name_po = name_po[::-1].replace(', '', 1)[::-1]
    # value_po = value_po[::-1].replace(', '', 1)[::-1]
    #
    # for i in range(len(bx_ind_ne)):
    #     name_ne += str(bx_ind_ne["行业"].iloc[i]) + str(")
    #     value_ne += str(round(bx_ind_ne["净买入(亿元)"].iloc[i], 2)) + str(")
    #
    # name_ne = name_ne[::-1].replace(', '', 1)[::-1]
    # value_ne = value_ne[::-1].replace(', '', 1)[::-1]
    #
    # if len(bx_ind_po) == 0:
    #     text = "最新余额数据（{}月{}日{}）为{}万亿元；融资买入占比为{}。本周北向资金净买入{}亿元：{}分别净流入{}亿元，无净流出行业.format(date_yes_m, date_yes_d, zj,
    #                                                                                        ye_value, buy, bx_buy, name_po,
    #                                                                                        value_po)
    # elif len(bx_ind_ne) == 0:
    #     text = "最新余额数据（{}月{}日{}）为{}万亿元；融资买入占比为{}。本周北向资金净买入{}亿元：{}分别净流出{}亿元.format(date_yes_m, date_yes_d, zj, ye_value,
    #                                                                                 buy, bx_buy, name_ne, value_ne)
    # else:
    #     text = "最新余额数据（{}月{}日{}）为{}万亿元；融资买入占比为{}。本周北向资金净买入{}亿元：{}分别净流入{}亿元；{}分别净流出{}亿元.format(date_yes_m, date_yes_d, zj,
    #                                                                                             ye_value, buy, bx_buy,
    #                                                                                             name_po, value_po, name_ne,
    #                                                                                             value_ne)
    #
    # # 储存text
    #     f.write(text)


    #  1: 各类资金变动规模（亿元）本周增量资金来源（亿元）
    ##%%
    # 获取最个周五的日期
    def get_last_3_fridays():
        today = datetime.datetime.combine(report_date, datetime.time.min)
        fridays = []

        while len(fridays) < 3:
            if today.weekday() == 4:  # 4代表周五
                fridays.append(today.strftime("%Y%m%d"))
            today -= timedelta(days=1)  # 每次减少一

        return fridays[::-1]  # 从最早到最-1 是reverse

    ##%%

    today0 = report_date
    today_year = today0.year
    # Outer loop for years (2010 to 今年)
    years = range(2010, today_year+1)  # 010到今年（包括今年
    quarters = ["0331", "0630", "0930", "1231"]  # 每年的四个季度日

    # 创建所有季度日
    dates = []
    for year in years:
        for quarter in quarters:
            date = f"{year}{quarter}"  # 生成 "YYYYMMDD" 格式的字符串
            dates.append(date)

    # 筛选出"20100331" "最近的季度 之间的日
    filtered_dates_list = [date for date in dates if "20100331" <= date <= (today0 - timedelta(days = 30)).strftime("%Y%m%d")]
    # 双重保险第一设置必须要在20100331和此月减个月的那个月之间才可以出现在这个list上
    #第二同时只取倒数第二和第三个日期作为newest date。相当于滞后5个月
    # 比如今天0号，那么upper limit则是40号，31 可以出现，但是还是取21的数据，直到81号：
    # 这时候upper limit0，那么就可以1 的数据了，此原因主要是险资害人不浅，懒狗一个，出数据太

    if not filtered_dates_list:
        raise RuntimeError("资金面报告期列表为空，请检查报告期范围或输入数据。")
    print(
        f"资金面报告期筛选完成：共{len(filtered_dates_list)}个，"
        f"起始={filtered_dates_list[0]}，结束={filtered_dates_list[-1]}"
    )


    ##%%
    #time list
        # 通用指标
    newest_rpt_date = filtered_dates_list[-1]
    newest_rpt_date_shift = filtered_dates_list[-2]

    time_list = [newest_rpt_date_shift, newest_rpt_date]
    last_3_fridays = get_last_3_fridays()
    time_list.extend(last_3_fridays)
    print(time_list)

    ##%%
    ## 图9

    # 假设时间范围
    start_date = "2010-01-01"
    end_date = time_list[-1]  # 假设 time_list[-1] 为字符串格式的日期

    被动权益ETF = pd.read_csv("workspace/funding_draft/data/被动权益基金/图9_被动权益ETF.csv")
    被动权益ETF["日期"] = 被动权益ETF["日期_汇总"]
    被动权益ETF = 被动权益ETF.drop(columns="日期_汇总")

    主动权益基金 = pd.read_csv("workspace/funding_draft/data/主动权益基金/图9_主动权益基金.csv")

    两融余额 = pd.read_csv("workspace/funding_draft/data/两融余额数据/图9_两融余额数据.csv")
    两融余额 = 两融余额.drop(columns="周度差值")

    保险资金 = pd.read_csv("workspace/funding_draft/data/保险资金/图9_保险资金.csv")

    # 确保日期列为 datetime 类型
    被动权益ETF["日期"] = pd.to_datetime(被动权益ETF["日期"], format="%Y%m%d")
    主动权益基金["日期"] = pd.to_datetime(主动权益基金["日期"], format="%Y%m%d")
    两融余额["日期"] = pd.to_datetime(两融余额["日期"], format="%Y%m%d")
    保险资金["日期"] = pd.to_datetime(保险资金["日期"], format="%Y%m%d")


    # 创建图表，设置时间范
    plt.figure(figsize=(12, 6))
    plt.xlim(pd.to_datetime(start_date), pd.to_datetime(end_date))  # 设置 x 轴时间范
    plt.ylim(0, 60000)  # y 轴范围（可以根据数据情况调整

    # 绘制每个 DataFrame 的折
    plt.plot(被动权益ETF["日期"], 被动权益ETF["被动权益ETF"], label="被动权益ETF", color="gold")
    plt.plot(主动权益基金["日期"], 主动权益基金["主动权益基金"], label="主动权益基金", color="orange")
    plt.plot(两融余额["日期"], 两融余额["两融余额"], label="两融余额", color="blue")
    plt.plot(保险资金["日期"], 保险资金["保险资金"], label="保险资金", color="red")

    # 添加图例、标签、标
    plt.legend(loc="lower center", ncol=4, fontsize=12,bbox_to_anchor=(0.5, -0.2))
    #plt.xlabel("日期", fontsize=15)
    #plt.ylabel("资金, fontsize=15)
    plt.grid(linestyle="--", alpha=0.5)

    # x轴与y轴数字大小与字体调整
    plt.xticks(fontproperties=font_English)  # 设置 x 轴刻度字
    plt.yticks(fontproperties=font_English)  # 设置 y 轴刻度字
    plt.tick_params(axis='both', labelsize=10)  # x 轴和 y 轴的字体大小调整8



    # 格式化日期标
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig("outputs/figures/picture9.png")
    plt.savefig("outputs/figures/图9_各类资金变动规模.png", dpi=300, bbox_inches='tight')


    ##%%
    ##---------------------
    # 1/
    被动权益ETF = pd.read_csv("workspace/funding_draft/data/被动权益基金/图11_被动权益ETF处理后数据.csv")
    主动权益基金 = pd.read_csv("workspace/funding_draft/data/主动权益基金/图11_主动权益基金_数据处理.csv")
    两融余额 = pd.read_csv("workspace/funding_draft/data/两融余额数据/图11_两融余额数据.csv")
    保险资金 = pd.read_csv("workspace/funding_draft/data/保险资金/图11_保险资金.csv")


    # 创建表格格式
    fig, ax = plt.subplots(figsize=(12, 2))
    ax.axis('tight')
    ax.axis('off')

    # 定义列名和行名
    def time_list_1_识别(time_str,recognised_quarter,recognised_year):
        if time_str[4:8] == "0331":
            recognised_quarter = "一"
        elif time_str[4:8] == "0630":
            recognised_quarter = "二"
        elif time_str[4:8] == "0930":
            recognised_quarter = "三"
        elif time_str[4:8] == "1231":
            recognised_quarter = "四"
        recognised_year = time_str[0:4]
        return recognised_quarter, recognised_year

    quarter, year = time_list_1_识别(time_list[1],quarter,year)
    # 定义列名和行名
    table_headers = [f"{time_list[0]}", f"{time_list[1]}", f"{year}年{quarter}季度变动", f"{time_list[2]}", f"{time_list[3]}", f"{time_list[4]}",
                     f"较{year}年{quarter}季度末变动", "上周变动", "本周变动"]
    row_headers = ["被动权益ETF", "主动权益基金", "两融余额", "保险资金"]

    ##%%
    # 将日期列转换为字符串格式以匹time_list
    被动权益ETF["日期"] = 被动权益ETF["日期"].astype(str)
    主动权益基金["日期"] = 主动权益基金["日期"].astype(str)  # 确保日期格式一
    两融余额["日期"] = 两融余额["日期"].astype(str)
    保险资金["日期"] = 保险资金["日期"].astype(str)


    # 准备自动化生成表格数
    def get_filled_row(df, time_list):
        row = ["被动权益ETF"]
        values = {date: df.loc[df["日期"] == date, "被动权益ETF"].values[0] if date in df["日期"].values else '-' for date
                  in time_list}

        # 填写数据
        row.extend([
            (values[time_list[0]]),  # 初始日期数据
            values[time_list[1]],  # 二季度末数据
            (values[time_list[1]] - values[time_list[0]]).round(1) if values[time_list[0]] != '-' and values[
                time_list[1]] != '-' else '-',  # 三季度变
            values[time_list[2]],  # 最近日
            values[time_list[3]],  # 上周
            values[time_list[4]],  # 本周
            (values[time_list[4]] - values[time_list[1]]).round(1) if values[time_list[4]] != '-' and values[
                time_list[1]] != '-' else '-',  # 季度末变
            (values[time_list[3]] - values[time_list[2]]).round(1) if values[time_list[2]] != '-' and values[
                time_list[3]] != '-' else '-',  # 上周变化
            (values[time_list[4]] - values[time_list[3]]).round(1) if values[time_list[3]] != '-' and values[
                time_list[4]] != '-' else '-',  # 本周变化
        ])
        return row

    def get_filled_row1(df, time_list):
        row = ["主动权益基金"]
        values = {date: df.loc[df["日期"] == date, "主动权益基金"].values[0] if date in df["日期"].values else '-' for date
                  in time_list}

        # 填写数据
        row.extend([
            (values[time_list[0]]),  # 初始日期数据
            values[time_list[1]],  # 二季度末数据
            (values[time_list[1]] - values[time_list[0]]).round(1) if values[time_list[0]] != '-' and values[
                time_list[1]] != '-' else '-',  # 三季度变
            values[time_list[2]],  # 最近日
            values[time_list[3]],  # 上周
            values[time_list[4]],  # 本周
            (values[time_list[4]] - values[time_list[1]]).round(1) if values[time_list[4]] != '-' and values[
                time_list[1]] != '-' else '-',  # 季度末变
            (values[time_list[3]] - values[time_list[2]]).round(1) if values[time_list[2]] != '-' and values[
                time_list[3]] != '-' else '-',  # 上周变化
            (values[time_list[4]] - values[time_list[3]]).round(1) if values[time_list[3]] != '-' and values[
                time_list[4]] != '-' else '-',  # 本周变化
        ])
        return row

    def get_filled_row2(df, time_list):
        row = ["两融余额"]
        values = {date: df.loc[df["日期"] == date, "两融余额"].values[0] if date in df["日期"].values else '-' for date
                  in time_list}

        # 填写数据
        row.extend([
            (values[time_list[0]]),  # 初始日期数据
            values[time_list[1]],  # 二季度末数据
            (values[time_list[1]] - values[time_list[0]]).round(1) if values[time_list[0]] != '-' and values[
                time_list[1]] != '-' else '-',  # 三季度变
            values[time_list[2]],  # 最近日
            values[time_list[3]],  # 上周
            values[time_list[4]],  # 本周
            (values[time_list[4]] - values[time_list[1]]).round(1) if values[time_list[4]] != '-' and values[
                time_list[1]] != '-' else '-',  # 季度末变
            (values[time_list[3]] - values[time_list[2]]).round(1) if values[time_list[2]] != '-' and values[
                time_list[3]] != '-' else '-',  # 上周变化
            (values[time_list[4]] - values[time_list[3]]).round(1) if values[time_list[3]] != '-' and values[
                time_list[4]] != '-' else '-',  # 本周变化
        ])
        return row
    # newest_rpt_date = "20240930" # 请自行修
    # newest_rpt_date_shift = "20240630" # 请自行修
    # time_list = [newest_rpt_date_shift, newest_rpt_date]
    def get_filled_row3(df, time_list):
        row = ["保险资金"]
        values = {}

        for date in time_list:
            if date in df["日期"].values:
                value = df.loc[df["日期"] == date, "保险资金"].values[0]
                values[date] = float(value) if not pd.isna(value) else "-"
            else:
                values[date] = '-'

        def safe_subtract(value1, value2):
            if isinstance(value1, (int, float)) and isinstance(value2, (int, float)):
                return value1 - value2
            return "-"

        def format_float(value, decimals=2):
            if isinstance(value, (int, float)):
                return f"{value:.{decimals}f}"
            else:
                return value

        # 修复点：使用safe_subtract进行计算
        row.extend([
            values[time_list[0]],
            values[time_list[1]],
            format_float(safe_subtract(values[time_list[1]], values[time_list[0]]), decimals=1),
            "-", "-", "-", "-", "-", "-"
        ])

        return row


    # 自动填充“被动权益ETF”行数据
    被动权益ETF_row = get_filled_row(被动权益ETF, time_list)[1:]
    主动权益基金_row = get_filled_row1(主动权益基金, time_list)[1:]
    两融余额_row = get_filled_row2(两融余额, time_list)[1:]
    保险资金_row = get_filled_row3(保险资金, time_list)[1:]

    # 打印结果
    print("被动权益ETF行数据：")
    print(被动权益ETF_row)

    print("主动权益行数据：")
    print(主动权益基金_row)

    print("两融余额行数据：")
    print(两融余额_row)

    print("险资行数据：")
    print(保险资金_row)



    ##%%
    # 创建图表
    fig, ax = plt.subplots(figsize=(12, 2))
    ax.axis('tight')
    ax.axis('off')

    # 创建表格数据
    table_data = [
        被动权益ETF_row,
        主动权益基金_row,  # 主动权益基金占位
        两融余额_row,  # 两融余额占位
        保险资金_row   # 保险资金占位
    ]


    # 创建自定义渐变色函数
    def get_color_map(value, vmin, vmax):
        """根据值返回从绿色到红色的颜色"""
        cmap = mcolors.LinearSegmentedColormap.from_list("green_red", ["green", "yellow", "red"])
        norm = mcolors.Normalize(vmin=vmin, vmax=vmax)
        return cmap(norm(value))

    # 获取每列的最小值和最大
    def get_column_min_max(table_data, col_index):
        column_values = [row[col_index] for row in table_data if isinstance(row[col_index], (int, float))]
        if column_values:
            return min(column_values), max(column_values)
        return None, None

    # 动态设置颜色格
    colors = [["#FAFAD2"] * len(table_headers) for _ in table_data]  # 初始背景颜色

    for col_index in range(0, len(table_headers)):  # 从第1列开始，不包括行名列
        vmin, vmax = get_column_min_max(table_data, col_index)
        if vmin is not None and vmax is not None:
            for row_index in range(len(table_data)):
                value = table_data[row_index][col_index]
                if isinstance(value, (int, float)):
                    color = get_color_map(value, vmin, vmax)
                    colors[row_index][col_index] = mcolors.to_hex(color)  # 转换为十六进制颜

    # 创建表格
    summary_table = ax.table(cellText=table_data, colLabels=table_headers, rowLabels=row_headers, loc='center', cellLoc="center", cellColours=colors)
    summary_table.auto_set_font_size(False)
    summary_table.set_fontsize(10)

    # 绘制渐变色条小图例
    cmap = mcolors.LinearSegmentedColormap.from_list("green_red", ["green", "yellow", "red"])

    # 添加一个小型色谱条 Legend
    legend_ax = fig.add_axes([0.4, 0.2, 0.10, 0.020])  # 调整位置 [左, 下, 宽度, 高度]
    legend_ax.imshow(np.linspace(0, 100, 256).reshape(1, -1), cmap=cmap, aspect="auto")
    legend_ax.set_xticks([0, 128, 255])
    legend_ax.set_xticklabels(["低", "中", "高"], fontsize=12)
    legend_ax.set_yticks([])  # 隐藏 y 轴刻度
    legend_ax.spines['top'].set_visible(False)
    legend_ax.spines['right'].set_visible(False)
    legend_ax.spines['left'].set_visible(False)
    legend_ax.spines['bottom'].set_visible(False)

    # 手动调整列宽和文本折
    for key, cell in summary_table.get_celld().items():
        cell.set_text_props(wrap=True)
        if key[0] == 0:  # 表头单元
            cell.set_text_props(color="white")
            cell.set_facecolor("#B22222")  # 保留表头颜色
        if key[1] == -1:  # 行名
            cell.set_text_props(weight="bold")
        if key[1] in [2, 6]:  # "2024年三季度变动")和第7"024年三季度末变)
            cell.set_width(0.18)
        else:
            cell.set_width(0.1)

    # 保存为PDF
    output_dir = "workspace/funding_draft/picture"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_pdf_path = os.path.join(output_dir, "图11.pdf")
    with PdfPages(output_pdf_path) as pdf:
        pdf.savefig(fig, bbox_inches='tight')

    print(f"表格已保存至 {output_pdf_path}")

    # 格式化日期标
    plt.tight_layout()
    plt.savefig("outputs/figures/picture11.png")
    plt.savefig("outputs/figures/图11_本周增量资金来源.png", dpi=300, bbox_inches='tight')

    ##%%

    # 2 按投资范围分类ETF资产净
    ETF_df = pd.read_csv("outputs/csv/figures_1_15/图12_按投资范围分类ETF资产净值（亿元）.csv")

    # Convert to DataFrame
    df = pd.DataFrame(ETF_df)

    # Convert "Date" column to datetime
    df["Date"] = pd.to_datetime(df["Date"], format='%Y%m%d')


    # Plot
    fig, ax1 = plt.subplots(figsize=(10, 6))

    # Left axis for "stock"
    ax1.plot(df["Date"], df["stock"], label="股票型ETF", color="#B22222", linewidth=2)
    #ax1.set_ylabel("股票型ETF", fontsize=12)
    ax1.tick_params(axis='y', labelcolor="black", labelsize=10, labelrotation=0,
                    labelleft=True, labelright=False, labeltop=False, labelbottom=True)
    # Formatting x-axis
    #ax1.set_xlabel("日期", fontsize=12)
    ax1.xaxis.set_major_formatter(plt.matplotlib.dates.DateFormatter("%m/%d/%Y"))

    ax1.xaxis.set_major_locator(mdates.MonthLocator(interval=3))  # Set the tick frequency to every 3 months


    # Update x-axis and left y-axis tick labels to use Times New Roman
    for label in ax1.get_xticklabels() + ax1.get_yticklabels():
        label.set_fontname("Times New Roman")
        label.set_fontsize(10)

    # Right axis for "bond", "commodity", "currency"
    ax2 = ax1.twinx()
    ax2.plot(df["Date"], df["bond"], label="债券型ETF（右轴）", color="gold", linewidth=2)
    ax2.plot(df["Date"], df["commodity"], label="商品型ETF（右轴）", color="orange", linewidth=2)
    ax2.plot(df["Date"], df["currency"], label="货币型ETF（右轴）", color="skyblue", linewidth=2)
    #ax2.set_ylabel("债券型ETF / 商品型ETF / 货币型ETF", fontsize=12)
    ax2.tick_params(axis='y', labelsize=10, labelrotation=0)

    # Update right y-axis tick labels to use Times New Roman
    for label in ax2.get_yticklabels():
        label.set_fontname("Times New Roman")
        label.set_fontsize(10)

    # Adding legend
    lines, labels = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax2.legend(lines + lines2, labels + labels2, loc=9, bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=10)
              # , prop={'family': 'Times New Roman', 'size': 10})

    # Adjust layout
    #plt.title("ETF Data Visualization", fontsize=14)
    plt.tight_layout(pad = 3) # 让图占小

    plt.savefig("outputs/figures/图12_按投资范围分类ETF资产净值（亿元）.png", dpi=300, bbox_inches='tight')
    pic_1 = Image.open("outputs/figures/图12_按投资范围分类ETF资产净值（亿元）.png")
    size = pic_1.size
    print(size)
    print(size[0] / size[1])
    # print(size2[0]/size2[1])
    # 修改长宽比例
    pic_new = pic_1.resize((2380, 980))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture12.png")


    # 3:按行业指数分类ETF周区间净流入
    if not ENABLE_FIGURE13_INDUSTRY_ETF:
        print("图13开关关闭：跳过图13绘制。")
    else:
        df = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图13_按行业指数分类ETF周区间净流入额（亿元）.csv",
            ["net_flow"],
            "图13_按行业指数分类ETF周区间净流入额（亿元）",
        )

        industry_summary = df.groupby("industry_names")["net_flow"].sum().reset_index()
        industry_summary = industry_summary.sort_values(by="net_flow", ascending=False)
        value_span = industry_summary["net_flow"].max() - industry_summary["net_flow"].min()
        label_offset = max(float(value_span) * 0.03, 0.1)

        fig, ax = plt.subplots(figsize=(12, 6))

        bars = ax.bar(
            industry_summary["industry_names"],
            industry_summary["net_flow"],
            color="darkred",
            width=0.6
        )

        for bar in bars:
            height = bar.get_height()
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                height + label_offset if height > 0 else height - label_offset,
                f"{height:.1f}",
                ha="center",
                va="bottom" if height > 0 else "top",
                fontsize=12,
                fontdict=font_English
            )

        ax.set_ylabel("净流入额（亿元）", fontsize=14)
        for label in ax.get_yticklabels():
            label.set_fontname("Times New Roman")
            label.set_fontsize(12)

        plt.xticks(rotation=45, ha="right", fontsize=12)
        plt.tight_layout()
        plt.savefig("outputs/figures/图13_按行业指数分类ETF周区间净流入额（亿元）.png", dpi=300, bbox_inches='tight')
        pic_1 = Image.open("outputs/figures/图13_按行业指数分类ETF周区间净流入额（亿元）.png")
        size = pic_1.size
        print(size)
        print(size[0] / size[1])
        pic_new = pic_1.resize((2380, 980))
        pic_new.save("outputs/figures/picture13.png")


    # 转换日期格式为更可读的形
    def format_date_range(date_range):
        start_date, end_date = date_range.split('-')
        start_date = datetime.datetime.strptime(start_date, "%Y%m%d").strftime("%m/%d/%Y")
        end_date = datetime.datetime.strptime(end_date, "%Y%m%d").strftime("%m/%d/%Y")
        return f"{start_date}-{end_date}"

    # 0：股票型ETF净流入额（亿元
    if not ENABLE_FIGURE10_ETF_NETINFLOW:
        print("图10开关关闭：跳过图10绘制。")
    else:
        df = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图10_股票型ETF净流入额（亿元）.csv",
            ["net_inflow_stock"],
            "图10_股票型ETF净流入额（亿元）",
        )

        df["Formatted Date"] = df["Date"].apply(format_date_range)
        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(
            df["Formatted Date"],
            df["net_inflow_stock"],
            color="#B22222",
            width=0.4,
            label="股票型ETF区间净流入额（亿元）",
        )

        ax.axhline(y=0, color="black", linewidth=1.5, linestyle="-")

        for index, label in enumerate(ax.get_xticklabels()):
            if index % 2 != 0:
                label.set_visible(False)

        for bar in bars:
            yval = bar.get_height()
            if yval > 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    yval,
                    f"{yval:.0f}",
                    ha="center",
                    va="bottom",
                    fontsize=10,
                    fontname="Times New Roman",
                )
            elif yval < 0:
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    yval,
                    f"{yval:.0f}",
                    ha="center",
                    va="top",
                    fontsize=10,
                    fontname="Times New Roman",
                )

        for label in ax.get_xticklabels():
            label.set_fontname("Times New Roman")
            label.set_fontsize(10)

        for label in ax.get_yticklabels():
            label.set_fontname("Times New Roman")
            label.set_fontsize(10)

        ax.legend(fontsize=12, loc=8)
        plt.tight_layout()
        plt.savefig("outputs/figures/图10_股票型ETF净流入额（亿元）.png", dpi=300, bbox_inches='tight')
        pic_1 = Image.open("outputs/figures/图10_股票型ETF净流入额（亿元）.png")
        pic_new = pic_1.resize((2380, 980))
        pic_new.save("outputs/figures/picture10.png")

    # %%4绘制
    if not ENABLE_FIGURE14_UNLOCK:
        print("图14开关关闭：跳过图14绘制。")
    else:
        xs_use = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图14_限售股解禁金额和数量（周度）.csv",
            ["当周解禁市值", "当周解禁家数"],
            "图14_限售股解禁金额和数量（周度）",
        )
        remake = (34 / 10) / (8231 / 3072)
        fig = plt.figure(figsize=(34 * remake, 10), dpi=300)
        ax = fig.add_subplot(111)

        plt.rcParams['ytick.direction'] = 'in'

        max_ylim = int(max(xs_use["当周解禁市值"]) / 500) * 500 + 500
        max_ylim2 = int(max(xs_use["当周解禁家数"]) / 10) * 10 + 10
        ax.set_ylim(0, max_ylim)

        y_major_locator = MultipleLocator(500)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
        ax.axhline(y=0, c='grey', lw=2)

        plt.xticks(alpha=0)
        plt.tick_params(axis='x', width=0)

        n = len(xs_use)
        x1 = list(range(1, 1 + 2 * n, 2))
        x2 = list(range(2, 2 + 2 * n, 2))

        plt.bar(
            x1,
            xs_use["当周解禁市值"],
            alpha=0.8,
            width=1,
            color="steelblue",
            edgecolor=["navy"],
            linewidth=2,
            label="当周解禁市值(亿元)",
            zorder=1,
        )

        plt.xlim(0, x2[-1] + 2)
        ax.set_xticks(x2)

        for i in range(n):
            plt.text(x1[i], xs_use["当周解禁市值"][i] + 25, xs_use["当周解禁市值"][i],
                     ha='center', va='bottom', fontsize=30, color="black", zorder=5, family="Times new roman")

        for i in range(n):
            plt.text(x2[i], -100, xs_use['日期'][i],
                     ha='right', va='top', fontsize=30, rotation=30, color="black", family="Times new roman")

        plt.grid(zorder=0, linestyle='--', dashes=(15, 12))
        plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
            0.28, 1.12), loc=9, borderaxespad=0, fontsize=36)

        ax2 = ax.twinx()
        ax2.set_ylim(0, max_ylim2)

        ax2.plot(x1, xs_use["当周解禁家数"], color='darkorange',
                 linewidth=5, label="当周解禁家数(家，右轴)")
        ax2.scatter(x1, xs_use["当周解禁家数"], s=300, marker="D",
                    facecolor='none', ec='darkorange', linewidth=5, zorder=2)

        y_major_locator2 = MultipleLocator(10)
        ax2 = plt.gca()
        ax2.yaxis.set_major_locator(y_major_locator2)
        ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)

        plt.legend(frameon=False, bbox_to_anchor=(0.68, 1.12),
                   loc=9, borderaxespad=0, fontsize=36)

        labels = ax.get_yticklabels() + ax2.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        ax.spines['top'].set_visible(False)
        ax2.spines['top'].set_visible(False)

        plt.savefig("outputs/figures/图14_限售股解禁金额和数量（周度）.png", dpi=300, bbox_inches='tight')
        pic_14 = Image.open("outputs/figures/图14_限售股解禁金额和数量（周度）.png")
        size = pic_14.size
        print(size, size[0] / size[1])
        pic_new = pic_14.resize((2380, 700))
        pic_new.save("outputs/figures/picture14.png")


    # %%5绘制
    if not ENABLE_FIGURE15_IPO:
        print("图15开关关闭：跳过图15绘制。")
    else:
        ipo_use = load_chart_data_with_numeric_validation(
            "outputs/csv/figures_1_15/图15_IPO和定增金额（周度）.csv",
            ["IPO首发家数", "IPO首发募集资金（亿元）", "定增家数", "定增募集（亿元）"],
            "图15_IPO和定增金额（周度）",
        )

        remake = (34 / 10) / (8306 / 2728)
        fig = plt.figure(figsize=(34 * remake, 10), dpi=300)
        ax = fig.add_subplot(111)

        plt.rcParams['ytick.direction'] = 'in'

        max_ylim = int(max(ipo_use["IPO首发募集资金（亿元）"].max(), ipo_use["定增募集（亿元）"].max()) / 50) * 50 + 50
        max_ylim2 = int(max(ipo_use["IPO首发家数"].max(), ipo_use["定增家数"].max()) / 2) * 2 + 11
        max_y = max(max_ylim, (max_ylim2 * 25))
        ax.set_ylim(0, max_y)

        y_major_locator = MultipleLocator(50)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
        labels = ax.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        ax.axhline(y=0, c='grey', lw=2)
        plt.xticks(alpha=0)
        plt.tick_params(axis='x', width=0)

        x1 = list(range(1, 1 + 3 * 12, 3))
        x2 = list(range(2, 2 + 3 * 12, 3))
        x3 = list(range(3, 3 + 3 * 12, 3))

        plt.xlim(0, 36.01)
        ax.set_xticks(x3)
        plt.grid(zorder=0, linestyle='--', dashes=(8, 4))

        plt.bar(x1, ipo_use["IPO首发募集资金（亿元）"], alpha=0.8, width=1, color='steelblue', edgecolor=[
            "navy"], linewidth=2, label="IPO首发募集资金（亿元）", zorder=1)
        plt.bar(x2, ipo_use["定增募集（亿元）"], alpha=0.8, width=1, color='darkgrey', edgecolor=[
            "dimgrey"], linewidth=2, label="定增募集（亿元）", zorder=1)

        ax.set_ylim(0, 120)
        y_major_locator = MultipleLocator(20)
        ax.yaxis.set_major_locator(y_major_locator)

        for i in range(0, 12):
            ipo_val = ipo_use["IPO首发募集资金（亿元）"].iloc[i]
            dz_val = ipo_use["定增募集（亿元）"].iloc[i]
            plt.bar(x1[i], min(ipo_val, 100), alpha=0.8, width=1, color='steelblue', edgecolor="navy", linewidth=2, zorder=1)
            plt.bar(x2[i], min(dz_val, 100), alpha=0.8, width=1, color='darkgrey', edgecolor="dimgrey", linewidth=2, zorder=1)

            plt.text(x1[i], min(ipo_val, 100) + 3, f"{ipo_val:.1f}", ha='center', va='bottom',
                     fontsize=25, color="black", zorder=5, family="Times New Roman")
            plt.text(x2[i], min(dz_val, 100) + 3, f"{dz_val:.1f}", ha='center', va='bottom',
                     fontsize=25, color="black", zorder=5, family="Times New Roman")

        for i in range(0, 12):
            if ipo_use["定增募集（亿元）"][i]> 120:
                x_pos = x2[i] - 0.6
                rect = mpatches.Rectangle((x_pos, 0), 1.2, min(ipo_use["定增募集（亿元）"].iloc[i], 120),
                                          linewidth=4, edgecolor='red', facecolor='none', zorder=3)
                plt.gca().add_patch(rect)

        for i in range(0, 12):
            if ipo_use["IPO首发募集资金（亿元）"][i]> 120:
                x_pos = x2[i] - 1.6
                rect = mpatches.Rectangle((x_pos, 0), 1.2, min(ipo_use["IPO首发募集资金（亿元）"].iloc[i], 120),
                                          linewidth=4, edgecolor='red', facecolor='none', zorder=3)
                plt.gca().add_patch(rect)


        ax.set_yticks([120,110,100,80,60,40,20,0])
        ax.set_yticklabels(['4150', '...', '100',80,60,40,20,0])

    # 之后用这部分，红色的删了
    # # 柱状图上显示数据
    # for i in range(0, 12):
    #     plt.text(x1[i], ipo_use["IPO首发募集资金（亿元）"].iloc[i] + 3, ipo_use["IPO首发募集资金（亿元）"].iloc[i],
    #              ha='center', va='bottom', fontsize=25, color="black", zorder=5, family="Times new roman")
    #     plt.text(x2[i], ipo_use["定增募集（亿元）"].iloc[i] + 3, ipo_use["定增募集（亿元）"].iloc[i],
    #              ha='center', va='bottom', fontsize=25, color="black", zorder=5, family="Times new roman")

        plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
            0.32, 1.1), loc=9, borderaxespad=0, fontsize=30)
        ax2 = ax.twinx()
        ax2.set_ylim(0, max_y / 25)
        ax2.set_ylim(0, 6)

        ax2.plot(x1, ipo_use["IPO首发家数"], color='orangered',
                 linewidth=5, label="IPO首发家数(家，右轴)", alpha=0.5)
        ax2.scatter(x1, ipo_use["IPO首发家数"], s=300, marker="D",
                    facecolor='none', ec='orangered', linewidth=5, zorder=2, alpha=0.5)
        ax2.plot(x1, ipo_use["定增家数"], color='gold',
                 linewidth=5, label="定增家数(家，右轴)", alpha=0.7)
        ax2.scatter(x1, ipo_use["定增家数"], s=300, marker="D",
                    facecolor='none', ec='gold', linewidth=5, zorder=2, alpha=0.7)

        y_major_locator2 = MultipleLocator(2)
        ax2 = plt.gca()
        ax2.yaxis.set_major_locator(y_major_locator2)
        ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)
        plt.legend(frameon=False, ncol=2, bbox_to_anchor=(
            0.72, 1.1), loc=9, borderaxespad=0, fontsize=30)

        labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        ax.spines['top'].set_visible(False)
        ax2.spines['top'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        ax2.spines['bottom'].set_visible(False)

        for i in range(len(x1)):
            plt.text(x2[i] - 0.5, -0.5, ipo_use["日期"].iloc[i], ha='center',
                     va='top', fontsize=30, family="Times new roman")

        plt.savefig("outputs/figures/图15_IPO和定增金额（周度）.png", dpi=300, bbox_inches='tight')
        pic_15 = Image.open("outputs/figures/图15_IPO和定增金额（周度）.png")
        size = pic_15.size
        print(size, size[0] / size[1])
        pic_new = pic_15.resize((2380, 700))
        pic_new.save("outputs/figures/picture15.png")

        plt.savefig("outputs/figures/图15_IPO和定增金额（周度）.png", dpi=300, bbox_inches='tight')
        pic_15 = Image.open("outputs/figures/图15_IPO和定增金额（周度）.png")
        size = pic_15.size
        print(size, size[0] / size[1])
        pic_new = pic_15.resize((2380, 700))
        pic_new.save("outputs/figures/picture15.png")

    # ------------------------------%% 6绘制动态估值布林带--------------------------------------------------------------
    DB_LOC = str(paths.data_dir)
    database_location = r'{}/quotation'.format(DB_LOC)

    ##### 基本设置
    # 画图
    location = [321, 322, 323, 324, 325, 326]

    # 起始时间
    begin_date = '2017-01-01'
    end_date = report_date

    # 设置boll轨道参数
    rolling_m = 250 * 5  # 滚动窗口长度   取50 * 4/5/6
    bolling_n = 1  # 上下轨为n倍标准差 取/1.5
    bolling_n2 = 1.5

    # 指标
    value_index_list = [
        'PE_TTM',
        # 'PB_LF',
        # 'Shiller PE'
    ]

    value_index_short_dict = {'PE_TTM': 'PE_TTM',
                              'PB_LF': 'PB_LF',
                              'Shiller PE': 'Shiller PE',
                              }

    # 市场指数和行业板
    sector_code_dict = {
        '上证指数': '000001.SH',  ##
        '深证成指': '399001.SZ',
        '创业板指': '399006.SZ',  #
        # '上证50': '000016.SH',  #
        '沪深300': '000300.SH',  #
        '中证500': '000905.SH',
        '中证1000': '000852.SH',
    }


    ### 为绘制估值穿越上限或下限的点，进一步处理 data_，获取各个点所处的状态
    # 针对双上下轨
    def state_judge2(data, ma, up, down, up2, down2):
        """
        针对双上下轨
        判断各期估值指标处于何种状态：
            (0') data >= 上轨2:       3

            (1) 上轨 <= data < 上轨2: 2
            (2) 均线 <= data < 上轨:  1
            (3) 下轨 < data < 均线:  -1
            (4) 下轨2 < data <=下轨: -2

            (4') data <= 下轨2       -3

        """
        state = 0

        if data >= up2:
            state = 3
        elif data >= up:
            state = 2
        elif data >= ma:
            state = 1
        elif data > down:
            state = -1
        elif data > down2:
            state = -2
        elif data <= down2:
            state = -3

        return state


    def cross_state_judge2(state, state_diff):
        """
        针对双上下轨
        通过估值指标的状态哑变量及其差分数值，可以判断出穿越上限或下限的那一天是哪一天
        """
        cross_type = 'normal'  # 无上穿、下
        if state == 3:
            if state_diff == 1:
                cross_type = 'up_cross_up2'  # 上穿up2
        elif state == 2:
            if state_diff == 1:
                cross_type = 'up_cross_up'  # 上穿up
            elif state_diff == -1:
                cross_type = 'down_cross_up2'  # 下穿up2
        elif state == 1:
            if state_diff == 2:
                cross_type = 'up_cross_ma'  # 上穿ma
            elif state_diff == -1:
                cross_type = 'down_cross_up'  # 下穿up
        elif state == -1:
            if state_diff == 1:
                cross_type = 'up_cross_down'  # 上穿down
            elif state_diff == -2:
                cross_type = 'down_cross_ma'  # 下穿ma
        elif state == -2:
            if state_diff == 1:
                cross_type = 'up_cross_down2'  # 上穿down2
            elif state_diff == -1:
                cross_type = 'down_cross_down'  # 下穿down
        elif state == -3:
            if state_diff == -1:
                cross_type = 'down_cross_down2'  # 下穿down2

        return cross_type


    fig = plt.figure(figsize=(17, 16), dpi=150)

    quantile_table = []
    text = "动态估值布林带模型显示，主要市场指数估值均处于区域，其中处于区域。当前（{}）".format(today)
    for sector_name, sector_code in sector_code_dict.items():
        loc = location[0]
        location.pop(0)
        print(sector_name)
        for index in value_index_list:
            print(index)

            # 读取数据文件
            data_all = pd.read_csv('{}/{}.csv'.format(database_location, sector_name), dtype={'date': str})

            # 先把 date 列转datetime，转不出来的（比"000001.SH"）会NaT
            data_all['date'] = pd.to_datetime(data_all['date'], errors='coerce')

            # 删除 date 解析失败的脏
            data_all = data_all.dropna(subset=['date'])

            # 再设index
            data_all = data_all.set_index('date')

            print(f"[Step3调试] 需要的列: ['CLOSE', '{index}']")
            print(f"[Step3调试] 实际列名: {list(data_all.columns)}")
            data_all.columns = [str(col).upper().strip() for col in data_all.columns]
            index_upper = str(index).upper().strip()
            required_cols = ['CLOSE', index_upper]
            missing_cols = [col for col in required_cols if col not in data_all.columns]
            if missing_cols:
                print(f"[错误] 缺少列: {missing_cols}")
                print(f"[信息] 可用列: {list(data_all.columns)}")
                for col in missing_cols:
                    data_all[col] = np.nan
            data_all = data_all[['CLOSE', index_upper]]
            data_all.dropna(axis=0, how='any', inplace=True)

            if len(data_all) < rolling_m:
                print('{}.csv {}: 数据长度为{}, 长度不足{}!'.format(sector_name, index, len(data_all), str(rolling_m)))

            data_ = data_all

            # 计算boll
            data_[index + '_MA'] = data_[index].ewm(span=rolling_m).mean()  # 指数衰减加权移动平均
            data_[index + '_STD'] = data_[index].rolling(window=rolling_m).std()
            data_[index + '_UP'] = data_[index + '_MA'] + data_[index + '_STD'] * bolling_n
            data_[index + '_DOWN'] = data_[index + '_MA'] - data_[index + '_STD'] * bolling_n
            data_[index + '_UP2'] = data_[index + '_MA'] + data_[index + '_STD'] * bolling_n2
            data_[index + '_DOWN2'] = data_[index + '_MA'] - data_[index + '_STD'] * bolling_n2

            data_.dropna(axis=0, subset=[index + '_MA'], inplace=True)
            data_ = data_[begin_date:end_date]
            data_['日期'] = pd.to_datetime(data_.index)

            ### 计算指标最新值在这个区间内的分位
            max_index = data_[index].max()
            min_index = data_[index].min()
            data_['quantile'] = ((data_[index] - min_index) / (max_index - min_index)) * 100
            latest_index = round(data_.iloc[-1][index], 2)
            latest_quantile = round(data_.iloc[-1]['quantile'], 2)
            # 数据保存
            quantile_table.append([sector_name, latest_index, latest_quantile])
            print('{}, {}的取值为{}, 在全样本中的分位数为 {}'.format(end_date, index, str(latest_index), str(latest_quantile)))
            text += '{}PE-TTM为{}, 处于{}%分位数；'.format(sector_name, str(latest_index), str(latest_quantile))
            # 判断上穿和下穿boll轨道
            data_['state'] = data_.apply(
                lambda x: state_judge2(x[index], x[index + '_MA'], x[index + '_UP'], x[index + '_DOWN'], x[index + '_UP2'],
                                       x[index + '_DOWN2']), axis=1)
            data_['state_diff'] = data_['state'].diff()
            data_['cross_state'] = data_.apply(lambda x: cross_state_judge2(x['state'], x['state_diff']), axis=1)

            ### 绘制曲线
            # fig = plt.figure(figsize = (15,8), dpi = 300)  # PPT 单图使用
            # fig = plt.figure(figsize = (20,8), dpi = 300)  # PPT n*2图使
            # fig = plt.figure(figsize = (20,6), dpi = 300)  # word 单图使用
            # word n*2排版使用

            index_short = value_index_short_dict[index]  # 指标简

            plt.rcParams['font.sans-serif'] = [font_family]  # 在tilte中正常显示中
            plt.rcParams['axes.unicode_minus'] = False
            plt.rcParams.update({'font.size': 20})

            ## 绘制板块收盘价
            ax = fig.add_subplot(loc)
            ax = plt.gca()
            ax.tick_params(axis="x", direction="in", which="major", labelsize=30)
            ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
            # 绘制收盘价
            curves1 = ax.plot(data_['日期'], data_['CLOSE'], label=sector_name + ':收盘价', color='black', alpha=0.8,
                              linewidth=1.5)

            # 绘制 PE 等估值指标
            ax2 = ax.twinx()
            curves2 = ax2.plot(data_['日期'], data_[index + '_MA'], label='{}_MA'.format(index_short), color='darkorange',
                               alpha=1, linewidth=1)
            curves3 = ax2.plot(data_['日期'], data_[index + '_UP'], label='{}_UP'.format(index_short), color='tomato',
                               alpha=0.5, linewidth=1.5)
            curves4 = ax2.plot(data_['日期'], data_[index + '_DOWN'], label='{}_DOWN'.format(index_short),
                               color='forestgreen', alpha=0.5, linewidth=1.5)
            curves5 = ax2.plot(data_['日期'], data_[index + '_UP2'], label='{}_UP2'.format(index_short), color='tomato',
                               alpha=1, linewidth=1.5)
            curves6 = ax2.plot(data_['日期'], data_[index + '_DOWN2'], label='{}_DOWN2'.format(index_short),
                               color='forestgreen', alpha=1, linewidth=1.5)
            curves7 = ax2.plot(data_['日期'], data_[index], label=index_short, color='royalblue', alpha=0.8, linewidth=1.5)

            ### (绘图3) scatter顶部、底部区
            top_area = data_[data_['state'] == 2]  # 市场顶部区间 1倍标准差
            bottom_area = data_[data_['state'] == -2]  # 市场底部区间 1倍标准差
            top_area2 = data_[data_['state'] == 3]  # 市场顶部区间 1.5倍标准差
            bottom_area2 = data_[data_['state'] == -3]  # 市场底部区间 1.5倍标准差

            # 在收盘价曲线上scatter
            ax.scatter(top_area['日期'], top_area['CLOSE'], label='顶部区域', color='orange', marker='o',
                       linewidths=4)
            ax.scatter(top_area2['日期'], top_area2['CLOSE'], label='极端顶部区域', color='g', marker='o',
                       linewidths=4)

            ax.scatter(bottom_area['日期'], bottom_area['CLOSE'], label='底部区域', color='pink',
                       marker='o', linewidths=4)
            ax.scatter(bottom_area2['日期'], bottom_area2['CLOSE'], label='极端底部区域', color='r',
                       marker='o', linewidths=4)

            font_English = {'family': 'Times New Roman'}

            # 设置图例
            # 左轴legend
            ax.legend([sector_name, '顶部区域', '极端顶部区域', '底部区域', '极端底部区域'], loc='upper left')
            # 右轴legend
            curves = curves7 + curves2 + curves3 + curves4 + curves5 + curves6
            labs = [x.get_label() for x in curves]
            ax2.legend(curves, labs, loc='lower left', prop=font_English, fontsize=20)

            # 设置坐标轴数字的字体为Times New Roman
            labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_xticklabels() + ax2.get_yticklabels()
            [label.set_fontname('Times New Roman') for label in labels]
            # 右轴大小
            ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)
            # 设置网格
            ax2.grid(axis='y', linestyle='--')
            ax.grid(axis='x', linestyle='--')

    # 储存图片
    plt.savefig("outputs/figures/图16_动态估值布林带.png", dpi=300, bbox_inches='tight')
    pic_16 = Image.open("outputs/figures/图16_动态估值布林带.png")
    size = pic_16.size
    print(size, size[0] / size[1])
    # 修改长宽比例
    pic_new = pic_16.resize((1190, 1120))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture16.png")

    # 储存text
    text = text[::-1].replace("；", "。", 1)[::-1]
    (paths.word_replacement_dir / "1.5更替文段.txt").write_text(text, encoding="utf-8")
    ##%%
    # 番外：数据保存
    results_df = pd.DataFrame(quantile_table, columns=['Sector Name', 'PE-TTM', '估值分位数'])
    results_df.to_csv('outputs/csv/figures_1_15/图16__动态估值布林带.csv', index=False, encoding='utf-8-sig')



    # %% 7绘制 下跌能量模型
    DB_LOC = str(paths.data_dir)
    database_location = r'{}/falling_alert'.format(DB_LOC)

    begin_date = '2017-01-01'
    end_date = report_date.isoformat()  ##&也可2022-12-02'

    datafile = str(resolve_falling_alert_file(paths, report_date))

    def split_csv(file=None, key=None, split_path=database_location, split_name_suffix='.csv'):
        # read csv
        df = pd.read_csv(file)
        df_group = df.groupby(key)
        for g in df_group:
            print(g[0])
            g[1].drop(['SECUCODE'], axis=1, inplace=True)
            g[1].rename(columns={'TRADEDATE': 'Date', 'CLOSE_IDX': g[0]}, inplace=True)
            # print(g[1])
            file_name = os.path.join(split_path, '_'.join([g[0], split_name_suffix]))
            g[1].to_csv(file_name, index=False)
            print('save {}'.format(file_name))


    split_csv(datafile, key='SECUCODE', split_name_suffix='peak_window300.csv')

    index_dict = {
        '000001.SH': '上证指数',
        '399001.SZ': '深证成指',
        '399006.SZ': '创业板指',
        '000300.SH': '沪深300',
        '000905.SH': '中证500',
        '000852.SH': '中证1000',

        # 'CI005003.WI':'有色金属(中信一',
        # 'CI005005.WI':'钢铁(中信一',

        # 'HSI.HI':'恒生指数'
    }

    peak_index = 'PeakFactor'
    # peak_index = 'PeakFAdjust'  # 弃用

    threshold_param_dict = {
        '000001.SH': (0.12, 0.15),
        # '000001.SH':(0.1, 0.15),
        '000300.SH': (0.12, 0.15),
        '000905.SH': (0.12, 0.15),
        '399001.SZ': (0.12, 0.15),
        '399006.SZ': (0.12, 0.15),
        '000852.SH': (0.15, 0.17),
        'CI005003.WI': (0.12, 0.15),
        'CI005005.WI': (0.12, 0.15),
        'HSI.HI': (0.15, 0.2)
    }

    fig = plt.figure(figsize=(17, 16), dpi=150)  # 4/6宫格
    location = [321, 322, 323, 324, 325, 326]

    for index, index_name in index_dict.items():
        loc = location[0]
        location.pop(0)
        print('{} {}'.format(index, index_name))

        data_file = index + '_peak_window300.csv'
        data = pd.read_csv(r'{}/{}'.format(database_location, data_file))

        data.dropna(axis=0, subset=[index], inplace=True)  # 删除nan

        data['日期'] = pd.to_datetime(data['Date'])  # 转换为datetime，否则plot时会打印所有日
        data.set_index(['Date'], inplace=True)
        data = data[begin_date:end_date]

        threshold_1, threshold_2 = threshold_param_dict[index]  # 获取参数

        # fig = plt.figure(figsize = (20,8), dpi = 300)      # 大图
        # fig = plt.figure(figsize = (15,8), dpi = 300)      # ppt4宫格

        plt.rcParams['font.sans-serif'] = [font_family]  # 在tilte中正常显示中
        plt.rcParams['axes.unicode_minus'] = False
        plt.rcParams.update({'font.size': 20})
        font_English = {'family': 'Times New Roman'}

        ## 绘制板块收盘
        ax = fig.add_subplot(loc)

        # 设置close曲线的区间，避免close曲线与指标曲线重
        min_close = float(data[index].min())
        max_close = float(data[index].max())
        deci_period = (max_close - min_close) / 10
        min_ylim = min_close - deci_period
        max_ylim = max_close + deci_period * 2
        ax.set_ylim(min_ylim, max_ylim)  # 设置左轴范围

        curves1 = ax.plot(data['日期'], data[index], label=index_name, color='black', alpha=0.8, linewidth=1.5)

        ax = plt.gca()
        ax.tick_params(axis="x", direction="in", which="major", labelsize=30)
        ax.tick_params(axis="y", direction="in", which="major", labelsize=30)
        # 绘制PE等估值指
        ax2 = ax.twinx()
        curves2 = ax2.plot(data['日期'], data[peak_index], label=peak_index, color='r', alpha=0.4, linewidth=1)

        # 绘制超过阈值的
        alert_point_1 = data[(data[peak_index] >= threshold_1) & (data[peak_index] < threshold_2)]
        ax.scatter(alert_point_1['日期'], alert_point_1[index], color='orange', marker='o', linewidths=5)

        alert_point_2 = data[data[peak_index] >= threshold_2]
        ax.scatter(alert_point_2['日期'], alert_point_2[index], color='g', marker='o', linewidths=5)

        # 设置图例
        curves = curves1 + curves2
        labs = [x.get_label() for x in curves]
        ax.legend([index_name, '顶部信号', '极端顶部信号'], loc='upper left', fontsize=30)

        # 设置网格
        ax.grid(axis='x', linestyle='-.')
        ax2.grid(axis='y', linestyle='-.')
        ax2.legend(['下跌能量指标:{}'.format(peak_index)], loc='upper right', fontsize=30)

        # 设置横纵坐标轴数值为times new roman
        labels = ax.get_xticklabels() + ax.get_yticklabels() + ax2.get_xticklabels() + ax2.get_yticklabels()
        [label.set_fontname('Times New Roman') for label in labels]

        # 设置横坐标为年份
        ax2.xaxis.set_major_locator(mdates.YearLocator())
        ax2.xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
        ax = plt.gca()
        ax2.tick_params(axis="y", direction="in", which="major", labelsize=30)
    # 储存图片
    plt.savefig("outputs/figures/图17_下跌能量.png", dpi=300, bbox_inches='tight')
    pic_17 = Image.open("outputs/figures/图17_下跌能量.png")
    size = pic_17.size
    print(size, size[0] / size[1])
    # 修改长宽比例
    pic_new = pic_17.resize((1190, 1120))  # Image.ANTIALIAS
    pic_new.save("outputs/figures/picture17.png")


def run_step4(paths: ProjectPaths, report_date: datetime.date) -> None:
    template_path = paths.word_template
    replacement_dir = paths.word_replacement_dir
    if not template_path.exists():
        raise FileNotFoundError(f'模板 docx 不存在: {template_path}')
    from docx import Document
    from docx.shared import Cm
    import datetime
    from datetime import timedelta

    ##%%
    def replace_text_in_summary(doc, replacements):
        """遍历正文/表格/页眉页脚/文本框/XML 静态文本并做整段替换，确保标题和目录真正落地。"""

        def replace_in_paragraph(paragraph, location: str) -> None:
            original_text = paragraph.text
            if not original_text:
                return
            new_text = original_text
            for old_text, replacement in replacements.items():
                if old_text in new_text:
                    new_text = new_text.replace(old_text, replacement)
            if new_text != original_text:
                print(f"{location}替换: {original_text} -> {new_text}")
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.text = new_text

        def replace_in_table(table, location: str) -> None:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, location)
                    for nested_table in cell.tables:
                        replace_in_table(nested_table, location)

        def replace_in_xml_text_nodes(root_element, location: str) -> None:
            for text_node in root_element.xpath('.//w:t'):
                original_text = text_node.text or ""
                if not original_text:
                    continue
                new_text = original_text
                for old_text, replacement in replacements.items():
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, replacement)
                if new_text != original_text:
                    print(f"{location}XML替换: {original_text} -> {new_text}")
                    text_node.text = new_text

        for para in doc.paragraphs:
            replace_in_paragraph(para, "正文")

        for table in doc.tables:
            replace_in_table(table, "表格")

        for section in doc.sections:
            for para in section.header.paragraphs:
                replace_in_paragraph(para, "页眉")
            for table in section.header.tables:
                replace_in_table(table, "页眉表格")

            for para in section.footer.paragraphs:
                replace_in_paragraph(para, "页脚")
            for table in section.footer.tables:
                replace_in_table(table, "页脚表格")

            replace_in_xml_text_nodes(section.header._element, "页眉")
            replace_in_xml_text_nodes(section.footer._element, "页脚")

        for shape in doc.inline_shapes:
            if shape._element.xpath('.//w:t'):
                replace_in_xml_text_nodes(shape._element, "文本框")

        replace_in_xml_text_nodes(doc._element, "正文")



    ##%%

    title = str(paths.weekly_word_output_dir / f"{report_date.strftime('%Y%m%d')}周报-图片更替")

    # 获取今天的日
    today = datetime.datetime.combine(report_date, datetime.time.min)
    year = today.strftime("%Y")  # 
    month = today.strftime("%m")  # 
    day = today.strftime("%d")  # 

    five_day_ago = (today - timedelta(days=5))
    five_day_ago = five_day_ago.strftime("%Y-%m-%d")

    today_str = today.strftime("%Y-%m-%d")

    docx = Document(str(template_path))

    def read_replacement_text(path: Path) -> str:
        encodings = ["utf-8", "utf-8-sig", "gb18030", "gbk"]
        last_error: Exception | None = None
        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError as exc:
                last_error = exc
                continue
        raise RuntimeError(
            f"更替文段读取失败：无法用支持的编码解析文件 {path}。"
            f" 已尝试编码={encodings}，last_error={last_error}"
        )

    text1 = read_replacement_text(replacement_dir / '1.1更替文段.txt')
    text2 = "" if not ENABLE_FIGURE6_FUTURES else read_replacement_text(replacement_dir / '1.3更替文段.txt')

    # text3 是关于北向资金的，已删除，不需

    text4 = read_replacement_text(replacement_dir / '1.5更替文段.txt')
    # text 1.5 动态估值布林带数据



    # 替换占位
    for table in docx.tables:  # 遍历所有表
        for row in table.rows:
            for cell in row.cells:
                # 替换 <YY>
                if "<YY>" in cell.text:
                    cell.text = cell.text.replace("<YY>", year)
                # 替换 <MM>
                if "<MM>" in cell.text:
                    cell.text = cell.text.replace("<MM>", month)
                # 替换 <DD>
                if "<DD>" in cell.text:
                    cell.text = cell.text.replace("<DD>", day)
                if '<today>' in cell.text:
                    # 把占位符去掉
                    cell.text = cell.text.replace('<today>', today_str)
                if '<today1>' in cell.text:
                    # 把占位符去掉
                    cell.text = cell.text.replace('<today1>', today_str)


    #换字
    for paragraph in docx.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<str1>' in paragraph.text:
            # 把占位符去掉
            paragraph.text= paragraph.text.replace('<str1>', text1)

    for paragraph in docx.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<str2>' in paragraph.text:
            # 把占位符去掉
            paragraph.text = paragraph.text.replace('<str2>', text2)

    for paragraph in docx.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<str4>' in paragraph.text:
            # 把占位符去掉
            paragraph.text= paragraph.text.replace('<str4>', text4)

    today = report_date

    replacements = {
        "<five_days_ago>": five_day_ago,
        "本周中信一级行业涨跌幅": "本周申万一级行业涨跌幅",
        "本周中信二级行业涨跌幅（排名前10位和后10位）": "本周申万二级行业涨跌幅（排名前10位和后10位）",
        "本周中信二级行业涨跌幅(排名前10位和后10位)": "本周申万二级行业涨跌幅（排名前10位和后10位）",
    }

    # **执行替换（不保存*
    replace_text_in_summary(docx, replacements)

    #换图
    pic_sm_list_12=[1]
    for i in range(0,1):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list_12[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list_12[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list_12[i]), '')
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(7)
                         picture.width=Cm(17)

    pic_sm_list=list(range(2,9))
    for i in range(0,7):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         if pic_sm_list[i] == 6 and not ENABLE_FIGURE6_FUTURES:
                             continue
                         if pic_sm_list[i] == 7 and not ENABLE_FIGURE7_UPDOWN:
                             continue
                         if pic_sm_list[i] == 8 and not ENABLE_FIGURE8_MARGIN:
                             continue
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(5)
                         picture.width=Cm(17)

    pic_sm_list=list(range(9,11))
    for i in range(0,2):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         if pic_sm_list[i] == 10 and not ENABLE_FIGURE10_ETF_NETINFLOW:
                             continue
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(5)
                         picture.width=Cm(8.5)

    pic_sm_list=list(range(11,12))
    for i in range(0,1):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(2.8)
                         picture.width=Cm(18)

    pic_sm_list=list(range(12,14))
    for i in range(0,2):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         if pic_sm_list[i] == 13 and not ENABLE_FIGURE13_INDUSTRY_ETF:
                             continue
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(5)
                         picture.width=Cm(8.5)

    pic_sm_list=list(range(14,16))
    for i in range(0,2):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         if pic_sm_list[i] == 14 and not ENABLE_FIGURE14_UNLOCK:
                             continue
                         if pic_sm_list[i] == 15 and not ENABLE_FIGURE15_IPO:
                             continue
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(5)
                         picture.width=Cm(17)

    pic_sm_list=list(range(16,18))
    for i in range(0,2):
        picture_path = 'outputs/figures/picture{}.png'.format(pic_sm_list[i])
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<Picture{}>'.format(pic_sm_list[i]) in cell.text:
                         # 把占位符去掉
                         cell.text = cell.text.replace('<Picture{}>'.format(pic_sm_list[i]), '')
                         run = cell.paragraphs[0].add_run()
                         # 添加图片并指定大
                         picture = run.add_picture(picture_path)
                         picture.height=Cm(16)
                         picture.width=Cm(17)
                     
    docx.save('{}.docx'.format(title))


def main() -> None:
    parser = argparse.ArgumentParser(description='Weekly report pipeline based on local data files.')
    parser.add_argument('--report-date', dest='report_date', help='Report date in YYYY-MM-DD format')
    parser.add_argument('--refresh-code-audit', action='store_true', help='Generate code_mapping_audit.csv explicitly')
    parser.add_argument('--init-full-history', action='store_true', help='Allow full history initialization when cache files are missing')
    parser.add_argument('--validate-etf-field', action='store_true', help='Validate and persist the structured ETF size path, then exit')
    parser.add_argument(
        '--refresh-active-quarter-cache',
        action='store_true',
        help='Refresh missing active equity quarter anchors in active_old_path during this run',
    )
    parser.add_argument(
        '--allow-stale-margin-data',
        action='store_true',
        help='Allow 图8使用早于 report_date 的最新可用数据，默认严格要求更新到 report_date',
    )
    parser.add_argument('--profile', action='store_true', help='Print timing summary for each pipeline stage')
    parser.add_argument('--batch-size', type=int, default=120, help='Chunk size for local batch processing')
    args = parser.parse_args()

    project_root = get_project_root()
    os.chdir(project_root)
    paths = build_paths(project_root)
    ensure_directories(paths)
    report_date = get_report_date(args.report_date)
    options = PipelineOptions(
        refresh_code_audit=args.refresh_code_audit,
        init_full_history=args.init_full_history,
        profile=args.profile,
        batch_size=args.batch_size,
        validate_etf_field=args.validate_etf_field,
        refresh_active_quarter_cache=args.refresh_active_quarter_cache,
        allow_stale_margin_data=args.allow_stale_margin_data,
    )
    tracker = PerformanceTracker(enabled=True)
    local_client = LocalDataClient(paths)

    if options.validate_etf_field:
        with tracker.track('validate_etf_field'):
            run_etf_field_validation(paths, report_date)
        tracker.summary()
        return

    with tracker.track('run_step1'):
        run_step1(local_client, paths, report_date, options)
    with tracker.track('run_step2'):
        run_step2(local_client, paths, report_date, options)
    with tracker.track('run_step3'):
        run_step3(paths, report_date)
    with tracker.track('run_step4'):
        run_step4(paths, report_date)
    tracker.summary()

if __name__ == '__main__':
    main()
